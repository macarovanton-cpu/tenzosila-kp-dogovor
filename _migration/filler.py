"""
filler.py — подстановка данных в Word-шаблон с плейсхолдерами.
Работает напрямую с python-docx, форматирование сохраняется 100%.
"""

import re
from docx import Document
from docx.oxml.ns import qn


def merge_runs(paragraph) -> None:
    """
    Склеивает соседние runs с одинаковым форматированием в одном параграфе.
    Нужно потому что Word иногда разбивает текст на несколько runs,
    из-за чего плейсхолдер {{КЛЮЧ}} может быть разбит на {{, КЛЮЧ, }}.
    """
    runs = paragraph.runs
    if len(runs) < 2:
        return
    
    i = 0
    while i < len(runs) - 1:
        curr = runs[i]
        next_r = runs[i + 1]
        
        # Склеиваем если форматирование совпадает
        curr_rpr = curr._r.find(qn('w:rPr'))
        next_rpr = next_r._r.find(qn('w:rPr'))
        
        curr_xml = '' if curr_rpr is None else curr_rpr.xml
        next_xml = '' if next_rpr is None else next_rpr.xml
        
        if curr_xml == next_xml:
            curr.text += next_r.text
            next_r._r.getparent().remove(next_r._r)
            # После удаления обновляем список
            runs = paragraph.runs
        else:
            i += 1


def replace_in_paragraph(paragraph, data: dict) -> None:
    """
    Заменяет все плейсхолдеры {{КЛЮЧ}} в параграфе на значения из data.
    Сначала склеивает runs, потом делает замену.
    """
    merge_runs(paragraph)
    
    for run in paragraph.runs:
        text = run.text
        if '{{' in text:
            for key, value in data.items():
                placeholder = f'{{{{{key}}}}}'
                if placeholder in text:
                    text = text.replace(placeholder, str(value) if value else '')
            run.text = text


def fill_template(template_path: str, data: dict, output_path: str) -> None:
    """
    Главная функция: открывает шаблон .docx, подставляет все плейсхолдеры,
    сохраняет результат в output_path.
    
    data — плоский словарь {КЛЮЧ: значение}.
    Если передан словарь с 'requisites' и 'specification' — автоматически объединяет.
    """
    # Если data содержит вложенные блоки — объединяем в плоский словарь
    if 'requisites' in data or 'specification' in data:
        flat = {}
        flat.update(data.get('requisites', {}))
        flat.update(data.get('specification', {}))
        data = flat
    
    doc = Document(template_path)
    
    # Обрабатываем все параграфы документа
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, data)
    
    # Обрабатываем все таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, data)
    
    # Обрабатываем колонтитулы (header/footer)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_in_paragraph(paragraph, data)
        for paragraph in section.footer.paragraphs:
            replace_in_paragraph(paragraph, data)
    
    doc.save(output_path)


def get_unfilled_placeholders(docx_path: str) -> list[str]:
    """
    Возвращает список незаполненных плейсхолдеров в документе.
    Удобно для проверки после генерации.
    """
    doc = Document(docx_path)
    found = set()
    
    def scan_text(text):
        for match in re.finditer(r'\{\{[^}]+\}\}', text):
            found.add(match.group())
    
    for p in doc.paragraphs:
        scan_text(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    scan_text(p.text)
    
    return sorted(found)


def remove_empty_paragraphs(doc_path: str, output_path: str = None) -> None:
    """
    Удаляет параграфы в которых после подстановки ничего не осталось.
    Используется для удаления лишних пунктов оплаты (если их меньше 6).
    """
    from lxml import etree

    doc = Document(doc_path)
    output_path = output_path or doc_path

    for paragraph in doc.paragraphs:
        # Удаляем параграфы которые полностью пустые после подстановки
        if paragraph.text.strip() == '':
            # Проверяем что это не структурный пустой параграф
            # (между разделами), а параграф-плейсхолдер оплаты
            p = paragraph._element
            parent = p.getparent()
            # Удаляем только если родитель — body (не таблица)
            if parent.tag.endswith('}body'):
                # Дополнительная проверка: параграф имеет нумерацию (список)
                ppr = p.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                if ppr is not None:
                    parent.remove(p)

    doc.save(output_path)
