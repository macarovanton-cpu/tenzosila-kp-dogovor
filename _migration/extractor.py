"""
extractor.py — извлечение текста из PDF/DOCX и данных через Gemini API
"""

import json
import re
import pdfplumber
from docx import Document
from google import genai
from google.genai import types
from config import GEMINI_API_KEY, PROMPT_PATH


def extract_pdf_text(pdf_path: str, pages: list[int] = None) -> str:
    """
    Извлекает текст из PDF.
    pages — список номеров страниц (с 1). Если None — все страницы.
    """
    with pdfplumber.open(pdf_path) as pdf:
        total = len(pdf.pages)
        if pages is None:
            pages = list(range(1, total + 1))
        
        parts = []
        for page_num in pages:
            if 1 <= page_num <= total:
                text = pdf.pages[page_num - 1].extract_text()
                if text:
                    parts.append(f"=== СТРАНИЦА {page_num} ===\n{text}")
        
        return "\n\n".join(parts)


def extract_kp_text(pdf_path: str) -> str:
    """
    Извлекает нужные страницы из КП:
    - Страница 3: технические характеристики
    - Страница 4: спецификация и условия оплаты
    - Страница 5: (если есть ОРИОН — третья спецификация)
    Если страниц меньше — берёт что есть.
    """
    with pdfplumber.open(pdf_path) as pdf:
        total = len(pdf.pages)
        # Берём страницы 3–5, но не больше чем есть
        target_pages = [p for p in [3, 4, 5] if p <= total]
        
        parts = []
        for page_num in target_pages:
            text = pdf.pages[page_num - 1].extract_text()
            if text:
                parts.append(f"=== СТРАНИЦА {page_num} ===\n{text}")
        
        return "\n\n".join(parts)


def extract_docx_text(docx_path: str) -> str:
    """
    Извлекает текст из Word-документа (карточка контрагента).
    Обрабатывает таблицы и параграфы.
    """
    doc = Document(docx_path)
    parts = []
    
    # Таблицы (основная часть карточки)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            # Убираем дубли (merged cells) и пустые
            unique = []
            seen = set()
            for c in cells:
                if c and c not in seen:
                    unique.append(c)
                    seen.add(c)
            if unique:
                parts.append(" | ".join(unique))
    
    # Отдельные параграфы (если есть)
    for p in doc.paragraphs:
        t = p.text.strip()
        if t and t not in parts:
            parts.append(t)
    
    return "\n".join(parts)


def extract_data_via_ai(kp_text: str, card_text: str) -> dict:
    """
    Отправляет тексты в Gemini API, получает JSON с данными для договора.
    Возвращает dict с ключами 'requisites' и 'specification'.
    """
    # Читаем промпт из файла
    with open(PROMPT_PATH, 'r', encoding='utf-8') as f:
        system_prompt = f.read()
    
    # Настраиваем Gemini (с поддержкой прокси для работы из РФ)
    from config import PROXY_URL
    if PROXY_URL:
        import os
        os.environ['HTTPS_PROXY'] = PROXY_URL
        os.environ['HTTP_PROXY']  = PROXY_URL

    client = genai.Client(api_key=GEMINI_API_KEY)

    user_message = (
        f"КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ:\n{kp_text}"
        f"\n\n---\n\nКАРТОЧКА КОНТРАГЕНТА:\n{card_text}"
    )

    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=user_message,
        config=types.GenerateContentConfig(
            system_instruction=system_prompt,
        )
    )
    raw = response.text.strip()
    
    # Убираем markdown-обёртку если есть
    raw = re.sub(r'^```json\s*', '', raw)
    raw = re.sub(r'\s*```$', '', raw)
    raw = raw.strip()
    
    data = json.loads(raw)
    return data


def extract_from_files(kp_path: str, card_path: str) -> dict:
    """
    Главная функция: принимает пути к файлам, возвращает dict с данными.
    card_path может быть .docx или .pdf
    """
    # Извлекаем текст КП
    kp_text = extract_kp_text(kp_path)
    
    # Извлекаем текст карточки (поддерживаем оба формата)
    if card_path.lower().endswith('.pdf'):
        card_text = extract_pdf_text(card_path)
    else:
        card_text = extract_docx_text(card_path)
    
    # Получаем данные через AI
    data = extract_data_via_ai(kp_text, card_text)
    return data
