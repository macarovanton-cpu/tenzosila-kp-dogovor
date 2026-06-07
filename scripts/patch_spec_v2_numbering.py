"""Одноразовый патч нумерации заголовков ТТХ/комплекта в spec_v2.docx."""
from __future__ import annotations

import shutil
from dataclasses import dataclass
from pathlib import Path

from docx import Document


TEMPLATE_PATH = Path("templates/contracts/spec_v2.docx")
BACKUP_PATH = TEMPLATE_PATH.with_suffix(".docx.bak")


@dataclass(frozen=True)
class PatchTarget:
    label: str
    needle: str
    placeholder: str


TARGETS = (
    PatchTarget(
        label="tth_heading",
        needle="Технические характеристики весов",
        placeholder="{{ТТХ_НОМЕР}}",
    ),
    PatchTarget(
        label="kit_heading",
        needle="Комплект поставки весов",
        placeholder="{{КОМПЛЕКТ_НОМЕР}}",
    ),
)


def _remove_num_pr(paragraph) -> bool:
    p_pr = paragraph._p.pPr
    if p_pr is None or p_pr.numPr is None:
        return False
    p_pr.remove(p_pr.numPr)
    return True


def _prepend_placeholder(paragraph, placeholder: str) -> bool:
    prefix = f"{placeholder}. "
    if paragraph.text.startswith(prefix) or prefix in paragraph.text:
        return False

    if not paragraph.runs:
        paragraph.add_run(prefix)
        return True

    for run in paragraph.runs:
        if run.text:
            run.text = prefix + run.text.lstrip()
            return True

    paragraph.runs[0].text = prefix + paragraph.runs[0].text
    return True


def _patch_target(document, target: PatchTarget) -> tuple[bool, bool, str, str]:
    for index, paragraph in enumerate(document.paragraphs):
        before = paragraph.text
        if target.needle not in before:
            continue

        num_pr_removed = _remove_num_pr(paragraph)
        text_changed = _prepend_placeholder(paragraph, target.placeholder)
        after = paragraph.text
        return num_pr_removed, text_changed, f"#{index}", after

    raise RuntimeError(f"Параграф не найден: {target.needle}")


def main() -> None:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(TEMPLATE_PATH)

    if BACKUP_PATH.exists():
        backup_status = f"backup_exists: {BACKUP_PATH}"
    else:
        shutil.copy2(TEMPLATE_PATH, BACKUP_PATH)
        backup_status = f"backup_created: {BACKUP_PATH}"

    document = Document(TEMPLATE_PATH)
    print(backup_status)

    for target in TARGETS:
        num_pr_removed, text_changed, paragraph_ref, after = _patch_target(document, target)
        print(
            f"{target.label}: found {paragraph_ref}; "
            f"numPr_removed={num_pr_removed}; "
            f"text_changed={text_changed}; "
            f"text={after!r}"
        )

    document.save(TEMPLATE_PATH)
    print(f"saved: {TEMPLATE_PATH}")


if __name__ == "__main__":
    main()
