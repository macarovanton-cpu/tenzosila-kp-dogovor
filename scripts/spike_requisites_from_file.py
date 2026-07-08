# -*- coding: utf-8 -*-
"""Feasibility-spike: реквизиты из файла (DOCX/XLSX/PDF-текст).

Одноразовый разведочный харнесс, НЕ часть приложения. Извлекает текст из
файлов и гонит через существующий пайплайн parse -> derive -> validate.
Выход: печать полей/ошибок по файлу + raw-дампы в outputs/spike_requisites/raw/.

Запуск: python scripts/spike_requisites_from_file.py [папка_с_образцами]
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

# Импорт из src — запускаем из корня репозитория
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from src.contracts.requisites_parser import parse_requisites
from src.contracts.requisites_transforms import derive_requisites
from src.contracts.requisites_validation import validate_requisites

DEFAULT_SAMPLES = Path("scripts/spike_data/requisites_samples")
RAW_DIR = Path("outputs/spike_requisites/raw")

# Порог детекта скана: меньше — считаем, что текстового слоя нет
MIN_NONSPACE_CHARS = 50

NO_TEXT_LAYER = "NO_TEXT_LAYER"
SKIP_UNSUPPORTED = "SKIP_UNSUPPORTED"

HARD_FIELDS = [
    "ЗАКАЗЧИК_КРАТКОЕ_НАИМЕНОВАНИЕ",
    "ЗАКАЗЧИК_ИНН",
    "ЗАКАЗЧИК_КПП",
    "ЗАКАЗЧИК_ОГРН",
    "ЗАКАЗЧИК_РС",
    "ЗАКАЗЧИК_КС",
    "ЗАКАЗЧИК_БИК",
]


# ---------------------------------------------------------------------------
# Извлечение текста по форматам
# ---------------------------------------------------------------------------

def _docx_table_lines(table) -> list[str]:
    """Строки таблицы DOCX: ячейки строки склеиваем пробелом.

    Вложенные таблицы обходим одним уровнем (частый случай в картах партнёра).
    """
    lines: list[str] = []
    for row in table.rows:
        parts: list[str] = []
        for cell in row.cells:
            cell_text = " ".join(
                p.text.strip() for p in cell.paragraphs if p.text.strip()
            )
            if cell_text:
                parts.append(cell_text)
            for nested in cell.tables:
                lines.extend(_docx_table_lines(nested))
        if parts:
            lines.append(" ".join(parts))
    return lines


def extract_docx(path: Path) -> str:
    import docx

    doc = docx.Document(str(path))
    lines: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())
    for table in doc.tables:
        lines.extend(_docx_table_lines(table))
    return "\n".join(lines)


def extract_xlsx(path: Path) -> str:
    """Непустые ячейки строки склеиваем пробелом по порядку столбцов.

    Merged cells: openpyxl кладёт значение в верхнюю-левую ячейку диапазона,
    остальные — None, т.е. отваливаются сами.
    """
    import openpyxl

    wb = openpyxl.load_workbook(str(path), data_only=True)
    lines: list[str] = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            parts = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if parts:
                lines.append(" ".join(parts))
    return "\n".join(lines)


def extract_pdf(path: Path, sort: bool) -> str:
    import fitz

    with fitz.open(str(path)) as doc:
        pages = [page.get_text(sort=sort) for page in doc]
    text = "\n".join(pages)
    if len("".join(text.split())) < MIN_NONSPACE_CHARS:
        return NO_TEXT_LAYER
    return text


def extract_variants(path: Path) -> dict[str, str]:
    """Вернуть {метка_режима: текст_или_маркер}. PDF даёт два режима."""
    ext = path.suffix.lower()
    if ext == ".docx":
        return {"docx": extract_docx(path)}
    if ext == ".xlsx":
        return {"xlsx": extract_xlsx(path)}
    if ext == ".pdf":
        return {
            "pdf_block": extract_pdf(path, sort=False),
            "pdf_sorted": extract_pdf(path, sort=True),
        }
    return {ext or "noext": SKIP_UNSUPPORTED}


# ---------------------------------------------------------------------------
# Прогон
# ---------------------------------------------------------------------------

def run_pipeline(text: str) -> dict:
    fields = parse_requisites(text)
    derived, derive_warnings = derive_requisites(fields)
    merged = {**derived, **fields}  # derived не перетирает распознанное
    errors, warnings = validate_requisites(merged)
    return {
        "fields": merged,
        "derive_warnings": derive_warnings,
        "errors": errors,
        "warnings": warnings,
    }


def main() -> None:
    samples = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_SAMPLES
    if not samples.is_dir():
        sys.exit(f"Папка не найдена: {samples}")
    RAW_DIR.mkdir(parents=True, exist_ok=True)

    results: dict[str, dict] = {}
    for path in sorted(samples.iterdir()):
        if not path.is_file():
            continue
        print(f"\n{'=' * 78}\nФАЙЛ: {path.name}")
        for mode, text in extract_variants(path).items():
            tag = f"{path.stem}__{mode}"
            print(f"\n--- режим: {mode} ---")
            if text in (NO_TEXT_LAYER, SKIP_UNSUPPORTED):
                print(f"  {text}")
                results[tag] = {"marker": text}
                continue
            (RAW_DIR / f"{tag}.txt").write_text(text, encoding="utf-8")
            res = run_pipeline(text)
            results[tag] = res
            print(f"  извлечено символов: {len(text)}")
            for key in HARD_FIELDS:
                print(f"  {key}: {res['fields'].get(key, '<пусто>')}")
            extra = {
                k: v for k, v in res["fields"].items() if k not in HARD_FIELDS
            }
            for k, v in extra.items():
                print(f"  (soft) {k}: {v}")
            for e in res["errors"]:
                print(f"  ERROR: {e}")
            for w in res["warnings"] + res["derive_warnings"]:
                print(f"  warn:  {w}")

    out = RAW_DIR.parent / "results.json"
    out.write_text(
        json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(f"\nИтоги сохранены: {out}")


if __name__ == "__main__":
    main()
