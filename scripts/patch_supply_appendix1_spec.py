"""Разовый хирургический патч spec-таблицы supply_appendix_1.docx.

Делает строку товара динамической: вместо статичных {{ТОВАР_НАИМЕНОВАНИЕ}} /
{{СУММА_ЦИФРАМИ}} ставит docxtpl row-цикл по spec_rows — весы и доставка
отдельными строками. Строку «Итого» ({{СУММА_ЦИФРАМИ}}) НЕ трогает.

Структура цикла — как в починенном supply_appendix_2.docx (3 строки): отдельная
строка-{%tr for%}, строка-данные, отдельная строка-{%tr endfor%}. Однострочный
вариант (for и endfor в одной строке) docxtpl рендерит с ошибкой — см. коммит
«fix: починить {%tr for/endfor %} в supply_appendix_2.docx».

Запуск: python scripts/patch_supply_appendix1_spec.py
"""
import copy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

DST_APP1 = Path("templates/contracts/supply_appendix_1.docx")


def _set_cell_text(cell, text: str) -> None:
    """Очистить ячейку и записать текст одним чистым run в первом параграфе."""
    paras = cell.paragraphs
    if not paras:
        return
    runs = paras[0].runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        paras[0].add_run(text)
    for p in paras[1:]:
        p._element.getparent().remove(p._element)


def patch_appendix1_spec() -> None:
    doc = Document(str(DST_APP1))
    spec = doc.tables[0]

    idx = next(
        (i for i, r in enumerate(spec.rows) if "ТОВАР_НАИМЕНОВАНИЕ" in r.cells[0].text),
        None,
    )
    if idx is None:
        raise ValueError("Не найдена строка товара ({{ТОВАР_НАИМЕНОВАНИЕ}}) в tables[0]")

    # Из одной строки делаем три: for-строка, строка-данные, endfor-строка.
    data_tr = spec.rows[idx]._tr
    for_tr = copy.deepcopy(data_tr)
    endfor_tr = copy.deepcopy(data_tr)
    data_tr.addprevious(for_tr)
    data_tr.addnext(endfor_tr)

    rows = spec.rows  # перечитать после вставки
    _set_cell_text(rows[idx].cells[0], "{%tr for row in spec_rows %}")
    _set_cell_text(rows[idx].cells[1], "")
    _set_cell_text(rows[idx + 1].cells[0], "{{row.name}}")
    _set_cell_text(rows[idx + 1].cells[1], "{{row.sum}}")
    _set_cell_text(rows[idx + 2].cells[0], "")
    _set_cell_text(rows[idx + 2].cells[1], "{%tr endfor %}")

    doc.save(str(DST_APP1))
    print(f"Пропатчен: {DST_APP1} (spec-цикл из 3 строк, idx={idx})")


if __name__ == "__main__":
    patch_appendix1_spec()
