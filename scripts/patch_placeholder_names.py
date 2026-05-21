"""
Патч плейсхолдеров наименования заказчика:
  ЗАКАЗЧИК_КРАТКОЕ_НАИМЕНОВАНИЕ → ЗАКАЗЧИК_ПОЛНОЕ_НАИМЕНОВАНИЕ
  в преамбуле, реквизитах и блоке подписей.

Idempotent: повторный запуск безопасен.
Запускать из корня проекта:
    python scripts/patch_placeholder_names.py
"""
import sys
sys.stdout.reconfigure(encoding="utf-8")
from pathlib import Path
import shutil

from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

OLD = "ЗАКАЗЧИК_КРАТКОЕ_НАИМЕНОВАНИЕ"
NEW = "ЗАКАЗЧИК_ПОЛНОЕ_НАИМЕНОВАНИЕ"
OLD_PH = f"{{{{{OLD}}}}}"
NEW_PH = f"{{{{{NEW}}}}}"

CONTRACTS = Path("templates/contracts")
BACKUP_DIR = Path("templates/contracts/backup")


def _merge_runs(para: Paragraph) -> None:
    """Склеивает соседние runs с одинаковым rPr (idempotent)."""
    runs = para.runs
    if len(runs) < 2:
        return
    i = 0
    while i < len(runs) - 1:
        curr, nxt = runs[i], runs[i + 1]
        curr_rpr = curr._r.find(qn("w:rPr"))
        nxt_rpr = nxt._r.find(qn("w:rPr"))
        cx = "" if curr_rpr is None else curr_rpr.xml
        nx = "" if nxt_rpr is None else nxt_rpr.xml
        if cx == nx:
            curr.text += nxt.text
            nxt._r.getparent().remove(nxt._r)
            runs = para.runs
        else:
            i += 1


def _replace_in_para(para: Paragraph) -> bool:
    """Заменяет OLD_PH → NEW_PH в runs параграфа. Возвращает True если изменено."""
    _merge_runs(para)
    changed = False
    for run in para.runs:
        if OLD_PH in run.text:
            run.text = run.text.replace(OLD_PH, NEW_PH)
            changed = True
    return changed


def _body_paras(doc: DocxDocument) -> list[Paragraph]:
    body_tag = qn("w:body")
    return [p for p in doc.paragraphs if p._p.getparent().tag == body_tag]


def patch_contract(path: Path) -> None:
    doc = Document(path)
    count = 0

    # Преамбула: body para содержащий «Заказчик»
    for p in _body_paras(doc):
        if "Заказчик" in p.text and OLD_PH in p.text:
            if _replace_in_para(p):
                count += 1
                print(f"  contract преамбула: {p.text[:100]!r}")

    # Реквизиты: TABLE[0] row[2] cell[1]
    if len(doc.tables) > 0:
        if len(doc.tables[0].rows) > 2:
            req_cell = doc.tables[0].rows[2].cells[1]
            for p in req_cell.paragraphs:
                if _replace_in_para(p):
                    count += 1
                    print(f"  contract реквизиты: {p.text[:80]!r}")
        else:
            print("WARNING: TABLE[0] имеет менее 3 строк — реквизиты не обработаны")

    # Подписи: TABLE[1] row[0] cell[1]
    if len(doc.tables) > 1:
        sig_cell = doc.tables[1].rows[0].cells[1]
        for p in sig_cell.paragraphs:
            if _replace_in_para(p):
                count += 1
                print(f"  contract подписи: {p.text[:80]!r}")

    doc.save(path)
    print(f"  Сохранено: {path}. Заменено: {count} вхождений.")


def patch_spec(path: Path) -> None:
    doc = Document(path)
    count = 0

    # Преамбула: body para содержащий «Подрядчик»
    for p in _body_paras(doc):
        if "Подрядчик" in p.text and OLD_PH in p.text:
            if _replace_in_para(p):
                count += 1
                print(f"  spec преамбула: {p.text[:100]!r}")

    # Подписи: TABLE[3] row[0] cell[1]
    if len(doc.tables) > 3:
        sig_cell = doc.tables[3].rows[0].cells[1]
        for p in sig_cell.paragraphs:
            if _replace_in_para(p):
                count += 1
                print(f"  spec подписи: {p.text[:80]!r}")

    doc.save(path)
    print(f"  Сохранено: {path}. Заменено: {count} вхождений.")


def main() -> None:
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)

    contract = CONTRACTS / "contract.docx"
    spec = CONTRACTS / "spec_foundation_install.docx"

    shutil.copy2(contract, BACKUP_DIR / "contract.docx")
    shutil.copy2(spec, BACKUP_DIR / "spec_foundation_install.docx")
    print(f"Бэкапы: {BACKUP_DIR}/")

    print("\nPatch contract.docx:")
    patch_contract(contract)

    print("\nPatch spec_foundation_install.docx:")
    patch_spec(spec)

    print("\nГотово. Проверь: python -m pytest tests/contracts/test_templates.py -v")


if __name__ == "__main__":
    main()
