"""Идемпотентный патч шаблона spec_v2.docx.

Заменяет захардкоженные значения на динамические маркеры/плейсхолдеры.
Запуск: python scripts/patch_spec_v2_template.py
"""
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


TEMPLATE = Path(__file__).resolve().parent.parent / "templates" / "contracts" / "spec_v2.docx"


def _find_para(doc, substring: str):
    """Первый параграф, содержащий substring."""
    for p in doc.paragraphs:
        if substring in p.text:
            return p
    return None


def _make_marker_para(text: str):
    """Создать XML-элемент параграфа с текстом-маркером."""
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    p.append(r)
    return p


def _set_cell_text_plain(cell, text: str) -> None:
    """Заменить весь текст ячейки, сохраняя форматирование первого run."""
    paras = cell.paragraphs
    if not paras:
        return
    for p in paras[1:]:
        p._element.getparent().remove(p._element)
    runs = paras[0].runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        paras[0].add_run(text)


def fix_quotes(doc) -> int:
    """Кавычки: «ТПК «Тензосила» → «ТПК «Тензосила»» (двойная закрывающая).

    Текст часто разбит на runs: Run[i]='Тензосила', Run[i+1]='»'.
    """
    count = 0

    def _fix_in_paragraphs(paragraphs):
        nonlocal count
        for para in paragraphs:
            runs = para.runs
            for i, run in enumerate(runs):
                if '«Тензосила»' in run.text and '«Тензосила»»' not in run.text:
                    run.text = run.text.replace('«Тензосила»', '«Тензосила»»')
                    count += 1
                elif run.text.rstrip() == 'Тензосила' and i + 1 < len(runs):
                    nxt = runs[i + 1]
                    if nxt.text.startswith('»') and not nxt.text.startswith('»»'):
                        nxt.text = '»' + nxt.text
                        count += 1

    _fix_in_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _fix_in_paragraphs(cell.paragraphs)
    return count


def patch_payment(doc) -> bool:
    """Заменить 6 слотов {{СПЕЦ_ОПЛАТА_П1..П6}} на единый {{PAYMENT_SECTION}}."""
    if _find_para(doc, '{{PAYMENT_SECTION}}'):
        return False
    header = _find_para(doc, 'Порядок оплаты')
    if not header:
        return False
    body = doc.element.body
    to_remove = [p._element for p in doc.paragraphs if '{{СПЕЦ_ОПЛАТА_П' in p.text]
    for el in to_remove:
        body.remove(el)
    header._element.addnext(_make_marker_para('{{PAYMENT_SECTION}}'))
    return True


def patch_terms(doc) -> bool:
    """Заменить 3 строки сроков на единый {{TERMS_SECTION}}."""
    if _find_para(doc, '{{TERMS_SECTION}}'):
        return False
    header = _find_para(doc, 'Срок поставки Весов')
    if not header:
        return False
    body = doc.element.body
    markers = ['СПЕЦ_СРОК_ПОСТАВКИ', 'СПЕЦ_СРОК_ФУНДАМЕНТ', 'СПЕЦ_СРОК_МОНТАЖ']
    to_remove = []
    for p in doc.paragraphs:
        if any(m in p.text for m in markers):
            to_remove.append(p._element)
    # Пустые параграфы между строками сроков (сразу после каждой)
    all_p = list(body.iterchildren(qn('w:p')))
    for el in to_remove:
        idx = all_p.index(el)
        if idx + 1 < len(all_p):
            next_el = all_p[idx + 1]
            next_text = ''.join(t.text or '' for t in next_el.findall('.//' + qn('w:t')))
            if next_text.strip() == '' and next_el not in to_remove:
                to_remove.append(next_el)
    for el in to_remove:
        body.remove(el)
    header._element.addnext(_make_marker_para('{{TERMS_SECTION}}'))
    return True


def patch_tth(doc) -> bool:
    """ТТХ: захардкоженные значения → плейсхолдеры."""
    table = doc.tables[1]
    replacements = {
        2: ('11', '{{ТТХ_НАГРУЗКА_НА_ОСЬ}}'),
        3: ('не более 50 м', '{{ТТХ_РАССТОЯНИЕ_ДО_ТЕРМИНАЛА}}'),
        # row 4 (дискретность) имеет 2 параграфа — управляется patch_tth_discreteness
        5: (None, '{{ТТХ_ГАБАРИТЫ}}'),
        6: (None, '{{ТТХ_ТЕМПЕРАТУРА}}'),
    }
    changed = False
    for row_idx, (old, new) in replacements.items():
        cell = table.rows[row_idx].cells[2]
        if new in cell.text:
            continue
        _set_cell_text_plain(cell, new)
        changed = True
    return changed


def patch_kit(doc) -> bool:
    """Комплект: оставить header + 1 шаблонную строку, удалить rows 2-7."""
    table = doc.tables[2]
    if len(table.rows) <= 2:
        return False
    tbl = table._tbl
    all_trs = list(tbl.iterchildren(qn('w:tr')))
    for tr in all_trs[2:]:
        tbl.remove(tr)
    # Очистить текст шаблонной строки (Row 1)
    template_row = table.rows[1]
    for cell in template_row.cells:
        _set_cell_text_plain(cell, '')
    return True


def patch_year(doc) -> bool:
    """Год 2026 → {{ТЕКУЩИЙ_ГОД}} в таблице подписей."""
    table = doc.tables[3]
    changed = False
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if '2026' in run.text and '{{ТЕКУЩИЙ_ГОД}}' not in run.text:
                        run.text = run.text.replace('2026', '{{ТЕКУЩИЙ_ГОД}}')
                        changed = True
    return changed


def patch_appendix(doc) -> bool:
    """Первый {{СПЕЦ_НОМЕР}} в строке приложения → {{ПРИЛОЖЕНИЕ_НОМЕР}}."""
    if _find_para(doc, '{{ПРИЛОЖЕНИЕ_НОМЕР}}'):
        return False
    for para in doc.paragraphs:
        if 'Приложение' in para.text and '{{СПЕЦ_НОМЕР}}' in para.text:
            for run in para.runs:
                if '{{СПЕЦ_НОМЕР}}' in run.text:
                    run.text = run.text.replace('{{СПЕЦ_НОМЕР}}', '{{ПРИЛОЖЕНИЕ_НОМЕР}}', 1)
                    return True
    return False


def patch_tth_discreteness(doc) -> bool:
    """Восстановить 2-параграфную ячейку дискретности (если была затёрта старым патчем)."""
    cell = doc.tables[1].rows[4].cells[2]
    if "{{ТТХ_ДИСКРЕТНОСТЬ_1}}" in cell.text:
        return False
    # Очистить всё кроме первого параграфа
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    paras = cell.paragraphs
    if paras[0].runs:
        paras[0].runs[0].text = "{{ТТХ_ДИСКРЕТНОСТЬ_1}}"
        for r in paras[0].runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        paras[0].add_run("{{ТТХ_ДИСКРЕТНОСТЬ_1}}")
    p2 = OxmlElement('w:p')
    r2 = OxmlElement('w:r')
    t2 = OxmlElement('w:t')
    t2.text = "{{ТТХ_ДИСКРЕТНОСТЬ_2}}"
    r2.append(t2)
    p2.append(r2)
    cell._tc.append(p2)
    return True


def patch_tth_boundaries(doc) -> bool:
    """Заменить захардкоженные границы диапазонов в метке строки дискретности."""
    cell = doc.tables[1].rows[4].cells[1]
    if "{{ТТХ_ГРАНИЦА_1}}" in cell.text:
        return False
    if "дискретность" not in cell.text or "от 0 до" not in cell.text:
        print("ПРЕДУПРЕЖДЕНИЕ: структура шаблона изменилась, patch_tth_boundaries пропущен")
        return False
    paras = cell.paragraphs
    if len(paras) >= 2:
        run = paras[1].runs[0] if paras[1].runs else paras[1].add_run("")
        run.text = "- от 0 до {{ТТХ_ГРАНИЦА_1}}т"
    if len(paras) >= 3:
        run = paras[2].runs[0] if paras[2].runs else paras[2].add_run("")
        run.text = "- от {{ТТХ_ГРАНИЦА_1}} до {{ТТХ_ГРАНИЦА_2}}т"
    return True


def patch_foundation_check(doc) -> bool:
    """Вставить маркер {{APPENDIX_FOUNDATION_CHECK}} после строительного задания."""
    if _find_para(doc, '{{APPENDIX_FOUNDATION_CHECK}}'):
        return False
    target = _find_para(doc, 'Строительное задание на фундамент')
    if not target:
        return False
    target._element.addnext(_make_marker_para('{{APPENDIX_FOUNDATION_CHECK}}'))
    return True


def _zip_replace(path: str, old: str, new: str) -> int:
    """ZIP-уровневая замена текста в XML-файлах DOCX (text box-ы и пр.)."""
    tmp = path + '.tmp'
    count = 0
    with zipfile.ZipFile(path, 'r') as zin:
        with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.endswith('.xml'):
                    text = data.decode('utf-8')
                    if old in text:
                        text = text.replace(old, new)
                        count += 1
                    data = text.encode('utf-8')
                zout.writestr(item, data)
    shutil.move(tmp, path)
    return count


def main():
    doc = Document(str(TEMPLATE))

    n = fix_quotes(doc)
    print(f"Кавычки: {n} замен")

    print(f"Оплата → маркер: {patch_payment(doc)}")
    print(f"Сроки → маркер: {patch_terms(doc)}")
    print(f"ТТХ → плейсхолдеры: {patch_tth(doc)}")
    print(f"ТТХ дискретность → 2 параграфа: {patch_tth_discreteness(doc)}")
    print(f"Комплект → шаблон: {patch_kit(doc)}")
    print(f"Год → плейсхолдер: {patch_year(doc)}")
    print(f"Приложение → номер: {patch_appendix(doc)}")
    print(f"Контрольный лист → маркер: {patch_foundation_check(doc)}")
    print(f"ТТХ границы → плейсхолдеры: {patch_tth_boundaries(doc)}")

    doc.save(str(TEMPLATE))
    print(f"Сохранено: {TEMPLATE}")


if __name__ == "__main__":
    main()
