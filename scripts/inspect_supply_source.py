"""Инспекция docs/source/Автовесы _Поставка.docx.

Выводит параграфы и таблицы с индексами для поиска границ разделов.
Запуск: python scripts/inspect_supply_source.py
"""
from docx import Document
from pathlib import Path

SRC = Path("docs/source/Автовесы _Поставка.docx")
OUT = Path("docs/source/inspect_supply.txt")
doc = Document(str(SRC))

lines = []
lines.append("=" * 70)
lines.append("ПАРАГРАФЫ")
lines.append("=" * 70)
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        style = p.style.name if p.style else "?"
        lines.append(f"[{i:3d}] [{style[:20]:20s}] {t[:100]}")

lines.append("")
lines.append("=" * 70)
lines.append(f"ТАБЛИЦЫ: {len(doc.tables)} шт.")
lines.append("=" * 70)
for ti, table in enumerate(doc.tables):
    lines.append(f"\n--- Таблица {ti} ({len(table.rows)} строк x {len(table.columns)} колонок) ---")
    for ri, row in enumerate(table.rows):
        cells_text = []
        for ci, cell in enumerate(row.cells):
            ct = cell.text.strip()[:60]
            cells_text.append(f"[{ci}]{ct}")
        lines.append(f"  Row {ri}: {' | '.join(cells_text)}")

OUT.write_text("\n".join(lines), encoding="utf-8")
print(f"Готово: {OUT}")
