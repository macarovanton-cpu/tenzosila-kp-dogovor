"""Ядро конвертации PDF-чертежей фундаментов в DOCX."""
from __future__ import annotations

import copy
import io
import struct
from pathlib import Path
from typing import Any

import fitz
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent.parent
SRC_DIR = ROOT / "data" / "fundament" / "pdf_source"
OUT_DIR = ROOT / "data" / "fundament" / "build_task"
CS_SRC = SRC_DIR / "control_sheet"
CS_OUT = ROOT / "data" / "fundament" / "control_sheet"
REFERENCE = OUT_DIR / "пандусный_С_Ф_3скц.docx"

DPI = 200
EMU_PER_MM = 36_000
USABLE_W_EMU = 6_390_640
FIRST_PAGE_IMAGE_MAX_H_EMU = (297 - 10 - 35) * EMU_PER_MM
NEXT_PAGE_IMAGE_MAX_H_EMU = (297 - 10 - 5) * EMU_PER_MM
TEXTBOX_DOC_PR_BASE_ID = 500_000
TEXTBOX_ANCHOR_BASE_ID = 0x47FCA4F2
TEXTBOX_WIDTH_EMU = 100 * EMU_PER_MM
TEXTBOX_HEIGHT_EMU = 22 * EMU_PER_MM

WPS_TXBX_TAG = "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}txbx"
WP14_ANCHOR_ID_ATTR = (
    "{http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing}anchorId"
)


def _load_reference_parts(doc: Any) -> tuple[list[Any], Any, Any]:
    paragraphs = doc.paragraphs
    start_idx = None
    for idx, paragraph in enumerate(paragraphs):
        if "Приложение" in paragraph.text and "к Спецификации" in paragraph.text:
            start_idx = idx
            break
    if start_idx is None:
        raise ValueError("В эталоне не найден заголовок приложения")

    image_idx = None
    textbox_drawing = None
    for idx, paragraph in enumerate(paragraphs[start_idx + 1:], start=start_idx + 1):
        for drawing in paragraph._p.iter(qn("w:drawing")):
            if drawing.find(".//" + WPS_TXBX_TAG) is not None:
                image_idx = idx
                textbox_drawing = drawing
                break
        if image_idx is not None:
            break
    if image_idx is None or textbox_drawing is None:
        raise ValueError("В эталоне не найден параграф картинки с floating TextBox")

    header_block = [copy.deepcopy(p._p) for p in paragraphs[start_idx:image_idx]]
    image_para = paragraphs[image_idx]
    img_ppr = image_para._p.find(qn("w:pPr"))
    if img_ppr is None:
        raise ValueError("В параграфе картинки эталона нет w:pPr")

    return header_block, copy.deepcopy(img_ppr), copy.deepcopy(textbox_drawing)


def _render_pages(pdf_path: Path) -> list[bytes]:
    pages: list[bytes] = []
    with fitz.open(str(pdf_path)) as pdf:
        for page in pdf:
            pixmap = page.get_pixmap(dpi=DPI)
            pages.append(pixmap.tobytes("png"))
    return pages


def _png_size(data: bytes) -> tuple[int, int]:
    if not data.startswith(b"\x89PNG\r\n\x1a\n"):
        raise ValueError("Ожидался PNG")
    return struct.unpack(">II", data[16:24])


def _image_emu(px_w: int, px_h: int, page_idx: int) -> tuple[int, int]:
    if px_w <= 0 or px_h <= 0:
        raise ValueError("Размеры изображения должны быть положительными")

    aspect = px_w / px_h
    max_h = FIRST_PAGE_IMAGE_MAX_H_EMU if page_idx == 0 else NEXT_PAGE_IMAGE_MAX_H_EMU
    target_w = USABLE_W_EMU
    target_h = round(target_w / aspect)
    if target_h > max_h:
        target_h = max_h
        target_w = round(target_h * aspect)
    return target_w, target_h


def _clone_textbox(drawing: Any, page_idx: int) -> Any:
    clone = copy.deepcopy(drawing)
    doc_pr = clone.find(".//" + qn("wp:docPr"))
    if doc_pr is None:
        raise ValueError("В TextBox не найден wp:docPr")
    doc_pr.set("id", str(TEXTBOX_DOC_PR_BASE_ID + page_idx))

    anchor = clone.find(".//" + qn("wp:anchor"))
    if anchor is None:
        raise ValueError("В TextBox не найден wp:anchor")
    anchor.set(WP14_ANCHOR_ID_ATTR, f"{TEXTBOX_ANCHOR_BASE_ID + page_idx:08X}")

    for extent in clone.findall(".//" + qn("wp:extent")):
        extent.set("cx", str(TEXTBOX_WIDTH_EMU))
        extent.set("cy", str(TEXTBOX_HEIGHT_EMU))
    for extent in clone.findall(".//" + qn("a:ext")):
        extent.set("cx", str(TEXTBOX_WIDTH_EMU))
        extent.set("cy", str(TEXTBOX_HEIGHT_EMU))
    return clone


def _remove_page_breaks(paragraph_xml: Any) -> None:
    ppr = paragraph_xml.find(qn("w:pPr"))
    if ppr is not None:
        for page_break_before in ppr.findall(qn("w:pageBreakBefore")):
            ppr.remove(page_break_before)

    for br in list(paragraph_xml.iter(qn("w:br"))):
        if br.get(qn("w:type")) == "page":
            parent = br.getparent()
            if parent is not None:
                parent.remove(br)


def _clear_body_keep_section(doc: Any) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _insert_paragraph_xml(doc: Any, paragraph_xml: Any) -> None:
    sect_pr = doc.element.body.find(qn("w:sectPr"))
    if sect_pr is None:
        doc.element.body.append(paragraph_xml)
    else:
        sect_pr.addprevious(paragraph_xml)


def _add_empty_paragraph(doc: Any) -> Paragraph:
    paragraph_xml = OxmlElement("w:p")
    _insert_paragraph_xml(doc, paragraph_xml)
    return Paragraph(paragraph_xml, doc._body)


def _add_image_paragraph(
    doc: Any,
    img_ppr: Any,
    textbox_drawing: Any,
    png: bytes,
    page_idx: int,
) -> None:
    paragraph = _add_empty_paragraph(doc)
    paragraph._p.append(copy.deepcopy(img_ppr))
    if page_idx > 0:
        paragraph.add_run().add_break(WD_BREAK.PAGE)

    textbox_run = paragraph.add_run()
    textbox_run._r.append(_clone_textbox(textbox_drawing, page_idx))

    px_w, px_h = _png_size(png)
    width_emu, height_emu = _image_emu(px_w, px_h, page_idx)
    paragraph.add_run().add_picture(
        io.BytesIO(png),
        width=Emu(width_emu),
        height=Emu(height_emu),
    )


def convert(pdf_path: Path, out_path: Path) -> None:
    doc = Document(str(REFERENCE))
    header_block, img_ppr, textbox_drawing = _load_reference_parts(doc)
    pages = _render_pages(pdf_path)

    _clear_body_keep_section(doc)
    for paragraph_xml in header_block:
        paragraph_copy = copy.deepcopy(paragraph_xml)
        _remove_page_breaks(paragraph_copy)
        _insert_paragraph_xml(doc, paragraph_copy)
    for page_idx, png in enumerate(pages):
        _add_image_paragraph(doc, img_ppr, textbox_drawing, png, page_idx)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def _pdf_files(path: Path) -> list[Path]:
    if not path.exists():
        return []
    return sorted([*path.glob("*.pdf"), *path.glob("*.PDF")], key=lambda item: item.name.lower())
