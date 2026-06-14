"""Создание трёх DOCX-шаблонов договора поставки из образца.

Запуск: python scripts/create_supply_templates.py
"""
import re
import shutil
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = Path("docs/source/Автовесы _Поставка.docx")
DST_CONTRACT = Path("templates/contracts/supply_contract.docx")
DST_APP1 = Path("templates/contracts/supply_appendix_1.docx")
DST_APP2 = Path("templates/contracts/supply_appendix_2.docx")


# ─── helpers ────────────────────────────────────────────────────────────────

def _para_text(elem):
    return "".join(t.text or "" for t in elem.findall(".//" + qn("w:t")))


def trim_body_after(doc, marker: str):
    """Удалить все body-элементы начиная с параграфа, содержащего marker."""
    body = doc.element.body
    marker_el = None
    for child in body.iterchildren():
        if child.tag == qn("w:p") and marker in _para_text(child):
            marker_el = child
            break
    if marker_el is None:
        raise ValueError(f"Маркер не найден: {marker!r}")
    to_remove = []
    found = False
    for child in list(body.iterchildren()):
        if child is marker_el:
            found = True
        if found and child.tag != qn("w:sectPr"):
            to_remove.append(child)
    for el in to_remove:
        body.remove(el)


def trim_body_before(doc, marker: str):
    """Удалить все body-элементы ДО параграфа, содержащего marker."""
    body = doc.element.body
    marker_el = None
    for child in body.iterchildren():
        if child.tag == qn("w:p") and marker in _para_text(child):
            marker_el = child
            break
    if marker_el is None:
        raise ValueError(f"Маркер не найден: {marker!r}")
    to_remove = []
    for child in list(body.iterchildren()):
        if child is marker_el:
            break
        if child.tag != qn("w:sectPr"):
            to_remove.append(child)
    for el in to_remove:
        body.remove(el)


def merge_para_runs(para) -> str:
    """Слить все runs в первый, вернуть полный текст."""
    runs = para.runs
    if not runs:
        return ""
    text = "".join(r.text or "" for r in runs)
    runs[0].text = text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)
    return text


def replace_in_para(para, old: str, new: str) -> bool:
    text = merge_para_runs(para)
    if old in text:
        para.runs[0].text = text.replace(old, new)
        return True
    return False


def regex_replace_in_para(para, pattern: str, new: str) -> bool:
    text = merge_para_runs(para)
    result, n = re.subn(pattern, new, text)
    if n and para.runs:
        para.runs[0].text = result
    return bool(n)


def replace_in_cell(cell, old: str, new: str):
    for para in cell.paragraphs:
        replace_in_para(para, old, new)


def set_cell_text(cell, text: str):
    """Очистить ячейку и задать текст в первом параграфе."""
    paras = cell.paragraphs
    if not paras:
        return
    runs = paras[0].runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        paras[0].add_run(text)
    for p in paras[1:]:
        p._element.getparent().remove(p._element)


def find_para(doc, substring: str):
    for p in doc.paragraphs:
        if substring in p.text:
            return p
    return None


def _make_plain_para(text: str):
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p.append(r)
    return p


# ─── reusable: подписи ───────────────────────────────────────────────────────

def patch_signature_cell_sup(cell):
    """Заменить подпись поставщика (col 0)."""
    replace_in_cell(cell, "О. А. Сенаторов", "{{ПОСТАВЩИК_ДИРЕКТОР_ФИО}}")
    replace_in_cell(cell, "2026 г.", "{{ТЕКУЩИЙ_ГОД}} г.")


def patch_signature_cell_cus(cell):
    """Заменить подпись покупателя (col 2)."""
    for para in cell.paragraphs:
        t = para.text.strip()
        if t == "Глава":
            merge_para_runs(para)
            if para.runs:
                para.runs[0].text = "{{ПОКУПАТЕЛЬ_ДИРЕКТОР_ДОЛЖНОСТЬ}}"
        elif "КФХ" in t and "Молчанов" in t:
            merge_para_runs(para)
            if para.runs:
                para.runs[0].text = "{{ПОКУПАТЕЛЬ_НАИМЕНОВАНИЕ}}"
        elif "В.Г. Молчанов" in t:
            replace_in_para(para, "В.Г. Молчанов", "{{ПОКУПАТЕЛЬ_ДИРЕКТОР_ФИО}}")
        elif "2026 г." in t:
            replace_in_para(para, "2026 г.", "{{ТЕКУЩИЙ_ГОД}} г.")


# ─── supply_contract.docx ────────────────────────────────────────────────────

def create_supply_contract():
    shutil.copy(SRC, DST_CONTRACT)
    doc = Document(str(DST_CONTRACT))

    # Обрезать: удалить Приложения
    trim_body_after(doc, "Приложение №1 к Договору")

    # P[0]: номер договора
    p = find_para(doc, "86/2026")
    if p and "ДОГОВОР" in p.text:
        replace_in_para(p, "86/2026", "{{ДОГОВОР_НОМЕР}}")

    # P[3]: стороны
    p = find_para(doc, "Сенаторова Олега Александровича")
    if p:
        replace_in_para(p, "Сенаторова Олега Александровича", "{{ПОСТАВЩИК_ДИРЕКТОР_ФИО_РП}}")
        replace_in_para(p, "ИП ГКФХ Молчанов Владимир Григорьевич", "{{ПОКУПАТЕЛЬ_НАИМЕНОВАНИЕ}}")

    # P[7]: наименование товара (1.1)
    p = find_para(doc, "ВЕСТА-СЛ-80-18-Ц")
    if p and "1.1." in p.text:
        regex_replace_in_para(
            p,
            r"Весы автомобильные ВЕСТА-[^,]+, max \d+т, размеры платформы\s+\S+м в количестве \d+шт\.",
            "{{ТОВАР_НАИМЕНОВАНИЕ}}",
        )

    # P[21]: срок производства (3.2)
    p = find_para(doc, "двенадцати")
    if p:
        replace_in_para(p, "12 (двенадцати)", "{{СРОК_ПРОИЗВОДСТВА_ДН}}")

    # P[22]: срок доставки (3.3)
    p = find_para(doc, "четырёх")
    if p:
        replace_in_para(p, "4 (четырёх)", "{{СРОК_ДОСТАВКИ_ДН}}")

    # P[23]: адрес поставки (3.4)
    p = find_para(doc, "Каменка")
    if p:
        replace_in_para(
            p,
            "Липецкая область, Тербунский район, с. Каменка",
            "{{АДРЕС_ПОСТАВКИ}}",
        )

    # P[36]: стоимость (4.1)
    p = find_para(doc, "4.1.")
    if p and "242" in p.text:
        text = merge_para_runs(p)
        text = text.replace(
            "2 242\xa0000 (два миллиона двести сорок две тысячи) рублей",
            "{{СУММА_ЦИФРАМИ}} ({{СУММА_ПРОПИСЬЮ}}) рублей",
        )
        if p.runs:
            p.runs[0].text = text

    # P[38]: блок оплаты → payment_lines loop
    p = find_para(doc, "Предоплата 100%")
    if p:
        merge_para_runs(p)
        if p.runs:
            p.runs[0].text = "{% for line in payment_lines %}{{line.text}}{% endfor %}"

    # P[64]: срок действия (9.1)
    p = find_para(doc, "9.1.")
    if p and "31" in p.text:
        text = merge_para_runs(p)
        text = re.sub(r"31\.05\.2026 г\.", "{{СРОК_ДЕЙСТВИЯ_ДО}}", text)
        if p.runs:
            p.runs[0].text = text

    # TABLE[0]: город и дата (шапка)
    t0 = doc.tables[0]
    replace_in_cell(t0.rows[0].cells[0], "г. Воронеж", "{{ДОГОВОР_ГОРОД}}")
    replace_in_cell(t0.rows[0].cells[1], "26 апреля 2026 года", "{{ДОГОВОР_ДАТА}}")

    # TABLE[1]: реквизиты сторон
    t1 = doc.tables[1]

    # Row 1: наименования
    replace_in_cell(
        t1.rows[1].cells[2],
        "ИП ГКФХ Молчанов Владимир Григорьевич",
        "{{ПОКУПАТЕЛЬ_НАИМЕНОВАНИЕ}}",
    )

    # Row 2: реквизиты поставщика (col 0)
    sup = t1.rows[2].cells[0]
    for old, new in [
        ("394005, Российская Федерация, город Воронеж, улица Владимира Невского, дом № 25/1, офис 5.", "{{ПОСТАВЩИК_ЮР_АДРЕС}}"),
        ("ИНН 3662257349", "ИНН {{ПОСТАВЩИК_ИНН}}"),
        ("КПП 366201001", "КПП {{ПОСТАВЩИК_КПП}}"),
        ("ОГРН 1173668061984", "ОГРН {{ПОСТАВЩИК_ОГРН}}"),
        ("Центрально-Черноземный банк Сбербанк России г. Воронеж", "{{ПОСТАВЩИК_БАНК}}"),
        ("40702810513000031419", "{{ПОСТАВЩИК_РС}}"),
        ("30101810600000000681", "{{ПОСТАВЩИК_КС}}"),
        ("042007681", "{{ПОСТАВЩИК_БИК}}"),
    ]:
        replace_in_cell(sup, old, new)

    # Row 2: реквизиты покупателя (col 2)
    cus = t1.rows[2].cells[2]
    for old, new in [
        ("399558, Липецкая область, Тербунский район, с. Вислая Поляна, ул. Ворошилова, 43", "{{ПОКУПАТЕЛЬ_ЮР_АДРЕС}}"),
        ("ИНН 481501477253", "ИНН {{ПОКУПАТЕЛЬ_ИНН}}"),
        ("ОГРН  312480711800010", "ОГРН {{ПОКУПАТЕЛЬ_ОГРН}}"),
        ("40802810535000009577", "{{ПОКУПАТЕЛЬ_РС}}"),
        ("Липецкое отделение №8593 ПАО Сбербанк г. Липецк", "{{ПОКУПАТЕЛЬ_БАНК}}"),
        ("044206604", "{{ПОКУПАТЕЛЬ_БИК}}"),
        ("30101810800000000604", "{{ПОКУПАТЕЛЬ_КС}}"),
        ("matveeva-lubov@mail.ru", "{{ПОКУПАТЕЛЬ_EMAIL}}"),
    ]:
        replace_in_cell(cus, old, new)

    # Row 3: подписи
    patch_signature_cell_sup(t1.rows[3].cells[0])
    patch_signature_cell_cus(t1.rows[3].cells[2])

    doc.save(str(DST_CONTRACT))
    print(f"Создан: {DST_CONTRACT}")


# ─── supply_appendix_1.docx ──────────────────────────────────────────────────

def create_supply_appendix_1():
    shutil.copy(SRC, DST_APP1)
    doc = Document(str(DST_APP1))

    trim_body_before(doc, "Приложение №1 к Договору")
    trim_body_after(doc, "Приложение №2 к Договору")

    # Заголовок приложения
    p = find_para(doc, "Приложение №1")
    if p:
        text = merge_para_runs(p)
        text = re.sub(r"№ 86/20\d+", "№ {{ДОГОВОР_НОМЕР}}", text)
        text = re.sub(r"\d{2}\.\d{2}\.\d{4} года", "{{ДОГОВОР_ДАТА}}", text)
        if p.runs:
            p.runs[0].text = text

    # TABLE[0] в этом документе = TABLE[2] из исходника (спецификация)
    spec = doc.tables[0]
    replace_in_cell(
        spec.rows[1].cells[0],
        "Весы автомобильные ВЕСТА-СЛ-80-18-Ц, max 80т, размеры платформы 18х3м",
        "{{ТОВАР_НАИМЕНОВАНИЕ}}",
    )
    replace_in_cell(spec.rows[1].cells[1], "2 242 000", "{{СУММА_ЦИФРАМИ}}")
    replace_in_cell(spec.rows[2].cells[1], "2 242 000", "{{СУММА_ЦИФРАМИ}}")

    # TABLE[1] = подписи
    if len(doc.tables) >= 2:
        sig = doc.tables[1]
        patch_signature_cell_sup(sig.rows[0].cells[0])
        patch_signature_cell_cus(sig.rows[0].cells[2])

    doc.save(str(DST_APP1))
    print(f"Создан: {DST_APP1}")


# ─── supply_appendix_2.docx ──────────────────────────────────────────────────

def create_supply_appendix_2():
    shutil.copy(SRC, DST_APP2)
    doc = Document(str(DST_APP2))

    trim_body_before(doc, "Приложение №2 к Договору")

    # Заголовок приложения
    p = find_para(doc, "Приложение №2")
    if p:
        text = merge_para_runs(p)
        text = re.sub(r"№ 86/20\d+", "№ {{ДОГОВОР_НОМЕР}}", text)
        text = re.sub(r"\d{2}\.\d{2}\.\d{4} года", "{{ДОГОВОР_ДАТА}}", text)
        if p.runs:
            p.runs[0].text = text

    # Параграфы с наименованием модели
    for keyword in ("Технические характеристики", "Комплект поставки"):
        p = find_para(doc, keyword)
        if p and "ВЕСТА" in p.text:
            regex_replace_in_para(p, r"ВЕСТА-[А-ЯЁA-Z\d-]+", "{{ТОВАР_НАИМЕНОВАНИЕ}}")

    # TABLE[0] = ТТХ
    ttx = doc.tables[0]
    for row_idx, placeholder in [
        (1, "{{ТТХ_MAX}}"),
        (2, "{{ТТХ_ОСЬ}}"),
        (3, "{{ТТХ_РАССТОЯНИЕ_ТЕРМИНАЛ}}"),
        (4, "{{ТТХ_ДИСКРЕТНОСТЬ}}"),
        (5, "{{ТТХ_ГАБАРИТЫ}}"),
        (6, "{{ТТХ_ТЕМПЕРАТУРА}}"),
        (7, "{{ТТХ_СВЯЗЬ}}"),
        (8, "{{ТТХ_ПИТАНИЕ}}"),
        (9, "{{ТТХ_МОЩНОСТЬ}}"),
    ]:
        set_cell_text(ttx.rows[row_idx].cells[2], placeholder)

    # Row 10: ГОСТ строка (объединённые ячейки)
    gost_row = ttx.rows[10]
    for ci in range(len(gost_row.cells)):
        cell = gost_row.cells[ci]
        if "ВЕСТА" in cell.text or "ГОСТ" in cell.text:
            set_cell_text(cell, "{{ТТХ_ГОСТ_СТРОКА}}")
            break  # merged cells — достаточно первой

    # TABLE[1] = комплект (loop)
    if len(doc.tables) >= 2:
        kit = doc.tables[1]
        tbl_el = kit._tbl
        all_trs = list(tbl_el.iterchildren(qn("w:tr")))
        for tr in all_trs[2:]:
            tbl_el.remove(tr)
        tmpl_row = kit.rows[1]
        set_cell_text(tmpl_row.cells[0], "{%tr for item in kit_rows %}{{item.name}}")
        set_cell_text(tmpl_row.cells[1], "{{item.qty}}{%tr endfor %}")

    # TABLE[2] = подписи
    if len(doc.tables) >= 3:
        sig = doc.tables[2]
        patch_signature_cell_sup(sig.rows[0].cells[0])
        patch_signature_cell_cus(sig.rows[0].cells[2])

    doc.save(str(DST_APP2))
    print(f"Создан: {DST_APP2}")


# ─── main ────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    create_supply_contract()
    create_supply_appendix_1()
    create_supply_appendix_2()
