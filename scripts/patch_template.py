"""Одноразовый патч шаблонов. Запускать из корня проекта: python scripts/patch_template.py"""
from pathlib import Path
import shutil

from docx import Document

TEMPLATES = Path("templates/contracts")
BACKUP = Path("templates/backup")


def _replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    if old not in paragraph.text:
        return False
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # Текст разбит по runs — склеиваем в первый
    combined = paragraph.text
    if paragraph.runs:
        paragraph.runs[0].text = combined.replace(old, new)
        for r in paragraph.runs[1:]:
            r.text = ""
        return True
    return False


def _patch_header(doc, old: str, new: str) -> bool:
    for section in doc.sections:
        for p in section.header.paragraphs:
            if _replace_in_paragraph(p, old, new):
                return True
    return False


def _patch_table_cells(doc, old: str, new: str) -> int:
    count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if _replace_in_paragraph(p, old, new):
                        count += 1
    return count


def patch_contract():
    src = TEMPLATES / "contract.docx"
    BACKUP.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, BACKUP / "contract.docx")
    doc = Document(src)
    ok = _patch_header(
        doc,
        "Договор № 29/2026 от 27.03.2026г",
        "Договор № {{ДОГОВОР_НОМЕР}} от {{ДОГОВОР_ДАТА_ПОЛНАЯ}}г",
    )
    if not ok:
        print("WARNING: contract.docx header — строка не найдена, возможно уже исправлено")
        return
    doc.save(src)
    print("contract.docx header OK")


def patch_spec():
    src = TEMPLATES / "spec_foundation_install.docx"
    BACKUP.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, BACKUP / "spec_foundation_install.docx")
    doc = Document(src)
    ok = _patch_header(
        doc,
        "Договор № 110/2026 от 29.05.2026г",
        "Договор № {{ДОГОВОР_НОМЕР}} от {{ДОГОВОР_ДАТА_ПОЛНАЯ}}г",
    )
    if not ok:
        print("WARNING: spec header — строка не найдена")
    else:
        print("spec header OK")
    n = _patch_table_cells(doc, "ООО «Компания Тензосила»", "ООО «ТПК «Тензосила»»")
    if n == 0:
        print("WARNING: «Компания Тензосила» — не найдено в ячейках")
    else:
        print(f"spec table cells ({n} шт.) OK")
    doc.save(src)


if __name__ == "__main__":
    patch_contract()
    patch_spec()
    print("Готово.")
