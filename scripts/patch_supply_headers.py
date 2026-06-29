"""Разовый хирургический патч колонтитулов договора поставки.

create_supply_templates.py шёл только по body+таблицам, section.header не трогал —
во всех трёх шаблонах в header остался литерал «№ 86/2026 от 26.04.2026 года».
Приложения склеиваются в один файл (compose_supply), поэтому патчим все три.
Заодно схлопываем мусорный прогон пробелов перед «п. 4.2» в теле договора.

Запуск: python scripts/patch_supply_headers.py
"""
import re
from pathlib import Path

from docx import Document

TEMPLATES = Path("templates/contracts")
HEADER_FILES = ["supply_contract.docx", "supply_appendix_1.docx", "supply_appendix_2.docx"]


def _merge_runs(para) -> str:
    """Слить все runs параграфа в первый, вернуть полный текст."""
    runs = para.runs
    if not runs:
        return ""
    text = "".join(r.text or "" for r in runs)
    runs[0].text = text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)
    return text


def _patch_header(doc) -> bool:
    """Заменить в header литералы номера/даты на плейсхолдеры. True если нашёл/уже пропатчен."""
    found = False
    for section in doc.sections:
        for para in section.header.paragraphs:
            if "{{ДОГОВОР_НОМЕР}}" in para.text:
                found = True  # уже пропатчен — идемпотентность
                continue
            if "86/2026" not in para.text:
                continue
            text = _merge_runs(para)
            text = re.sub(r"86/2026", "{{ДОГОВОР_НОМЕР}}", text)
            # Поглощаем хвостовое «года»: ДОГОВОР_ДАТА = «26 апреля 2026 года» уже с ним
            text = re.sub(r"\d{2}\.\d{2}\.\d{4}\s*года", "{{ДОГОВОР_ДАТА}}", text)
            if para.runs:
                para.runs[0].text = text
            found = True
    return found


def _collapse_spaces_before_42(doc) -> int:
    """Схлопнуть двойные+ пробелы в параграфах со ссылкой на п.4.2. Вернуть кол-во правок."""
    n = 0
    for para in doc.paragraphs:
        if "4.2" not in para.text:
            continue
        text = _merge_runs(para)
        new = re.sub(r" {2,}", " ", text)
        if new != text:
            if para.runs:
                para.runs[0].text = new
            n += 1
    return n


def patch_headers() -> None:
    for fname in HEADER_FILES:
        path = TEMPLATES / fname
        doc = Document(str(path))
        ok = _patch_header(doc)
        spaces = _collapse_spaces_before_42(doc) if fname == "supply_contract.docx" else 0
        if not ok:
            print(f"⚠ {fname}: литерал 86/2026 в header не найден")
        doc.save(str(path))
        print(f"Пропатчен: {path} (header={ok}, схлопнуто пробелов={spaces})")


if __name__ == "__main__":
    patch_headers()
