"""Создание шаблона spec_v2.docx из spec_foundation_install.docx.

Удаляет хардкодированные пункты обязательств,
заменяя их маркерами {{CLAUSE_SECTION_*}} для динамической подстановки.

Запуск: python scripts/create_spec_v2_template.py
"""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = "templates/contracts/spec_foundation_install.docx"
DST = "templates/contracts/spec_v2.docx"

CLAUSE_SECTIONS = [
    "obligations_supplier",
    "obligations_customer",
    "special_conditions",
    "final",
]


def main():
    doc = Document(SRC)
    paras = doc.paragraphs
    body = doc.element.body

    # Границы: конец секции сроков → начало ТТХ
    terms_end_idx = None
    ttx_start_idx = None

    for i, p in enumerate(paras):
        if "Выезд строительной" in p.text:
            terms_end_idx = i
        if "Технические характеристики" in p.text:
            ttx_start_idx = i
            break

    assert terms_end_idx is not None, "Не найден 'Выезд строительной'"
    assert ttx_start_idx is not None, "Не найден 'Технические характеристики'"
    print(f"Удаляю параграфы {terms_end_idx + 1}..{ttx_start_idx - 1}")

    to_remove = [paras[i]._element for i in range(terms_end_idx + 1, ttx_start_idx)]
    for el in to_remove:
        body.remove(el)
    print(f"Удалено {len(to_remove)} параграфов")

    # Найти ТТХ заново (после удаления индексы сдвинулись)
    ttx_el = None
    for p in doc.paragraphs:
        if "Технические характеристики" in p.text:
            ttx_el = p._element
            break
    assert ttx_el is not None

    for section_id in CLAUSE_SECTIONS:
        marker_p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "{{CLAUSE_SECTION_" + section_id + "}}"
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        marker_p.append(r)
        ttx_el.addprevious(marker_p)
    print(f"Вставлено {len(CLAUSE_SECTIONS)} маркеров")

    doc.save(DST)
    print(f"Шаблон сохранён: {DST}")


if __name__ == "__main__":
    main()
