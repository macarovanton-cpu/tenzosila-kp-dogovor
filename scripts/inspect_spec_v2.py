"""Inspect spec_v2.docx template structure."""
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from pathlib import Path

doc = Document(Path(r"D:\Projects\Tenzosila_KP_Dogovor\templates\contracts\spec_v2.docx"))

# 1. All paragraphs
print("=" * 80)
print("PARAGRAPHS")
print("=" * 80)
for i, p in enumerate(doc.paragraphs):
    text = p.text.replace("\n", "\\n")
    print(f"[{i:3d}] style={p.style.name:<30s} | {text[:100]}")

# 2. All tables summary
print("\n" + "=" * 80)
print("TABLES SUMMARY")
print("=" * 80)
for ti, table in enumerate(doc.tables):
    print(f"\n--- Table {ti} ({len(table.rows)} rows x {len(table.columns)} cols) ---")
    for ri, row in enumerate(table.rows):
        cells_text = " | ".join(c.text.replace("\n", "\\n")[:60] for c in row.cells)
        print(f"  Row {ri}: {cells_text}")

# 3. Table 1 (ТТХ) — full text of column 2
print("\n" + "=" * 80)
print("TABLE 1 — COLUMN 2 (values) FULL TEXT")
print("=" * 80)
if len(doc.tables) > 1:
    t1 = doc.tables[1]
    for ri, row in enumerate(t1.rows):
        if len(row.cells) > 2:
            print(f"  Row {ri}: {row.cells[2].text!r}")
        else:
            print(f"  Row {ri}: (only {len(row.cells)} cells)")

# 4. Table 3 (signatures) — full text of each cell
print("\n" + "=" * 80)
print("TABLE 3 — SIGNATURES FULL TEXT")
print("=" * 80)
if len(doc.tables) > 3:
    t3 = doc.tables[3]
    for ri, row in enumerate(t3.rows):
        for ci, cell in enumerate(row.cells):
            print(f"  Row {ri}, Col {ci}: {cell.text!r}")
else:
    print(f"(only {len(doc.tables)} tables)")

# 5. Paragraph [7] full text
print("\n" + "=" * 80)
print("PARAGRAPH [7] FULL TEXT")
print("=" * 80)
if len(doc.paragraphs) > 7:
    print(doc.paragraphs[7].text)

# 6. Paragraph [48] full text
print("\n" + "=" * 80)
print("PARAGRAPH [48] FULL TEXT")
print("=" * 80)
if len(doc.paragraphs) > 48:
    print(doc.paragraphs[48].text)
else:
    print(f"(only {len(doc.paragraphs)} paragraphs, index 48 out of range)")
