"""Проверка плейсхолдеров в сгенерированных шаблонах."""
import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

TEMPLATES = {
    "supply_contract": Path("templates/contracts/supply_contract.docx"),
    "supply_appendix_1": Path("templates/contracts/supply_appendix_1.docx"),
    "supply_appendix_2": Path("templates/contracts/supply_appendix_2.docx"),
}

EXPECTED = {
    "supply_contract": [
        "{{ДОГОВОР_НОМЕР}}", "{{ДОГОВОР_ДАТА}}", "{{ДОГОВОР_ГОРОД}}",
        "{{ПОКУПАТЕЛЬ_НАИМЕНОВАНИЕ}}", "{{ПОСТАВЩИК_ДИРЕКТОР_ФИО_РП}}",
        "{{ТОВАР_НАИМЕНОВАНИЕ}}", "{{СРОК_ПРОИЗВОДСТВА_ДН}}", "{{СРОК_ДОСТАВКИ_ДН}}",
        "{{АДРЕС_ПОСТАВКИ}}", "{{СУММА_ЦИФРАМИ}}", "{{СУММА_ПРОПИСЬЮ}}",
        "{{PAYMENT_SECTION}}",
        "{{СРОК_ДЕЙСТВИЯ_ДО}}",
        "{{ПОСТАВЩИК_ИНН}}", "{{ПОСТАВЩИК_КПП}}", "{{ПОСТАВЩИК_ОГРН}}",
        "{{ПОСТАВЩИК_БАНК}}", "{{ПОСТАВЩИК_РС}}", "{{ПОСТАВЩИК_КС}}", "{{ПОСТАВЩИК_БИК}}",
        "{{ПОКУПАТЕЛЬ_ИНН}}", "{{ПОКУПАТЕЛЬ_ОГРН}}",
        "{{ПОКУПАТЕЛЬ_РС}}", "{{ПОКУПАТЕЛЬ_БИК}}", "{{ПОКУПАТЕЛЬ_КС}}",
        "{{ПОСТАВЩИК_ДИРЕКТОР_ФИО}}", "{{ПОКУПАТЕЛЬ_ДИРЕКТОР_ДОЛЖНОСТЬ}}",
        "{{ПОКУПАТЕЛЬ_ДИРЕКТОР_ФИО}}", "{{ТЕКУЩИЙ_ГОД}}",
    ],
    "supply_appendix_1": [
        "{{ДОГОВОР_НОМЕР}}", "{{ДОГОВОР_ДАТА}}",
        # Спецификация — row-цикл по spec_rows (весы + доставка)
        "{{row.name}}", "{{row.sum}}", "{{СУММА_ЦИФРАМИ}}",
        "{{ПОСТАВЩИК_ДИРЕКТОР_ФИО}}", "{{ПОКУПАТЕЛЬ_ДИРЕКТОР_ДОЛЖНОСТЬ}}",
        "{{ПОКУПАТЕЛЬ_ДИРЕКТОР_ФИО}}", "{{ТЕКУЩИЙ_ГОД}}",
    ],
    "supply_appendix_2": [
        "{{ДОГОВОР_НОМЕР}}", "{{ДОГОВОР_ДАТА}}", "{{ТОВАР_НАИМЕНОВАНИЕ}}",
        "{{ТТХ_MAX}}", "{{ТТХ_ОСЬ}}", "{{ТТХ_РАССТОЯНИЕ_ТЕРМИНАЛ}}",
        "{{ТТХ_ДИСКРЕТНОСТЬ}}", "{{ТТХ_ГАБАРИТЫ}}", "{{ТТХ_ТЕМПЕРАТУРА}}",
        "{{ТТХ_СВЯЗЬ}}", "{{ТТХ_ПИТАНИЕ}}", "{{ТТХ_МОЩНОСТЬ}}",
        "{{ТТХ_ГОСТ_СТРОКА}}",
        "kit_rows",
        "{{ПОСТАВЩИК_ДИРЕКТОР_ФИО}}", "{{ПОКУПАТЕЛЬ_ДИРЕКТОР_ДОЛЖНОСТЬ}}",
        "{{ПОКУПАТЕЛЬ_ДИРЕКТОР_ФИО}}", "{{ТЕКУЩИЙ_ГОД}}",
    ],
}

BANNED = ["86/2026", "26 апреля 2026 года", "26.04.2026", "Молчанов Владимир Григорьевич",
          "3662257349", "40702810513000031419", "481501477253", "40802810535000009577"]

OUT = Path("docs/source/verify_templates.txt")
lines = []

for name, path in TEMPLATES.items():
    doc = Document(str(path))
    # Собрать весь текст документа (body + таблицы + колонтитулы)
    all_text_parts = []
    for p in doc.paragraphs:
        all_text_parts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    all_text_parts.append(p.text)
    for section in doc.sections:
        for hf in (section.header, section.footer):
            for p in hf.paragraphs:
                all_text_parts.append(p.text)
            for t in hf.tables:
                for row in t.rows:
                    for cell in row.cells:
                        all_text_parts.append(cell.text)
    full_text = "\n".join(all_text_parts)

    lines.append(f"\n{'='*60}")
    lines.append(f"ШАБЛОН: {name}")
    lines.append(f"{'='*60}")

    # Проверить ожидаемые плейсхолдеры
    lines.append("\n  [ОЖИДАЕМЫЕ ПЛЕЙСХОЛДЕРЫ]")
    for ph in EXPECTED[name]:
        found = ph in full_text
        status = "OK" if found else "MISSING"
        lines.append(f"  [{status}] {ph}")

    # Проверить запрещённые значения
    lines.append("\n  [ЗАПРЕЩЁННЫЕ ЗНАЧЕНИЯ (не должны остаться)]")
    for val in BANNED:
        found = val in full_text
        status = "LEAKED!" if found else "OK"
        lines.append(f"  [{status}] {val}")

    # Найти все плейсхолдеры
    placeholders = sorted(set(re.findall(r"\{\{[^}]+\}\}", full_text)))
    lines.append(f"\n  [ВСЕ ПЛЕЙСХОЛДЕРЫ В ДОКУМЕНТЕ] ({len(placeholders)} уникальных)")
    for ph in placeholders:
        lines.append(f"  - {ph}")

    # Найти jinja-теги
    jinja_tags = sorted(set(re.findall(r"\{[%][^%]+[%]\}", full_text)))
    if jinja_tags:
        lines.append(f"\n  [JINJA ТЕГИ]")
        for tag in jinja_tags:
            lines.append(f"  - {tag}")

OUT.write_text("\n".join(lines), encoding="utf-8")
print(f"Готово: {OUT}")
