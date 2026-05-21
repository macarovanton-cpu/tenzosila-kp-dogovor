"""
Патч вёрстки спецификации — баги 5 и 6.
- п.14 (ТХ) начинается с новой страницы и не отрывается от таблицы
- п.15 (Комплект поставки) заголовок не отрывается от таблицы
- Приложение №1 начинается с новой страницы

Запускать из корня проекта:
    python scripts/patch_spec_template.py
"""
import sys
sys.stdout.reconfigure(encoding="utf-8")
from pathlib import Path
import shutil
import zipfile

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE = Path("templates/contracts/spec_foundation_install.docx")
BACKUP_DIR = Path("templates/contracts/backup")


def _set_para_prop(para, prop_name: str) -> None:
    """Установить boolean-свойство параграфа (idempotent)."""
    pPr = para._p.get_or_add_pPr()
    if pPr.find(qn(prop_name)) is None:
        el = OxmlElement(prop_name)
        el.set(qn("w:val"), "1")
        pPr.append(el)


def set_page_break_before(para) -> None:
    _set_para_prop(para, "w:pageBreakBefore")


def set_keep_with_next(para) -> None:
    _set_para_prop(para, "w:keepNext")


def set_table_no_split(table) -> None:
    """Запретить разрыв строк таблицы между страницами (idempotent)."""
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        if trPr.find(qn("w:cantSplit")) is None:
            cant_split = OxmlElement("w:cantSplit")
            cant_split.set(qn("w:val"), "1")
            trPr.append(cant_split)


def _body_paras(doc):
    """Только параграфы верхнего уровня тела документа (не из ячеек таблиц)."""
    body_tag = qn("w:body")
    return [p for p in doc.paragraphs if p._p.getparent().tag == body_tag]


def _find_body_para(doc, contains: str):
    for p in _body_paras(doc):
        if contains in p.text:
            return p
    return None


def _next_body_para(doc, para):
    """Следующий body-level параграф после данного."""
    paras = _body_paras(doc)
    for i, p in enumerate(paras):
        if p._p is para._p and i + 1 < len(paras):
            return paras[i + 1]
    return None


def _replace_in_docx_xml(docx_path: Path, old: str, new: str) -> int:
    """Заменяет строку old→new во всех XML внутри DOCX. Возвращает число замен."""
    tmp = docx_path.with_suffix(".tmp.docx")
    count = 0
    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.endswith(".xml"):
                    text = data.decode("utf-8")
                    count += text.count(old)
                    text = text.replace(old, new)
                    data = text.encode("utf-8")
                zout.writestr(item, data)
    tmp.replace(docx_path)
    return count


def main() -> None:
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(TEMPLATE, BACKUP_DIR / TEMPLATE.name)
    print(f"Бэкап: {BACKUP_DIR / TEMPLATE.name}")

    doc = Document(TEMPLATE)

    # Баг 5: п.14 (ТХ) — с новой страницы, заголовок держится с таблицей
    p14 = _find_body_para(doc, "Технические характеристики")
    p15 = _find_body_para(doc, "Комплект поставки")
    if p14:
        set_page_break_before(p14)
        set_keep_with_next(p14)
        print(f"п.14: pageBreakBefore + keepNext: {p14.text[:60]!r}")
        # Пустой параграф сразу перед TABLE[1] — продолжаем цепочку keepNext
        p14_next = _next_body_para(doc, p14)
        if p14_next is not None and not p14_next.text.strip():
            set_keep_with_next(p14_next)
            print("п.14 next (пустой): keepNext")
    else:
        print("WARNING: п.14 не найден")

    # TABLE[1] — ТХ: keepNext на всех строках создаёт цепочку внутри таблицы
    if len(doc.tables) > 1:
        th_table = doc.tables[1]
        for row in th_table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    set_keep_with_next(p)
        print(f"TABLE[1] все строки: keepNext на {len(th_table.rows)} строках")
    else:
        print("WARNING: TABLE[1] (ТХ) не найдена")

    # Цепочка keepNext: параграфы [55-57] между TABLE[1] и п.15
    if len(doc.tables) > 1 and p14 and p15:
        all_bp = _body_paras(doc)
        p14_i = next(i for i, p in enumerate(all_bp) if p._p is p14._p)
        p15_i = next(i for i, p in enumerate(all_bp) if p._p is p15._p)
        intermediates = all_bp[p14_i + 2 : p15_i]
        for p in intermediates:
            set_keep_with_next(p)
        print(f"Параграфы между TABLE[1] и п.15: keepNext на {len(intermediates)} параграфах")

    # п.15 (Комплект поставки): заголовок не отрывается от таблицы
    if p15:
        set_keep_with_next(p15)
        print(f"п.15: keepNext: {p15.text[:60]!r}")
    else:
        print("WARNING: п.15 не найден")

    # Продолжаем keepNext-цепь: п.15 → пустой → TABLE[2]
    if p15:
        p15_next = _next_body_para(doc, p15)
        if p15_next is not None and not p15_next.text.strip():
            set_keep_with_next(p15_next)
            print("п.15 next (пустой): keepNext")

    # TABLE[2] — Комплект: keepNext на всех строках, последняя замыкает цепочку на [61]
    if len(doc.tables) > 2:
        kp_table = doc.tables[2]
        for row in kp_table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    set_keep_with_next(p)
        print(f"TABLE[2] все строки: keepNext на {len(kp_table.rows)} строках")
    else:
        print("WARNING: TABLE[2] (Комплект поставки) не найдена")

    # TABLE[1] и TABLE[2]: cantSplit на всех строках
    for tbl_idx, tbl_name in [(1, "ТХ"), (2, "Комплект поставки")]:
        if len(doc.tables) > tbl_idx:
            set_table_no_split(doc.tables[tbl_idx])
            print(f"TABLE[{tbl_idx}] ({tbl_name}): cantSplit на {len(doc.tables[tbl_idx].rows)} строках")
        else:
            print(f"WARNING: TABLE[{tbl_idx}] ({tbl_name}) не найдена")

    # Параграф [61] между TABLE[2] и таблицей подписей — замыкаем цепочку
    if p14 and p15:
        all_bp = _body_paras(doc)
        p15_i = next(i for i, p in enumerate(all_bp) if p._p is p15._p)
        if p15_i + 2 < len(all_bp):
            set_keep_with_next(all_bp[p15_i + 2])
            print("Параграф после TABLE[2] (перед подписями): keepNext")
        else:
            print("WARNING: параграф [61] после TABLE[2] не найден")

    # TABLE[3] (таблица подписей): cantSplit
    if len(doc.tables) > 3:
        set_table_no_split(doc.tables[3])
        print(f"TABLE[3] (подписи): cantSplit на {len(doc.tables[3].rows)} строках")
    else:
        print("WARNING: TABLE[3] (подписи) не найдена")

    # Баг 6: Приложение №1 — с новой страницы
    # Ищем параграф-заголовок (начинается с «Приложение №{{...»), а не ссылку внутри текста
    pril = _find_body_para(doc, "Приложение №{{СПЕЦ_НОМЕР}}")
    if pril:
        set_page_break_before(pril)
        print(f"Приложение: pageBreakBefore: {pril.text[:60]!r}")
    else:
        print("WARNING: Приложение не найдено")

    doc.save(TEMPLATE)
    print(f"\nСохранено: {TEMPLATE}")

    # Баг 9: устаревший плейсхолдер в text box Приложения (не доступен через python-docx API)
    n = _replace_in_docx_xml(
        TEMPLATE,
        "ЗАКАЗЧИК_ДИРЕКТОР_ФИО_КРАТКОЕ",
        "ЗАКАЗЧИК_ДИРЕКТОР_ИНИЦИАЛЫ",
    )
    if n > 0:
        print(f"Баг 9: плейсхолдер ФИО_КРАТКОЕ → ИНИЦИАЛЫ ({n} вхождений)")
    else:
        print("INFO: плейсхолдер ФИО_КРАТКОЕ не найден (уже исправлен)")


if __name__ == "__main__":
    main()
