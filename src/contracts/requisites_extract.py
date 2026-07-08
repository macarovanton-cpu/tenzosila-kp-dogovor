"""Извлечение текста из файлов реквизитов (DOCX / PDF с текстовым слоем).

Единая точка входа — extract_text(): текст идёт дальше в общий пайплайн
parse_requisites -> derive_requisites. Скан (PDF без текстового слоя)
сигнализируется исключением NoTextLayerError — вызывающий код отличает его
от валидного текста структурно, не сравнением строк.

Осознанное ограничение: валидный PDF с < MIN_NONSPACE_CHARS непробельных
символов (обрезок в одну строку) пометится как скан — менеджер получит
громкое «вставьте текстом», не тихую порчу. Дешёвого способа отличить его
от скана нет. К DOCX порог не применяется: текстовый слой там есть по
построению, короткий валидный документ легитимен.
"""
from __future__ import annotations

from pathlib import Path

# Порог детекта скана: меньше непробельных символов — считаем, что
# текстового слоя нет (только для PDF)
MIN_NONSPACE_CHARS = 50


class NoTextLayerError(Exception):
    """PDF без текстового слоя (скан) — распознавание невозможно."""


def _docx_table_lines(table) -> list[str]:
    """Строки таблицы DOCX: ячейки строки склеиваем пробелом.

    Вложенные таблицы обходим одним уровнем (частый случай в картах партнёра).
    """
    lines: list[str] = []
    for row in table.rows:
        parts: list[str] = []
        for cell in row.cells:
            cell_text = " ".join(
                p.text.strip() for p in cell.paragraphs if p.text.strip()
            )
            if cell_text:
                parts.append(cell_text)
            for nested in cell.tables:
                lines.extend(_docx_table_lines(nested))
        if parts:
            lines.append(" ".join(parts))
    return lines


def extract_docx(path: Path) -> str:
    """Текст DOCX: абзацы + таблицы. Никогда не бросает NoTextLayerError."""
    import docx  # noqa: PLC0415

    doc = docx.Document(str(path))
    lines: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())
    for table in doc.tables:
        lines.extend(_docx_table_lines(table))
    return "\n".join(lines)


def extract_pdf(path: Path) -> str:
    """Текст PDF (pymupdf, блочный get_text). Скан → NoTextLayerError."""
    import fitz  # noqa: PLC0415

    with fitz.open(str(path)) as doc:
        pages = [page.get_text() for page in doc]
    text = "\n".join(pages)
    if len("".join(text.split())) < MIN_NONSPACE_CHARS:
        raise NoTextLayerError(f"PDF без текстового слоя: {path.name}")
    return text


def extract_text(path: Path) -> str:
    """Диспетчер по расширению: .docx / .pdf → текст, иначе ValueError."""
    ext = path.suffix.lower()
    if ext == ".docx":
        return extract_docx(path)
    if ext == ".pdf":
        return extract_pdf(path)
    raise ValueError(f"Неподдерживаемый формат файла: {ext or '<без расширения>'}")
