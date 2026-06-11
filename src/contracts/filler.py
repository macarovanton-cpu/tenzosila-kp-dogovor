"""
filler.py — подстановка данных в Word-шаблон с плейсхолдерами.
Работает напрямую с python-docx, форматирование сохраняется 100%.
"""

import copy
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def merge_runs(paragraph) -> None:
    """
    Склеивает соседние runs с одинаковым форматированием в одном параграфе.
    Нужно потому что Word иногда разбивает текст на несколько runs,
    из-за чего плейсхолдер {{КЛЮЧ}} может быть разбит на {{, КЛЮЧ, }}.
    """
    runs = paragraph.runs
    if len(runs) < 2:
        return

    i = 0
    while i < len(runs) - 1:
        curr = runs[i]
        next_r = runs[i + 1]

        # run.text= вызывает CT_R.clear_content() и уничтожает w:drawing; пропускаем
        if (curr._r.find(qn('w:drawing')) is not None
                or next_r._r.find(qn('w:drawing')) is not None):
            i += 1
            continue

        # Склеиваем если форматирование совпадает
        curr_rpr = curr._r.find(qn('w:rPr'))
        next_rpr = next_r._r.find(qn('w:rPr'))

        curr_xml = '' if curr_rpr is None else curr_rpr.xml
        next_xml = '' if next_rpr is None else next_rpr.xml

        if curr_xml == next_xml:
            curr.text += next_r.text
            next_r._r.getparent().remove(next_r._r)
            # После удаления обновляем список
            runs = paragraph.runs
        else:
            i += 1


def replace_in_paragraph(paragraph, data: dict) -> None:
    """
    Заменяет все плейсхолдеры {{КЛЮЧ}} в параграфе на значения из data.
    Сначала склеивает runs, потом делает замену.
    """
    merge_runs(paragraph)

    for run in paragraph.runs:
        text = run.text
        if '{{' in text:
            for key, value in data.items():
                placeholder = f'{{{{{key}}}}}'
                if placeholder in text:
                    text = text.replace(placeholder, str(value) if value else '')
            run.text = text


def _replace_textbox_placeholders(docx_path: str, data: dict) -> None:
    """Заменяет {{KEY}} в text box-ах (w:txbxContent), которые python-docx пропускает."""
    import shutil
    import zipfile

    tmp = docx_path + '.tmp'
    with zipfile.ZipFile(docx_path, 'r') as zin:
        with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data_bytes = zin.read(item.filename)
                if item.filename.endswith('.xml'):
                    text = data_bytes.decode('utf-8')
                    for key, value in data.items():
                        placeholder = '{{' + key + '}}'
                        if placeholder in text:
                            text = text.replace(placeholder, str(value) if value else '')
                    data_bytes = text.encode('utf-8')
                zout.writestr(item, data_bytes)
    shutil.move(tmp, docx_path)


def fill_template(template_path: str, data: dict, output_path: str) -> None:
    """
    Главная функция: открывает шаблон .docx, подставляет все плейсхолдеры,
    сохраняет результат в output_path.

    data — плоский словарь {КЛЮЧ: значение}.
    Если передан словарь с 'requisites' и 'specification' — автоматически объединяет.
    """
    # Если data содержит вложенные блоки — объединяем в плоский словарь
    if 'requisites' in data or 'specification' in data:
        flat = {}
        flat.update(data.get('requisites', {}))
        flat.update(data.get('specification', {}))
        data = flat

    doc = Document(template_path)

    # Обрабатываем все параграфы документа
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, data)

    # Обрабатываем все таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, data)

    # Обрабатываем колонтитулы — только параграфы с плейсхолдерами.
    # Guard нужен: merge_runs уничтожает field-runs (fldChar/instrText) если у них
    # одинаковый rPr, а footer содержит поле PAGE без плейсхолдеров {{...}}.
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            if '{{' in paragraph.text:
                replace_in_paragraph(paragraph, data)
        for paragraph in section.footer.paragraphs:
            if '{{' in paragraph.text:
                replace_in_paragraph(paragraph, data)

    # Удаляем пустые нумерованные параграфы (СПЕЦ_ОПЛАТА_П5/П6 без значения → «2.5 »)
    _NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    _body_tag = f'{{{_NS}}}body'
    for para in list(doc.paragraphs):
        if para.text.strip() == '':
            p_el = para._element
            if (p_el.getparent().tag == _body_tag
                    and p_el.find(f'.//{{{_NS}}}numPr') is not None):
                p_el.getparent().remove(p_el)

    doc.save(output_path)
    _replace_textbox_placeholders(output_path, data)


def get_unfilled_placeholders(docx_path: str) -> list[str]:
    """
    Возвращает список незаполненных плейсхолдеров в документе.
    Удобно для проверки после генерации.
    """
    doc = Document(docx_path)
    found = set()

    def scan_text(text):
        for match in re.finditer(r'\{\{[^}]+\}\}', text):
            found.add(match.group())

    for p in doc.paragraphs:
        scan_text(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    scan_text(p.text)

    return sorted(found)


def remove_empty_paragraphs(doc_path: str, output_path: str = None) -> None:
    """
    Удаляет параграфы в которых после подстановки ничего не осталось.
    Используется для удаления лишних пунктов оплаты (если их меньше 6).
    """

    doc = Document(doc_path)
    output_path = output_path or doc_path

    for paragraph in doc.paragraphs:
        # Удаляем параграфы которые полностью пустые после подстановки
        if paragraph.text.strip() == '':
            # Проверяем что это не структурный пустой параграф
            # (между разделами), а параграф-плейсхолдер оплаты
            p = paragraph._element
            parent = p.getparent()
            # Удаляем только если родитель — body (не таблица)
            if parent.tag.endswith('}body'):
                # Дополнительная проверка: параграф имеет нумерацию (список)
                ppr = p.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                if ppr is not None:
                    parent.remove(p)

    doc.save(output_path)


def _fmt(amount: int) -> str:
    return f"{amount:,}".replace(",", " ") if amount else ""


def _clear_row_text(tr_el) -> None:
    """Обнулить все w:t элементы в строке таблицы."""
    for t_el in tr_el.findall('.//' + qn('w:t')):
        t_el.text = ''


def _set_cell_text(tc_el, text: str) -> None:
    """Записать текст в первый w:t элемент ячейки таблицы."""
    t_els = tc_el.findall('.//' + qn('w:t'))
    if t_els:
        t_els[0].text = text
    else:
        p_els = tc_el.findall('.//' + qn('w:p'))
        if p_els:
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = text
            t.set(qn('xml:space'), 'preserve')
            r.append(t)
            p_els[0].append(r)


def fill_spec_with_items(
    template_path: str,
    data: dict,
    items: list[dict],
    output_path: str,
) -> None:
    """Рендер шаблона спецификации с динамическим массивом позиций.

    Двухшаговый подход:
    1. fill_template() для всех плейсхолдеров кроме строк таблицы.
    2. python-docx: заменить строки Table[0] на позиции из items.

    Итого вычисляется из items и передаётся в fill_template.
    """
    from src.contracts.utils import number_to_words

    grand_total = sum(
        int(item.get("total", 0))
        for item in items
        if not item.get("metadata", {}).get("customer_side")
    )

    fill_data = dict(data)
    fill_data["СПЕЦ_ИТОГО"] = _fmt(grand_total)
    fill_data["СПЕЦ_ИТОГО_ПРОПИСЬ"] = number_to_words(grand_total)
    for i in range(1, 6):
        fill_data.setdefault(f"СПЕЦ_П{i}_НАИМЕНОВАНИЕ", "")
        fill_data.setdefault(f"СПЕЦ_П{i}_СУММА", "")
        fill_data.setdefault(f"СПЕЦ_П{i}_СРОК", "")

    fill_template(template_path, fill_data, output_path)

    doc = Document(output_path)
    table = doc.tables[0]
    tbl = table._tbl

    all_trs = [c for c in tbl if c.tag == qn('w:tr')]
    if len(all_trs) < 2:
        doc.save(output_path)
        return

    header_tr = all_trs[0]
    template_tr = copy.deepcopy(all_trs[1])

    for tr in all_trs[1:-1]:
        tbl.remove(tr)

    for item in reversed(items):
        new_tr = copy.deepcopy(template_tr)
        _clear_row_text(new_tr)

        customer_side = item.get("metadata", {}).get("customer_side", False)
        name_text = item.get("name", "")
        total_val = int(item.get("total", 0))
        total_text = "ЗАКАЗЧИК" if customer_side else (_fmt(total_val) if total_val else "")

        tcs = [c for c in new_tr if c.tag == qn('w:tc')]
        if tcs:
            _set_cell_text(tcs[0], name_text)
        if len(tcs) > 1:
            _set_cell_text(tcs[1], total_text)

        header_tr.addnext(new_tr)

    doc.save(output_path)
