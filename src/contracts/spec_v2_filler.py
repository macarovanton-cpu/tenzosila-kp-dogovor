"""spec_v2_filler.py — рендер спецификации v2 с динамическими clauses."""
import copy
import logging
from datetime import datetime

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.contracts.filler import _set_cell_text, fill_spec_with_items

_logger = logging.getLogger(__name__)

_CLAUSE_SECTION_ORDER = [
    "obligations_supplier",
    "obligations_customer",
    "special_conditions",
    "final",
]

_TTH_KEYS = (
    "ТТХ_НАГРУЗКА_НА_ОСЬ",
    "ТТХ_РАССТОЯНИЕ_ДО_ТЕРМИНАЛА",
    "ТТХ_ДИСКРЕТНОСТЬ_1",
    "ТТХ_ДИСКРЕТНОСТЬ_2",
    "ТТХ_ГРАНИЦА_1",
    "ТТХ_ГРАНИЦА_2",
    "ТТХ_ГАБАРИТЫ",
    "ТТХ_ТЕМПЕРАТУРА",
)


# Порядок дочерних элементов CT_RPr по схеме OOXML (wml.xsd, EG_RPrBase) —
# нужен, чтобы вставлять <w:b/> в скопированный rPr на корректную позицию.
_RPR_CHILD_ORDER = [
    "rStyle", "rFonts", "b", "bCs", "i", "iCs", "caps", "smallCaps", "strike",
    "dstrike", "outline", "shadow", "emboss", "imprint", "noProof", "snapToGrid",
    "vanish", "webHidden", "color", "spacing", "w", "kern", "position", "sz",
    "szCs", "highlight", "u", "effect", "bdr", "shd", "fitText", "vertAlign",
    "rtl", "cs", "em", "lang", "eastAsianLayout", "specVanish", "oMath",
]


def _insert_bold(rpr) -> None:
    """Добавить <w:b/> в rPr на позицию по схеме OOXML, если его там ещё нет."""
    if rpr.find(qn('w:b')) is not None:
        return
    b_pos = _RPR_CHILD_ORDER.index('b')
    insert_before = None
    for child in rpr:
        local = child.tag.split('}')[-1]
        if local in _RPR_CHILD_ORDER and _RPR_CHILD_ORDER.index(local) > b_pos:
            insert_before = child
            break
    b_el = OxmlElement('w:b')
    if insert_before is not None:
        insert_before.addprevious(b_el)
    else:
        rpr.append(b_el)


def _build_run_rpr(rpr, bold: bool):
    """Собрать rPr для нового рана: deepcopy формата маркера + (опционально) bold."""
    if rpr is not None:
        new_rpr = copy.deepcopy(rpr)
        if bold:
            _insert_bold(new_rpr)
        return new_rpr
    if bold:
        new_rpr = OxmlElement('w:rPr')
        new_rpr.append(OxmlElement('w:b'))
        return new_rpr
    return None


def _make_clause_para(
    text: str, bold: bool = False, justify: bool = False, rpr=None, ppr=None,
):
    """Создать XML-элемент параграфа для clause секции.

    rpr — rPr-элемент (например, рана с маркером), который копируется в каждый
    создаваемый ран, чтобы унаследовать форматирование маркера (шрифт, размер).
    ppr — pPr-элемент параграфа с маркером: если передан, копируется целиком
    и является источником истины для форматирования параграфа (интервал,
    выравнивание и т.п.) — параметр justify в этом случае игнорируется,
    т.к. сам ppr уже несёт нужное выравнивание маркера.
    """
    p = OxmlElement('w:p')
    if ppr is not None:
        p.append(copy.deepcopy(ppr))
    elif justify:
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'both')
        pPr.append(jc)
        p.append(pPr)
    if not text:
        return p
    lines = text.split('\n')
    for i, line in enumerate(lines):
        r = OxmlElement('w:r')
        run_rpr = _build_run_rpr(rpr, bold)
        if run_rpr is not None:
            r.append(run_rpr)
        t_el = OxmlElement('w:t')
        t_el.text = line
        t_el.set(qn('xml:space'), 'preserve')
        r.append(t_el)
        p.append(r)
        if i < len(lines) - 1:
            br_r = OxmlElement('w:r')
            br_rpr = _build_run_rpr(rpr, bold)
            if br_rpr is not None:
                br_r.append(br_rpr)
            br_r.append(OxmlElement('w:br'))
            p.append(br_r)
    return p


def _get_marker_rpr(para, marker: str):
    """Найти rPr рана с маркером — для наследования форматирования новыми ранами."""
    for run in para.runs:
        if marker in run.text:
            return run._element.find(qn('w:rPr'))
    if para.runs:
        return para.runs[0]._element.find(qn('w:rPr'))
    return None


def _get_marker_ppr(para):
    """Найти pPr параграфа с маркером — для наследования интервала/выравнивания."""
    return para._element.find(qn('w:pPr'))


def _replace_marker_with_paragraphs(
    doc, marker: str, lines: list[str],
) -> bool:
    """Заменить {{MARKER}} на список параграфов. Возвращает True если нашёл."""
    for para in doc.paragraphs:
        if marker in para.text:
            p_el = para._element
            marker_rpr = _get_marker_rpr(para, marker)
            marker_ppr = _get_marker_ppr(para)
            for line in reversed(lines):
                p_el.addnext(_make_clause_para(line, rpr=marker_rpr, ppr=marker_ppr))
            p_el.getparent().remove(p_el)
            return True
    return False


def _remove_marker(doc, marker: str) -> bool:
    """Удалить параграф с маркером. Возвращает True если нашёл."""
    for para in doc.paragraphs:
        if marker in para.text:
            para._element.getparent().remove(para._element)
            return True
    return False


def _fill_kit_table(doc, kit_items: list[dict]) -> None:
    """Заполнить таблицу комплекта поставки (Table 2) клонированием шаблонной строки."""
    if len(doc.tables) < 3:
        return
    table = doc.tables[2]
    tbl = table._tbl
    all_trs = list(tbl.iterchildren(qn('w:tr')))
    if len(all_trs) < 2:
        return

    template_tr = copy.deepcopy(all_trs[1])
    tbl.remove(all_trs[1])

    all_trs = list(tbl.iterchildren(qn('w:tr')))
    header_tr = all_trs[0]

    for item in reversed(kit_items):
        new_tr = copy.deepcopy(template_tr)
        tcs = [c for c in new_tr if c.tag == qn('w:tc')]
        if tcs:
            _set_cell_text(tcs[0], item.get("name", ""))
        if len(tcs) > 1:
            _set_cell_text(tcs[1], item.get("qty", ""))
        header_tr.addnext(new_tr)


def _payment_lines_from_data(data: dict) -> list[str]:
    """Собрать строки оплаты из СПЕЦ_ОПЛАТА_П1-6 в data."""
    return [data[k] for k in (f"СПЕЦ_ОПЛАТА_П{i}" for i in range(1, 7)) if data.get(k)]


def _load_model_and_deps(deal: dict):
    """Загрузить (model, line_defaults, sensor, indicator) из JSON-справочников.

    Возвращает None если lookup не удался (модель не найдена, файлы недоступны).
    """
    import json
    from pathlib import Path
    try:
        models_json = json.loads(Path("data/models.json").read_text(encoding="utf-8"))
        specs_json = json.loads(Path("data/equipment_specs.json").read_text(encoding="utf-8"))

        items = deal.get("items", [])
        if not items:
            return None
        meta = items[0].get("metadata", {})
        line = meta.get("line")
        max_t = meta.get("max")
        length = meta.get("length")
        if not line or max_t is None or length is None:
            return None

        model = next(
            (m for m in models_json.get("models", [])
             if m.get("line") == line
             and m.get("max_load_t") == max_t
             and m.get("length_m") == length),
            None,
        )
        if not model:
            return None

        ld = models_json.get("line_defaults", {}).get(line, {})
        default_sensor_name = ld.get("default_sensor", "")

        sensor: dict = {}
        for s in specs_json.get("sensors", []):
            full = f"{s.get('manufacturer', '')} {s.get('model', '')}"
            if default_sensor_name.startswith(full):
                sensor = s
                break

        default_indicator_name = ld.get("default_indicator", "")
        indicator: dict = {}
        for t in specs_json.get("terminals", []):
            if t.get("model", "") == default_indicator_name:
                indicator = t
                break

        return model, ld, sensor, indicator
    except Exception:
        return None


def fill_spec_v2(
    template_path: str,
    data: dict,
    items: list[dict],
    deal: dict,
    output_path: str,
    model_qty: int = 1,
) -> None:
    """Рендер спецификации v2: таблица позиций + динамические секции.

    data может содержать переопределения (если переданы явно):
      _payment_lines: list[str] — строки оплаты
      _terms_lines:   list[str] — строки сроков
      _kit_items:     list[dict] — [{name, qty}] для комплекта
      ТТХ_*:          str — значения ТТХ-плейсхолдеров
      ТЕКУЩИЙ_ГОД:    str — год в подписном блоке

    Если перечисленные ключи отсутствуют — вычисляются автоматически из deal.
    """
    from src.contracts.clauses_renderer import build_contract_clauses

    # Не мутировать dict вызывающего
    data = {**data}

    # --- Константы ---
    data.setdefault("ТЕКУЩИЙ_ГОД", str(datetime.now().year))

    # --- Clauses: заранее посчитать номера для плейсхолдеров ТТХ/комплекта ---
    clauses_by_section = build_contract_clauses(deal)
    last_clause_number = max(
        (
            int(clause.auto_number)
            for section_clauses in clauses_by_section.values()
            for clause in section_clauses
        ),
        default=3,
    )
    data["ТТХ_НОМЕР"] = str(last_clause_number + 1)
    data["КОМПЛЕКТ_НОМЕР"] = str(last_clause_number + 2)

    # --- ТТХ из models.json (дозаполняет отсутствующие ручные значения) ---
    deps = None
    if any(key not in data for key in _TTH_KEYS):
        deps = _load_model_and_deps(deal)
        if deps:
            from src.contracts.tth_context import build_tth_data
            for key, value in build_tth_data(deps[0], deps[2]).items():
                data.setdefault(key, value)

    fill_spec_with_items(template_path, data, items, output_path, model_qty=model_qty)

    doc = Document(output_path)

    # --- Payment ---
    payment_lines = data.get("_payment_lines") or _payment_lines_from_data(data)
    if payment_lines:
        found = _replace_marker_with_paragraphs(doc, "{{PAYMENT_SECTION}}", payment_lines)
    else:
        found = _remove_marker(doc, "{{PAYMENT_SECTION}}")
    if not found:
        _logger.warning("маркер {{PAYMENT_SECTION}} не найден в шаблоне спецификации v2")

    # --- Terms ---
    if "_terms_lines" in data:
        terms_lines = data["_terms_lines"]
    else:
        from src.contracts.terms_renderer import render_terms_section
        terms_lines = render_terms_section(deal, deal.get("items", []))
    if terms_lines:
        found = _replace_marker_with_paragraphs(doc, "{{TERMS_SECTION}}", terms_lines)
    else:
        found = _remove_marker(doc, "{{TERMS_SECTION}}")
    if not found:
        _logger.warning("маркер {{TERMS_SECTION}} не найден в шаблоне спецификации v2")

    # --- Kit ---
    if "_kit_items" in data:
        kit_items = data["_kit_items"]
    else:
        if deps is None:
            deps = _load_model_and_deps(deal)
        kit_items = []
        if deps:
            from src.contracts.kit_renderer import build_kit_items
            cable_m = deps[1].get("default_cable_length_m", 20)
            kit_items = build_kit_items(deps[0], deps[1], deps[2], deps[3], cable_m)
    if kit_items:
        _fill_kit_table(doc, kit_items)

    # --- Clauses ---
    for section_id in _CLAUSE_SECTION_ORDER:
        marker = "{{CLAUSE_SECTION_" + section_id + "}}"
        clauses = clauses_by_section.get(section_id, [])
        for para in doc.paragraphs:
            if marker in para.text:
                p_el = para._element
                if clauses:
                    for clause in clauses:
                        text = f"{clause.auto_number}. {clause.text}"
                        p_el.addprevious(
                            _make_clause_para(text, justify=True)
                        )
                p_el.getparent().remove(p_el)
                break

    doc.save(output_path)
