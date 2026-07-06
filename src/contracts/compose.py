"""Склейка спецификации с внешними DOCX-приложениями."""
import logging
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docxcompose.composer import Composer
from docxtpl import DocxTemplate

from src.contracts.filler import (
    fill_template,
    split_total_row_by_qty,
)
from src.contracts.spec_v2_filler import _remove_marker, _replace_marker_with_paragraphs

_logger = logging.getLogger(__name__)

_SUPPLY_TEMPLATES = Path("templates/contracts")


def compose_supply(context: dict, output_path: Path) -> None:
    """Рендер 3 шаблонов договора поставки и склейка в один DOCX.

    context — результат supply_filler.build_supply_context().
    Ключ _payment_lines используется для двупрохода (docxtpl → python-docx).
    Остальные ключи передаются docxtpl напрямую.
    """
    payment_lines: list[str] = context.get("_payment_lines", [])
    # Контекст для docxtpl: без служебных ключей с подчёркиванием
    tpl_context = {k: v for k, v in context.items() if not k.startswith("_")}

    tmp_files: list[Path] = []
    try:
        # --- 1. supply_contract (двупроход) ---
        contract_tpl = _SUPPLY_TEMPLATES / "supply_contract.docx"
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            contract_tmp1 = Path(f.name)
        tmp_files.append(contract_tmp1)

        tpl = DocxTemplate(str(contract_tpl))
        tpl.render(tpl_context)
        tpl.save(str(contract_tmp1))

        # Второй проход: заменить {{PAYMENT_SECTION}} реальными параграфами.
        # Пустой список оплаты → удалить маркер, иначе в п.4.2 остаётся пустой блок.
        contract_doc = Document(str(contract_tmp1))
        if payment_lines:
            found = _replace_marker_with_paragraphs(
                contract_doc, "{{PAYMENT_SECTION}}", payment_lines
            )
        else:
            found = _remove_marker(contract_doc, "{{PAYMENT_SECTION}}")
        if not found:
            _logger.warning(
                "маркер {{PAYMENT_SECTION}} не найден в supply_contract.docx"
            )
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            contract_tmp2 = Path(f.name)
        tmp_files.append(contract_tmp2)
        contract_doc.save(str(contract_tmp2))

        # --- 2. supply_appendix_1 ---
        appendix1_tpl = _SUPPLY_TEMPLATES / "supply_appendix_1.docx"
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            app1_tmp = Path(f.name)
        tmp_files.append(app1_tmp)
        tpl1 = DocxTemplate(str(appendix1_tpl))
        tpl1.render(tpl_context)
        tpl1.save(str(app1_tmp))

        # При model_qty > 1 — вторая строка ИТОГО в спецификации appendix_1:
        # «ИТОГО за 1 весы» (per-unit) + «ИТОГО за N весов» (× qty).
        qty_scales = int(context.get("_qty_scales", 1) or 1)
        if qty_scales > 1:
            _add_scale_qty_total_row(
                app1_tmp,
                qty_scales,
                context.get("_itogo_per_1", ""),
                context.get("_itogo_n", ""),
            )

        # --- 3. supply_appendix_2 ---
        appendix2_tpl = _SUPPLY_TEMPLATES / "supply_appendix_2.docx"
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            app2_tmp = Path(f.name)
        tmp_files.append(app2_tmp)
        tpl2 = DocxTemplate(str(appendix2_tpl))
        tpl2.render(tpl_context)
        tpl2.save(str(app2_tmp))

        # --- Склейка через docxcompose ---
        composer = Composer(Document(str(contract_tmp2)))

        for app_tmp in (app1_tmp, app2_tmp):
            app_doc = Document(str(app_tmp))
            _ensure_first_page_break_before(app_doc)
            composer.append(app_doc)

        composer.save(str(output_path))
    finally:
        for tmp_file in tmp_files:
            tmp_file.unlink(missing_ok=True)


def _add_scale_qty_total_row(
    app1_path: Path, qty_scales: int, itogo_per_1: str, itogo_n: str
) -> None:
    """Вставить вторую строку ИТОГО в спецификацию appendix_1 (Table[0]).

    После docxtpl-рендера последняя строка Table[0] — ИТОГО (2 колонки:
    наименование | сумма). Переписываем её в «за 1 весы» и добавляем «за N весов».
    """
    doc = Document(str(app1_path))
    if not doc.tables:
        return
    tbl = doc.tables[0]._tbl
    rows = [c for c in tbl if c.tag == qn('w:tr')]
    if len(rows) < 2:
        return
    itogo_tr = rows[-1]
    split_total_row_by_qty(
        itogo_tr,
        per1_cells={0: "ИТОГО за 1 весы", 1: itogo_per_1},
        n_cells={0: f"ИТОГО за {qty_scales} весов", 1: itogo_n},
    )
    doc.save(str(app1_path))


def _collect_appendices(attachments: dict) -> list[tuple[Path, int]]:
    """Вернуть пути и фиксированные номера: build_task -> 1, control_sheet -> 2."""
    out: list[tuple[Path, int]] = []
    build_task = (attachments.get("build_task_path") or "").strip()
    if build_task and attachments.get("build_task_source") != "none":
        out.append((Path(build_task), 1))

    control_sheet = (attachments.get("control_sheet_path") or "").strip()
    if attachments.get("include_control_sheet") and control_sheet:
        out.append((Path(control_sheet), 2))
    return out


def _ensure_first_page_break_before(doc) -> None:
    """Поставить pageBreakBefore на первый параграф приложения перед docxcompose."""
    if not doc.paragraphs:
        return
    doc.paragraphs[0].paragraph_format.page_break_before = True


def compose_spec_with_attachments(
    spec_path: Path,
    attachments: dict,
    data: dict,
) -> None:
    """Подклеить приложения к спецификации на месте.

    Порядок: Спецификация -> Приложение N1 (build_task) -> N2 (control_sheet).
    Пути берутся из attachments. Если приложений нет, файл не меняется.
    """
    appendices = _collect_appendices(attachments)
    if not appendices:
        return

    spec_path = Path(spec_path)
    composer = Composer(Document(str(spec_path)))
    tmp_files: list[Path] = []
    try:
        for src, number in appendices:
            filled = spec_path.parent / f"_attach_{number}_{spec_path.stem}.docx"
            fill_template(
                str(src),
                {**data, "ПРИЛОЖЕНИЕ_НОМЕР": str(number)},
                str(filled),
            )
            tmp_files.append(filled)
            appendix_doc = Document(str(filled))
            _ensure_first_page_break_before(appendix_doc)
            appendix_doc.save(str(filled))
            composer.append(appendix_doc)
        composer.save(str(spec_path))
    finally:
        for tmp_file in tmp_files:
            tmp_file.unlink(missing_ok=True)
