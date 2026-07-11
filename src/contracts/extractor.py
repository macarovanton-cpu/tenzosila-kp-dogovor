"""
extractor.py — извлечение текста из PDF/DOCX и данных через OpenRouter API
"""

import json
import re
from pathlib import Path

import pdfplumber
import streamlit as st
from docx import Document
from openai import OpenAI

from src.contracts.requisites_extract import MIN_NONSPACE_CHARS, NoTextLayerError

PROMPT_PATH = Path(__file__).parent / "prompts" / "extract_contract_data.txt"
CARD_PROMPT_PATH = Path(__file__).parent / "prompts" / "extract_card_data.txt"


def _require_text(
    text: str, doc_name: str, min_nonspace: int = MIN_NONSPACE_CHARS
) -> None:
    """Пустой/почти пустой текст (скан без текстового слоя) → громкая ошибка
    ДО вызова LLM: на пустом входе модель фантазирует правдоподобные данные.
    Для DOCX передавать min_nonspace=1: текстовый слой там есть по построению,
    короткий валидный документ легитимен (как в requisites_extract)."""
    if len("".join(text.split())) < min_nonspace:
        raise NoTextLayerError(
            f"{doc_name}: текст не извлечён — пустой документ или скан без "
            "текстового слоя. Распознавание сканов не поддерживается."
        )


def extract_pdf_text(pdf_path: str, pages: list[int] = None) -> str:
    """
    Извлекает текст из PDF.
    pages — список номеров страниц (с 1). Если None — все страницы.
    """
    with pdfplumber.open(pdf_path) as pdf:
        total = len(pdf.pages)
        if pages is None:
            pages = list(range(1, total + 1))

        parts = []
        for page_num in pages:
            if 1 <= page_num <= total:
                text = pdf.pages[page_num - 1].extract_text()
                if text:
                    parts.append(f"=== СТРАНИЦА {page_num} ===\n{text}")

        return "\n\n".join(parts)


def extract_docx_text(docx_path: str) -> str:
    """
    Извлекает текст из Word-документа (карточка контрагента).
    Обрабатывает таблицы и параграфы.
    """
    doc = Document(docx_path)
    parts = []

    # Таблицы (основная часть карточки)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            # Убираем дубли (merged cells) и пустые
            unique = []
            seen = set()
            for c in cells:
                if c and c not in seen:
                    unique.append(c)
                    seen.add(c)
            if unique:
                parts.append(" | ".join(unique))

    # Отдельные параграфы (если есть)
    for p in doc.paragraphs:
        t = p.text.strip()
        if t and t not in parts:
            parts.append(t)

    return "\n".join(parts)


def extract_data_via_ai(kp_text: str, card_text: str) -> dict:
    """
    Отправляет тексты в OpenRouter API, получает JSON с данными для договора.
    Возвращает dict с ключами 'requisites' и 'specification'.
    """
    with open(PROMPT_PATH, encoding='utf-8') as f:
        system_prompt = f.read()

    client = OpenAI(
        base_url="https://openrouter.ai/api/v1",
        api_key=st.secrets["OPENROUTER_API_KEY"],
    )

    user_message = (
        f"КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ:\n{kp_text}"
        f"\n\n---\n\nКАРТОЧКА КОНТРАГЕНТА:\n{card_text}"
    )

    response = client.chat.completions.create(
        model="qwen/qwen3-235b-a22b",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_message},
        ],
    )
    raw = response.choices[0].message.content.strip()

    # Убираем markdown-обёртку если есть
    raw = re.sub(r'^```json\s*', '', raw)
    raw = re.sub(r'\s*```$', '', raw)
    raw = raw.strip()

    return json.loads(raw)


def extract_kp_data_legacy(kp_path: str, card_path: str) -> dict:
    """
    Legacy путь: принимает пути к КП (PDF) и карточке, возвращает dict
    с 'requisites' и 'specification'. Используется только в режиме B.
    card_path может быть .docx или .pdf.
    """
    kp_text = extract_pdf_text(kp_path)
    _require_text(kp_text, "КП")
    if card_path.lower().endswith('.pdf'):
        card_text = extract_pdf_text(card_path)
        _require_text(card_text, "Карточка контрагента")
    else:
        card_text = extract_docx_text(card_path)
        _require_text(card_text, "Карточка контрагента", min_nonspace=1)
    return extract_data_via_ai(kp_text, card_text)


def extract_card_data(card_path: str) -> dict:
    """
    Парсит только карточку контрагента через AI.
    Возвращает {"requisites": {...}} с 19 полями ЗАКАЗЧИК_*.
    Используется в режиме A (КП из базы).
    card_path может быть .docx или .pdf.
    """
    if card_path.lower().endswith('.pdf'):
        card_text = extract_pdf_text(card_path)
        _require_text(card_text, "Карточка контрагента")
    else:
        card_text = extract_docx_text(card_path)
        _require_text(card_text, "Карточка контрагента", min_nonspace=1)

    with open(CARD_PROMPT_PATH, encoding='utf-8') as f:
        system_prompt = f.read()

    client = OpenAI(
        base_url="https://openrouter.ai/api/v1",
        api_key=st.secrets["OPENROUTER_API_KEY"],
    )
    response = client.chat.completions.create(
        model="qwen/qwen3-235b-a22b",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"КАРТОЧКА КОНТРАГЕНТА:\n{card_text}"},
        ],
    )
    raw = response.choices[0].message.content.strip()
    raw = re.sub(r'^```json\s*', '', raw)
    raw = re.sub(r'\s*```$', '', raw)
    return json.loads(raw.strip())


# Алиас для обратной совместимости
extract_from_files = extract_kp_data_legacy
