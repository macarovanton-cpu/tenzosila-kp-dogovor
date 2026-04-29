#!/usr/bin/env python
"""
make_template.py — собирает kp_template.docx из эталонного КП Гипсобетон.

Идемпотентен: при повторном запуске перезаписывает templates/kp_template.docx
с тем же результатом.

Запуск:
    python src/generators/make_template.py
"""
import copy
import os
from lxml import etree
from docx import Document

BASE = os.path.normpath(os.path.join(os.path.dirname(os.path.abspath(__file__)), "../.."))
SRC  = os.path.join(BASE, "03_knowledge_base/sample_kps/Гипсобетон_ВЕСТА-С-80-18.docx")
DST  = os.path.join(BASE, "templates/kp_template.docx")

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def qn(tag):
    return f"{{{W}}}{tag}"


# ---------------------------------------------------------------------------
# Run helpers
# ---------------------------------------------------------------------------

def get_para_text(para):
    return "".join(r.text or "" for r in para.runs)


def merge_runs(para):
    """Merge all runs in paragraph into the first; keep first run's formatting."""
    runs = para.runs
    if len(runs) <= 1:
        return
    full = "".join(r.text or "" for r in runs)
    runs[0].text = full
    for run in runs[1:]:
        run._r.getparent().remove(run._r)


def set_para_text(para, text):
    """Merge runs, then set text of the single remaining run."""
    merge_runs(para)
    if para.runs:
        para.runs[0].text = text
    else:
        from docx.oxml import OxmlElement
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = text
        r.append(t)
        para._p.append(r)


def set_cell_text(cell, text: str) -> None:
    """Перезаписать содержимое ячейки многострочным текстом (\\n → <w:br/>).

    Сохраняет форматирование первого run первого параграфа, удаляет
    остальные параграфы, остальные строки вставляет line break'ами в
    тот же первый run.
    """
    from docx.oxml import OxmlElement
    paras = cell.paragraphs
    if not paras:
        return
    first_para = paras[0]
    merge_runs(first_para)
    for p in paras[1:]:
        p._p.getparent().remove(p._p)
    if not first_para.runs:
        first_para.add_run("")
    first_run = first_para.runs[0]
    lines = text.split("\n")
    first_run.text = lines[0]
    for line in lines[1:]:
        br = OxmlElement("w:br")
        first_run._r.append(br)
        t = OxmlElement("w:t")
        t.text = line
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        first_run._r.append(t)


# ---------------------------------------------------------------------------
# Footer placeholders (manager_*)
# ---------------------------------------------------------------------------

def _replace_chunk_text(chunk: list, new_text: str) -> None:
    """Сжать список runs в первый, заменить текст, удалить остальные."""
    if not chunk:
        return
    chunk[0].text = new_text
    for r in chunk[1:]:
        r._r.getparent().remove(r._r)


def _split_runs_by_breaks(para) -> list[list]:
    """Разбить runs параграфа на чанки по соседним '\\n'-runs (line breaks)."""
    chunks: list[list] = [[]]
    for r in para.runs:
        # python-docx возвращает '\n' для <w:br/> элементов
        if r.text == "\n":
            chunks.append([])
        else:
            chunks[-1].append(r)
    return [c for c in chunks if c]


def replace_footer_placeholders(doc) -> None:
    """Заменить статические данные менеджера в колонтитуле на плейсхолдеры.

    В эталонном Гипсобетон-КП футер представлен таблицей: cell[0] содержит
    'Макаров Антон\\nМенеджер', cell[1] — телефон\\nemail. Между строками —
    line break (<w:br/>), который python-docx показывает как run.text == '\\n'.
    Заменяем чанки между линейными разрывами, line break'ы сохраняем.
    """
    line_replacements = {
        "Макаров Антон": "{{ manager_full_name }}",
        "+7 903 651-85-77": "{{ manager_phone }}",
        "a.makarov@tenzosila.ru": "{{ manager_email }}",
    }

    for section in doc.sections:
        footer = section.footer
        if footer is None:
            continue
        # Параграфы напрямую в footer
        for para in footer.paragraphs:
            _apply_chunk_replacements(para, line_replacements)
        # Таблицы в footer
        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _apply_chunk_replacements(para, line_replacements)


def _apply_chunk_replacements(para, line_replacements: dict[str, str]) -> None:
    """Найти чанки runs (между line break'ами), сопоставить с line_replacements
    и заменить текст на плейсхолдер."""
    chunks = _split_runs_by_breaks(para)
    if not chunks:
        return
    for chunk in chunks:
        full = "".join(r.text or "" for r in chunk)
        if not full:
            continue
        for old, placeholder in line_replacements.items():
            if old in full:
                _replace_chunk_text(chunk, full.replace(old, placeholder))
                break


# ---------------------------------------------------------------------------
# ОРИОН block removal
# ---------------------------------------------------------------------------

def remove_orion_block(doc):
    """
    Удаляет блок 'Спецификация ПАК ОРИОН' (заголовок + таблица + сноска).
    Сохраняет заголовок 'Конструкция весов ВЕСТА' и всё после него.
    """
    body = doc.element.body
    children = list(body)

    start_idx = None
    end_idx   = None

    for i, elem in enumerate(children):
        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if tag not in ("p", "tbl"):
            continue
        text = "".join(t.text or "" for t in elem.findall(f".//{qn('t')}"))

        if start_idx is None and "Спецификация ПАК ОРИОН" in text:
            start_idx = i

        if start_idx is not None and "* Цены и сроки подлежат уточнению" in text:
            end_idx = i
            break

    if start_idx is None or end_idx is None:
        print(f"WARNING: ОРИОН-блок не найден (start={start_idx}, end={end_idx})")
        return

    to_remove = children[start_idx : end_idx + 1]
    for elem in to_remove:
        body.remove(elem)
    print(f"  Удалено {len(to_remove)} элементов (блок ОРИОН)")


# ---------------------------------------------------------------------------
# Spec table → jinja loop
# ---------------------------------------------------------------------------

def _extract_first_rpr(tc):
    """Найти rPr первого run-а в tc, у которого rPr задано. None — иначе."""
    for r in tc.findall(f'.//{qn("r")}'):
        rpr = r.find(qn("rPr"))
        if rpr is not None:
            return copy.deepcopy(rpr)
    return None


def _force_pt_sans_on_run(r_elem, sz_halfpt: int = 22) -> None:
    """Принудительно установить PT Sans + размер на <w:r>.

    sz_halfpt в half-points: 22 = 11pt, 18 = 9pt, 16 = 8pt.
    rFonts ставим первым ребёнком rPr (по схеме OOXML), sz — после.
    Существующие значения перезаписываются, не дублируются.
    """
    from docx.oxml import OxmlElement
    rpr = r_elem.find(qn("rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        r_elem.insert(0, rpr)
    rfonts = rpr.find(qn("rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("ascii"), "PT Sans")
    rfonts.set(qn("hAnsi"), "PT Sans")
    rfonts.set(qn("cs"), "PT Sans")
    sz = rpr.find(qn("sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        rpr.append(sz)
    sz.set(qn("val"), str(sz_halfpt))


def _force_pt_sans_on_cell(tc, sz_halfpt: int = 22) -> None:
    """Применить _force_pt_sans_on_run ко всем <w:r> в ячейке."""
    for r in tc.findall(f'.//{qn("r")}'):
        _force_pt_sans_on_run(r, sz_halfpt=sz_halfpt)


def _set_tc_text(tc, text):
    """Установить текст первой ячейки <w:tc>, сохраняя rPr первого run-а.

    rPr (шрифт, цвет, размер) копируется с первого подходящего run-а ячейки
    и применяется к новому run-у. Это сохраняет PT Sans / 11pt / etc. вместо
    дефолтного шрифта Word после рендера.
    """
    from docx.oxml import OxmlElement
    saved_rpr = _extract_first_rpr(tc)
    ps = tc.findall(f'.//{qn("p")}')
    if not ps:
        return
    p = ps[0]
    for r in list(p.findall(qn("r"))):
        p.remove(r)
    r_el = OxmlElement('w:r')
    if saved_rpr is not None:
        r_el.insert(0, saved_rpr)  # rPr ДОЛЖЕН быть первым ребёнком <w:r>
    t_el = OxmlElement('w:t')
    t_el.text = text
    r_el.append(t_el)
    p.append(r_el)


def transform_spec_table(doc):
    """Заменяет 6 фиксированных строк спецификации на jinja-цикл {%tr%}.

    Итоговая структура таблицы в шаблоне:
      row 0: заголовок (header)
      row 1: {%tr for item in spec_items %}  ← маркер начала (удаляется при рендере)
      row 2: {{ item.name }} | {{ item.price }} | {{ item.term_days }}  ← шаблон (повторяется)
      row 3: {%tr endfor %}                  ← маркер конца (удаляется при рендере)
      row 4: ИТОГО | {{ total_price }}        ← статическая строка

    При рендере с N позициями → header + N строк + ИТОГО = N+2 строк.
    """
    from docx.oxml import OxmlElement

    spec_table = None
    for table in doc.tables:
        if table.rows and table.rows[0].cells:
            header_text = "".join(
                get_para_text(p) for p in table.rows[0].cells[0].paragraphs
            ).strip()
            if header_text == "Наименование":
                spec_table = table
                break
    if spec_table is None:
        raise RuntimeError("Таблица спецификации (заголовок 'Наименование') не найдена")

    # Клонируем row 1 ДО изменений — будет шаблоном контента
    content_tr = copy.deepcopy(spec_table.rows[1]._tr)

    # row 1 → маркер {%tr for %}: первая ячейка = тег, остальные — пусто
    for_row = spec_table.rows[1]
    tcs_for = for_row._tr.findall(qn("tc"))
    for i, tc in enumerate(tcs_for):
        _set_tc_text(tc, '{%tr for item in spec_items %}' if i == 0 else '')

    # content_tr → шаблон повторяемой строки
    tcs_content = content_tr.findall(qn("tc"))
    if len(tcs_content) >= 3:
        _set_tc_text(tcs_content[0], '{{ item.name }}')
        _set_tc_text(tcs_content[1], '{{ item.price }}')
        _set_tc_text(tcs_content[2], '{{ item.term_days }}')
        # Снимаем унаследованный <w:vMerge> со всех ячеек loop-template-row.
        # Иначе docxtpl клонирует строку с vMerge="restart" из эталона
        # на каждую отрисованную позицию, и spec_vmerge.py не сможет
        # отличить scales-row от foundation-row (обе имеют restart).
        for tc in tcs_content[:3]:
            tcPr = tc.find(qn("tcPr"))
            if tcPr is not None:
                for vm in tcPr.findall(qn("vMerge")):
                    tcPr.remove(vm)

    # endfor_tr → маркер {%tr endfor %}
    endfor_tr = copy.deepcopy(content_tr)
    for tc in endfor_tr.findall(qn("tc")):
        _set_tc_text(tc, '')
    endfor_tcs = endfor_tr.findall(qn("tc"))
    if endfor_tcs:
        _set_tc_text(endfor_tcs[0], '{%tr endfor %}')

    # Удаляем строки 2–6 (лишние позиции), строку 7 (ИТОГО) сохраняем
    rows_to_delete = list(spec_table.rows[2:7])
    for row in reversed(rows_to_delete):
        row._tr.getparent().remove(row._tr)
    # После удаления: rows[0]=header, rows[1]=for-marker, rows[2]=ИТОГО

    # Вставляем content_tr и endfor_tr перед ИТОГО
    itogo_row = spec_table.rows[2]
    itogo_row._tr.addprevious(endfor_tr)
    endfor_tr.addprevious(content_tr)
    # Итог: header | for-marker | content | endfor | ИТОГО

    print("  Таблица спецификации: transform OK (for-marker + content + endfor + ИТОГО)")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def make_template():
    print(f"Источник : {SRC}")
    print(f"Шаблон   : {DST}")

    doc = Document(SRC)

    # -----------------------------------------------------------------------
    # 1. ШАПКА — body paragraphs
    # -----------------------------------------------------------------------
    for para in doc.paragraphs:
        full = get_para_text(para)

        # client_name
        if "АО Гипсобетон" in full:
            merge_runs(para)
            para.runs[0].text = full.replace("АО Гипсобетон", "{{ client_name }}")

        # kp_number + kp_date (в одном абзаце)
        elif "Коммерческое предложение № 47141" in full:
            # НЕ merge_runs целиком — иначе пропадёт оранжевая «ВЕСТА»
            # (последний run с w:color=D04514). Сжимаем все runs ДО оранжевого
            # в первый, подставляем плейсхолдеры; оранжевый run не трогаем.
            runs = list(para.runs)
            orange_idx = None
            for i, r in enumerate(runs):
                rpr = r._r.find(qn("rPr"))
                if rpr is None:
                    continue
                color = rpr.find(qn("color"))
                if color is not None and color.get(qn("val")) == "D04514":
                    orange_idx = i
                    break
            if orange_idx is None:
                # Fallback: оранжевого run-а нет — обычный merge.
                merge_runs(para)
                para.runs[0].text = (get_para_text(para)
                    .replace("47141", "{{ kp_number }}")
                    .replace("22.04.2026", "{{ kp_date }}"))
            else:
                before_runs = runs[:orange_idx]
                if before_runs:
                    full_before = "".join((r.text or "") for r in before_runs)
                    new_text = (full_before
                        .replace("47141", "{{ kp_number }}")
                        .replace("22.04.2026", "{{ kp_date }}"))
                    before_runs[0].text = new_text
                    for r in before_runs[1:]:
                        r._r.getparent().remove(r._r)
                # runs[orange_idx] (оранжевая «ВЕСТА») остаётся нетронутой

        # vat_percent (body paragraph, not table)
        elif "НДС 22%" in full:
            merge_runs(para)
            para.runs[0].text = get_para_text(para).replace("22%", "{{ vat_percent }}%")

        # payment_terms_block (заменяет payment_line_1 в эталоне) — RichText с многострочным
        # содержимым, абзацные переносы добавляет docxtpl при рендере.
        # 1.5b-fix4: снять numPr — иначе Word добавляет list-маркер (тире) перед первой
        # строкой блока, а у нас уже есть ручной "— " из payment_renderer → двойное тире.
        elif full.startswith("Предоплата:"):
            merge_runs(para)
            para.runs[0].text = "{{ payment_terms_block }}"
            p_elem = para._p
            pPr = p_elem.find(qn("pPr"))
            if pPr is not None:
                numPr = pPr.find(qn("numPr"))
                if numPr is not None:
                    pPr.remove(numPr)

        # kp_valid_days — переиспользуем абзац "Доплата:" под строку срока действия КП.
        # Перед параграфом вставляем 2 пустых абзаца, текст делаем жирным.
        # Параграф «Доплата:» в эталоне унаследовал стиль списка (numPr) —
        # очищаем, чтобы Word не рисовал тире-маркер.
        elif full.startswith("Доплата:"):
            from docx.oxml import OxmlElement
            p_elem = para._p
            parent = p_elem.getparent()
            idx = list(parent).index(p_elem)
            for _ in range(2):
                empty_p = OxmlElement("w:p")
                parent.insert(idx, empty_p)
                idx += 1
            merge_runs(para)

            # Сменить стиль на Normal (если ListParagraph) и убрать numPr
            pPr = p_elem.find(qn("pPr"))
            if pPr is not None:
                numPr = pPr.find(qn("numPr"))
                if numPr is not None:
                    pPr.remove(numPr)
                pStyle = pPr.find(qn("pStyle"))
                if pStyle is not None:
                    style_val = pStyle.get(qn("val")) or ""
                    if "List" in style_val:
                        pPr.remove(pStyle)
            try:
                para.style = doc.styles["Normal"]
            except KeyError:
                pass

            run = para.runs[0]
            run.text = (
                "Срок действия настоящего коммерческого предложения — "
                "{{ kp_valid_days }}."
            )
            run.bold = True
            # После сброса стиля в Normal дефолтный шрифт — Times New Roman.
            # Явно ставим rFonts=PT Sans в rPr, чтобы не зависеть от стиля абзаца.
            r_elem = run._r
            rpr = r_elem.find(qn("rPr"))
            if rpr is None:
                rpr = OxmlElement("w:rPr")
                r_elem.insert(0, rpr)
            rfonts = rpr.find(qn("rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.insert(0, rfonts)  # rFonts — первый ребёнок rPr (схема OOXML)
            rfonts.set(qn("ascii"), "PT Sans")
            rfonts.set(qn("hAnsi"), "PT Sans")
            rfonts.set(qn("cs"), "PT Sans")

        # Заголовок раздела «Характеристики ВЕСТА...» → плейсхолдер с моделью
        elif full.strip().startswith("Характеристики ВЕСТА"):
            import re
            merge_runs(para)
            para.runs[0].text = re.sub(
                r"Характеристики ВЕСТА[\w\-]*",
                "Характеристики {{ model_full_name }}",
                get_para_text(para),
            )

    # -----------------------------------------------------------------------
    # 2. ТАБЛИЦЫ
    # -----------------------------------------------------------------------
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if not cells:
                continue

            # Текст первой ячейки строки — для контекстного матчинга
            first_cell_text = "".join(
                get_para_text(p) for p in cells[0].paragraphs
            ).strip()

            for ci, cell in enumerate(cells):
                paras     = cell.paragraphs
                cell_text = "".join(get_para_text(p) for p in paras).strip()

                # --- warranty_text ---
                if cell_text == "36 месяца":
                    set_para_text(paras[0], "{{ warranty_text }}")

                # --- platform_size ---
                elif cell_text == "18×3":
                    set_para_text(paras[0], "{{ platform_size }}")

                # --- construction_description ---
                elif cell_text.startswith("Конструкция сплошная 09Г2С"):
                    set_para_text(paras[0], "{{ construction_description }}")


                # --- vat_percent ---
                elif "НДС 22%" in cell_text:
                    merge_runs(paras[0])
                    if paras[0].runs:
                        paras[0].runs[0].text = get_para_text(paras[0]).replace(
                            "22%", "{{ vat_percent }}%"
                        )

            # --- total_price + total_term_days (ИТОГО — split runs) ---
            if first_cell_text == "ИТОГО" and len(cells) >= 2:
                val_paras = cells[1].paragraphs
                for p in val_paras:
                    merge_runs(p)
                set_para_text(val_paras[0], "{{ total_price }}")
            if first_cell_text == "ИТОГО" and len(cells) >= 3:
                term_paras = cells[2].paragraphs
                for p in term_paras:
                    merge_runs(p)
                set_para_text(term_paras[0], "{{ total_term_days }}")

    # -----------------------------------------------------------------------
    # 2b. Прицельная обработка таблицы ТХ (Tbl 2) — по индексам строк
    # -----------------------------------------------------------------------
    tx_table = doc.tables[2]

    # Sanity-check: подтвердить ожидаемую структуру эталона
    _expected_labels = {
        1: "Рабочий диапазон температур",
        9: "Описание весов",
        10: "Максимальная нагрузка",
        11: "Цена поверочного деления",
    }
    for r_idx, expected in _expected_labels.items():
        actual = tx_table.rows[r_idx].cells[0].text.strip()
        if expected not in actual:
            raise RuntimeError(
                f"Tbl 2 row {r_idx}: ожидался текст '{expected}', "
                f"в шаблоне '{actual[:80]}'. Структура эталона изменилась?"
            )

    # Row 1 — Рабочий диапазон температур (динамические датчик и терминал)
    set_cell_text(tx_table.rows[1].cells[0], (
        "Рабочий диапазон температур, °С:\n"
        "- для Тензодатчиков {{ sensor_label }}\n"
        "- для Весового индикатора (терминала) {{ indicator_label }}"
    ))
    set_cell_text(
        tx_table.rows[1].cells[2],
        "{{ sensor_temp_range }}\n{{ indicator_temp_range }}",
    )
    # Принудительно PT Sans 11pt на runs обеих ячеек: первый run в эталоне
    # не имеет явных rFonts/sz → Word рендерит Calibri/11pt по дефолту.
    _force_pt_sans_on_cell(tx_table.rows[1].cells[0]._tc, sz_halfpt=22)
    _force_pt_sans_on_cell(tx_table.rows[1].cells[2]._tc, sz_halfpt=22)

    # Row 9 — Описание весов: подпись восстанавливаем, значение очищаем
    set_cell_text(tx_table.rows[9].cells[0], "Описание весов")
    set_cell_text(tx_table.rows[9].cells[2], "")

    # Row 10 — Максимальная нагрузка: подпись + плейсхолдер в [2]
    set_cell_text(tx_table.rows[10].cells[0], "Максимальная нагрузка, т")
    set_cell_text(tx_table.rows[10].cells[2], "{{ max_load_t }}")

    # Row 11 — Цена поверочного деления: подпись + плейсхолдер в [2].
    # Единица измерения «кг» теперь в значении (см. build_main_scale_label).
    set_cell_text(tx_table.rows[11].cells[0], "Цена поверочного деления:")
    set_cell_text(tx_table.rows[11].cells[2], "{{ main_scale_label }}")

    # -----------------------------------------------------------------------
    # 3. Таблица спецификации → jinja-цикл
    # -----------------------------------------------------------------------
    transform_spec_table(doc)

    # -----------------------------------------------------------------------
    # 4. Удаляем блок ОРИОН
    # -----------------------------------------------------------------------
    remove_orion_block(doc)

    # -----------------------------------------------------------------------
    # 4b. Подставляем плейсхолдеры менеджера в колонтитул
    # -----------------------------------------------------------------------
    replace_footer_placeholders(doc)

    # -----------------------------------------------------------------------
    # 5. Сохраняем
    # -----------------------------------------------------------------------
    os.makedirs(os.path.dirname(DST), exist_ok=True)
    doc.save(DST)
    print(f"  Шаблон сохранён: {DST}")

    # -----------------------------------------------------------------------
    # 6. Sanity-check: 18 статических body + 3 footer + 3 loop-плейсхолдера
    # -----------------------------------------------------------------------
    import re, zipfile
    expected_static_body = {
        "client_name", "kp_number", "kp_date", "kp_valid_days",
        "warranty_text", "platform_size", "max_load_t",
        "construction_description", "main_scale_label",
        "total_price", "total_term_days", "vat_percent",
        "payment_terms_block",
        "model_full_name",
        "sensor_label", "sensor_temp_range",
        "indicator_label", "indicator_temp_range",
    }
    expected_static_footer = {
        "manager_full_name", "manager_phone", "manager_email",
    }
    expected_static = expected_static_body | expected_static_footer
    saved = Document(DST)
    body_text = " ".join(
        "".join(r.text or "" for r in p.runs)
        for p in saved.paragraphs
    )
    for table in saved.tables:
        for row in table.rows:
            for cell in row.cells:
                body_text += " " + "".join(
                    "".join(r.text or "" for r in p.runs)
                    for p in cell.paragraphs
                )
    footer_text = ""
    for section in saved.sections:
        footer = section.footer
        if footer is None:
            continue
        for p in footer.paragraphs:
            footer_text += " " + "".join(r.text or "" for r in p.runs)
        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        footer_text += " " + "".join(r.text or "" for r in p.runs)

    all_text = body_text + " " + footer_text
    found_static = set(re.findall(r"\{\{\s*(\w+)\s*\}\}", all_text))
    problems = []
    missing_static = expected_static - found_static
    if missing_static:
        problems.append(f"Отсутствуют статические плейсхолдеры: {sorted(missing_static)}")
    for loop_str in ("{{ item.name }}", "{{ item.price }}", "{{ item.term_days }}"):
        if loop_str not in all_text:
            problems.append(f"Отсутствует loop-плейсхолдер: {loop_str}")
    with zipfile.ZipFile(DST) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    if "{%tr for" not in xml:
        problems.append("{%tr for} не найден в XML")
    if "{%tr endfor" not in xml:
        problems.append("{%tr endfor} не найден в XML")
    # 1.5b-fix3: spec-row run должен иметь rPr (PT Sans), не пустой <w:r>
    name_idx = xml.find("{{ item.name }}")
    if name_idx > 0:
        # Ищем именно открывающий <w:r> или <w:r ...>, не <w:rPr>/<w:rFonts>.
        run_start = max(xml.rfind("<w:r>", 0, name_idx), xml.rfind("<w:r ", 0, name_idx))
        run_end = xml.find("</w:r>", name_idx) + len("</w:r>")
        run_xml = xml[run_start:run_end]
        if "<w:rPr>" not in run_xml:
            problems.append("Run с {{ item.name }} без rPr — шрифты будут дефолтными")
        elif 'w:ascii="PT Sans"' not in run_xml:
            problems.append("Run с {{ item.name }} без PT Sans в rPr")
    # 1.5b-fix3: оранжевая «ВЕСТА» должна быть в параграфе с {{ kp_number }}
    kp_idx = xml.find("{{ kp_number }}")
    if kp_idx > 0:
        para_start = xml.rfind("<w:p ", 0, kp_idx)
        para_end = xml.find("</w:p>", kp_idx) + len("</w:p>")
        para_xml = xml[para_start:para_end]
        if 'w:val="D04514"' not in para_xml:
            problems.append("Оранжевый цвет D04514 в заголовке КП пропал")
    if problems:
        raise RuntimeError("Sanity-check FAILED:\n" + "\n".join(f"  - {p}" for p in problems))
    n_static = len(expected_static)
    print(f"  Проверка: {n_static} статических + 3 loop-плейсхолдера на месте [OK]")


if __name__ == "__main__":
    make_template()
