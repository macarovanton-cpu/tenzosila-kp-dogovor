#!/usr/bin/env python
"""test_generate.py — генерирует тестовые КП из шаблона для ручной проверки в Word.

Использует тот же путь, что и UI: build_template_context(state, prices) →
DocxTemplate.render(). Это даёт уверенность, что DOCX-генерация согласована
с реальным state менеджера.

Кейсы:
  1. Гипсобетон   — ВЕСТА-С-80-18, dual_range, фундамент, split_by_items
  2. Кирова       — ВЕСТА-ФЛ-80-18, single, без фундамента, split_by_items
  3. Стресс-тест  — ВЕСТА-С-100-24, dual_range, все опции, custom

Запуск:
    python src/generators/test_generate.py
Результаты сохраняются в output/.
"""
from __future__ import annotations

import os
from datetime import date
from pathlib import Path

from docx import Document

from src.data_loader import load_prices
from src.generators.kp_generator import build_filename, generate_kp

BASE: Path = Path(__file__).resolve().parent.parent.parent
OUT_DIR: Path = BASE / "output"


def _gipsobeton_state() -> dict:
    """Большая сделка: dual_range, фундамент, ОРИОН, split_by_items с правкой."""
    return {
        "kp_number": "47141",
        "kp_date": date(2026, 4, 22),
        "kp_valid_days": 15,
        "total_term_days": 35,
        "manager_id": "makarov_av",
        "client_name": "АО «Гипсобетон»",
        "model_line": "С",
        "model_max": 80,
        "model_length": 18,
        "model_id": "vesta-с-80-18",
        "model_price": 2_450_000,
        "warranty_months": 36,
        "is_dual_range": True,
        "construction_beam": "Двутавр 30Б1",
        "construction_beam_count": 8,
        "construction_center_beam": "Швеллер №12",
        "construction_center_beam_count": 2,
        "construction_deck_mm": 10,
        "construction_underlining_mm": 4,
        "options": {
            "foundation_s_f_18": {
                "enabled": True, "price": 1_900_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 1_900_000, "dealer_is_synthetic": False,
                "block": "foundations",
            },
            "orion_standard_plus": {
                "enabled": True, "price": 1_200_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 1_200_000, "dealer_is_synthetic": False,
                "block": "pak_orion",
            },
            "install_default": {
                "enabled": True, "price": 250_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 250_000, "dealer_is_synthetic": False,
                "block": "install",
            },
            "delivery_default": {
                "enabled": True, "price": 70_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 70_000, "dealer_is_synthetic": False,
                "block": "delivery",
            },
            "verification_default": {
                "enabled": True, "price": 60_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 60_000, "dealer_is_synthetic": False,
                "block": "verification",
            },
        },
        "spec_items_overrides": {},
        "payment_preset_id": "split_by_items",
        "payment_split_state": {},
        "payment_percents": {},
        "payment_days": 5,
        "payment_custom_text": "",
        "payment_v1_prepay": 50,
        "payment_v2_prepay": 30,
        "payment_v2_preship": 40,
        "payment_v3_days": 15,
        "payment_v3_trigger_id": "after_installation",
    }


def _kirova_state() -> dict:
    """Без фундамента, без ОРИОН, поверка силами заказчика."""
    return {
        "kp_number": "47215",
        "kp_date": date(2026, 4, 22),
        "kp_valid_days": 15,
        "total_term_days": 30,
        "manager_id": "makarov_av",
        "client_name": "АО «Совхоз имени Кирова»",
        "model_line": "ФЛ",
        "model_max": 80,
        "model_length": 18,
        "model_id": "vesta-фл-80-18",
        "model_price": 1_800_000,
        "warranty_months": 24,
        "is_dual_range": False,
        "construction_beam": "Двутавр 25Б1",
        "construction_beam_count": 8,
        "construction_center_beam": "",
        "construction_center_beam_count": 0,
        "construction_deck_mm": 8,
        "construction_underlining_mm": 3,
        "options": {
            "install_default": {
                "enabled": True, "price": 200_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 200_000, "dealer_is_synthetic": False,
                "block": "install",
            },
            "delivery_default": {
                "enabled": True, "price": 75_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 75_000, "dealer_is_synthetic": False,
                "block": "delivery",
            },
        },
        "spec_items_overrides": {},
        # Кейс v0.4: showcase V1 (Аванс+Постоплата) — менеджер ввёл 30/70.
        "payment_preset_id": "v1_prepay_postpay",
        "payment_split_state": {},
        "payment_percents": {},
        "payment_days": 5,
        "payment_custom_text": "",
        "payment_v1_prepay": 30,
        "payment_v2_prepay": 30,
        "payment_v2_preship": 40,
        "payment_v3_days": 15,
        "payment_v3_trigger_id": "after_installation",
    }


def _stress_state() -> dict:
    """Максимум: 100т 24м dual_range, custom-оплата."""
    return {
        "kp_number": "99999",
        "kp_date": date(2026, 4, 23),
        "kp_valid_days": 21,
        "total_term_days": 60,
        "manager_id": "makarov_av",
        "client_name": "ООО «Стресс-тест Макс»",
        "model_line": "С",
        "model_max": 100,
        "model_length": 24,
        "model_id": "vesta-с-100-24",
        "model_price": 3_200_000,
        "warranty_months": 36,
        "is_dual_range": True,
        "construction_beam": "Двутавр 35Б1",
        "construction_beam_count": 10,
        "construction_center_beam": "Швеллер №14",
        "construction_center_beam_count": 2,
        "construction_deck_mm": 12,
        "construction_underlining_mm": 5,
        "options": {
            "foundation_s_f_24": {
                "enabled": True, "price": 2_800_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 2_800_000, "dealer_is_synthetic": False,
                "block": "foundations",
            },
            "orion_auto_plus": {
                "enabled": True, "price": 1_850_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 1_850_000, "dealer_is_synthetic": False,
                "block": "pak_orion",
            },
            "install_default": {
                "enabled": True, "price": 380_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 380_000, "dealer_is_synthetic": False,
                "block": "install",
            },
            "delivery_default": {
                "enabled": True, "price": 90_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 90_000, "dealer_is_synthetic": False,
                "block": "delivery",
            },
            "verification_default": {
                "enabled": True, "price": 70_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 70_000, "dealer_is_synthetic": False,
                "block": "verification",
            },
        },
        "spec_items_overrides": {},
        "payment_preset_id": "custom",
        "payment_split_state": {},
        "payment_percents": {},
        "payment_days": 5,
        "payment_custom_text": (
            "— 30% предоплаты в течение 10 банковских дней с момента подписания.\n"
            "— 40% после готовности фундамента.\n"
            "— 30% по факту запуска весов и подписания акта."
        ),
        "payment_v1_prepay": 50,
        "payment_v2_prepay": 30,
        "payment_v2_preship": 40,
        "payment_v3_days": 15,
        "payment_v3_trigger_id": "after_installation",
    }


CASES = [
    ("gipsobeton", _gipsobeton_state),
    ("kirova", _kirova_state),
    ("stress_max", _stress_state),
]


def _doc_text(doc) -> str:
    parts: list[str] = []
    for p in doc.paragraphs:
        parts.append("".join(r.text or "" for r in p.runs))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append("".join(r.text or "" for r in p.runs))
    for section in doc.sections:
        f = section.footer
        if f is None:
            continue
        for p in f.paragraphs:
            parts.append("".join(r.text or "" for r in p.runs))
        for t in f.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        parts.append("".join(r.text or "" for r in p.runs))
    return " ".join(parts)


def generate_case(name: str, state_factory) -> Path:
    state = state_factory()
    prices = load_prices()
    docx_bytes = generate_kp(state, prices)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out_path = OUT_DIR / f"КП_тест_{name}.docx"
    out_path.write_bytes(docx_bytes)

    # --- автопроверки ---
    doc = Document(str(out_path))
    text = _doc_text(doc)
    assert "{{" not in text, f"[{name}] Остались незаполненные плейсхолдеры {{{{ }}}}"
    assert "{%" not in text, f"[{name}] Остались jinja-теги {{%}}"

    # Имя файла строится по тем же правилам, что и в sidebar
    expected_filename = build_filename(state)
    assert expected_filename.endswith(".docx"), f"[{name}] Неверное имя: {expected_filename}"

    return out_path


def main() -> None:
    for name, factory in CASES:
        out = generate_case(name, factory)
        print(f"  Сохранён: {out}")
    print("Готово. Откройте файлы в output/ для ручной проверки.")


if __name__ == "__main__":
    main()
