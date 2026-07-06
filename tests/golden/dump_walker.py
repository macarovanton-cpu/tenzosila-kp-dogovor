"""Рекурсивный обход OOXML для канонического дампа DOCX (P0-01).

Единый walker для body / header / footer / w:txbxContent (textbox).
Обход строго в document order (python-docx paragraphs/tables теряют
взаимный порядок — поэтому прямая итерация XML-детей).

Ключевая нормализация: соседние раны с одинаковой ЗНАЧИМОЙ rPr-сводкой
склеиваются в одну строку — Word/docxtpl дробят раны произвольно
(rsid/proofErr-границы), без склейки дифф эталонов шумел бы ложными
расхождениями. w:tab/w:br кодируются как \\t/\\n внутри текста рана
(склейку не рвут); разрыв страницы — отдельная строка [PAGEBREAK].
"""

from __future__ import annotations

from dataclasses import dataclass, field

from lxml import etree

from docx.oxml.ns import qn

from tests.golden.dump_props import (
    ppr_summary,
    rpr_summary,
    sectpr_summary,
    tbl_summary,
    tcpr_summary,
    trpr_summary,
)

_IND = "  "
_MC = "{http://schemas.openxmlformats.org/markup-compatibility/2006}"


@dataclass
class DumpContext:
    """Сквозное состояние дампа: имена стилей, num-алиасы, использованные стили."""

    style_names: dict[str, str] = field(default_factory=dict)
    num_aliases: dict[str, str] = field(default_factory=dict)
    used_style_ids: set[str] = field(default_factory=set)


def _register_style(pr: etree._Element | None, tag: str, ctx: DumpContext) -> None:
    """Запомнить styleId (pStyle/rStyle) для секции [STYLE] шапки дампа."""
    if pr is None:
        return
    el = pr.find(qn(tag))
    if el is not None and el.get(qn("w:val")):
        ctx.used_style_ids.add(el.get(qn("w:val")))


def walk_block(parent: etree._Element, ctx: DumpContext, level: int = 0) -> list[str]:
    """Обойти block-level детей (w:p / w:tbl / w:sectPr) в document order."""
    lines: list[str] = []
    for child in parent:
        tag = child.tag
        if tag == qn("w:p"):
            lines.extend(_walk_paragraph(child, ctx, level))
        elif tag == qn("w:tbl"):
            lines.extend(_walk_table(child, ctx, level))
        elif tag == qn("w:sectPr"):
            lines.append(_IND * level + "SECT" + sectpr_summary(child))
    return lines


def _walk_paragraph(p_el: etree._Element, ctx: DumpContext, level: int) -> list[str]:
    ppr = p_el.find(qn("w:pPr"))
    _register_style(ppr, "w:pStyle", ctx)
    lines = [_IND * level + "P" + ppr_summary(ppr, ctx.style_names, ctx.num_aliases)]
    runs = _RunWalker(ctx, level + 1)
    for child in p_el:
        tag = child.tag
        if tag == qn("w:r"):
            runs.add_run(child)
        elif tag == qn("w:hyperlink"):
            # гиперссылка прозрачна: дампим её раны, rId не дампим
            for r_el in child.findall(qn("w:r")):
                runs.add_run(r_el)
        elif tag == qn("w:fldSimple"):
            runs.add_fld_simple(child)
    runs.flush()
    lines.extend(runs.lines)
    if ppr is not None:
        sect_pr = ppr.find(qn("w:sectPr"))
        if sect_pr is not None:
            # разрыв секции посреди документа живёт в pPr параграфа
            lines.append(_IND * level + "SECT" + sectpr_summary(sect_pr))
    return lines


def _walk_table(tbl: etree._Element, ctx: DumpContext, level: int) -> list[str]:
    lines = [_IND * level + "TBL" + tbl_summary(tbl)]
    for tr in tbl.findall(qn("w:tr")):
        lines.append(_IND * (level + 1) + "TR" + trpr_summary(tr))
        for tc in tr.findall(qn("w:tc")):
            lines.append(_IND * (level + 2) + "TC" + tcpr_summary(tc))
            # вложенные таблицы поддержаны рекурсией walk_block
            lines.extend(walk_block(tc, ctx, level + 3))
    return lines


class _RunWalker:
    """Накопитель ранов одного параграфа со склейкой по значимой rPr-сводке."""

    def __init__(self, ctx: DumpContext, level: int) -> None:
        self.ctx = ctx
        self.level = level
        self.lines: list[str] = []
        self._summary: str | None = None
        self._text: list[str] = []
        self._field: dict | None = None  # состояние w:fldChar begin..end

    @property
    def _prefix(self) -> str:
        return _IND * self.level

    def flush(self) -> None:
        """Выдать накопленный ран одной строкой (пустой текст не эмитится)."""
        joined = "".join(self._text)
        if joined:
            self.lines.append(f"{self._prefix}R{self._summary} {joined!r}")
        self._summary = None
        self._text = []

    def _push(self, summary: str, text: str) -> None:
        if summary != self._summary:
            self.flush()
            self._summary = summary
        self._text.append(text)

    def add_run(self, r_el: etree._Element) -> None:
        rpr = r_el.find(qn("w:rPr"))
        _register_style(rpr, "w:rStyle", self.ctx)
        summary = rpr_summary(rpr, self.ctx.style_names)
        for el in r_el:
            tag = el.tag
            if tag == qn("w:fldChar"):
                self._on_fld_char(el, summary)
            elif tag == qn("w:instrText"):
                if self._field is not None:
                    self._field["instr"].append(el.text or "")
            elif tag == qn("w:t"):
                text = el.text or ""
                if self._field is not None and self._field["stage"] == "cached":
                    self._field["cached"].append(text)
                else:
                    self._push(summary, text)
            elif tag == qn("w:tab"):
                self._push(summary, "\t")
            elif tag in (qn("w:br"), qn("w:cr")):
                if el.get(qn("w:type")) == "page":
                    self.flush()
                    self.lines.append(f"{self._prefix}[PAGEBREAK]")
                else:
                    self._push(summary, "\n")
            elif tag == qn("w:drawing"):
                self._emit_object(el, summary, is_drawing=True)
            elif tag in (qn("w:pict"), qn("w:object")):
                self._emit_object(el, summary, is_drawing=False)
            elif tag == _MC + "AlternateContent":
                self._on_alternate(el, summary)

    def _on_alternate(self, el: etree._Element, summary: str) -> None:
        """mc:AlternateContent: дампим только mc:Choice (Fallback — дубль
        того же содержимого в VML, иначе текст textbox удвоится)."""
        choice = el.find(_MC + "Choice")
        target = choice.find(qn("w:drawing")) if choice is not None else None
        if target is not None:
            self._emit_object(target, summary, is_drawing=True)
            return
        fallback = el.find(_MC + "Fallback")
        target = fallback.find(qn("w:pict")) if fallback is not None else None
        if target is not None:
            self._emit_object(target, summary, is_drawing=False)

    def add_fld_simple(self, fld: etree._Element) -> None:
        """w:fldSimple: инструкция в атрибуте, кэш — вложенные раны."""
        self.flush()
        first_run = fld.find(qn("w:r"))
        rpr = first_run.find(qn("w:rPr")) if first_run is not None else None
        summary = rpr_summary(rpr, self.ctx.style_names)
        instr = (fld.get(qn("w:instr")) or "").strip()
        cached = "".join(t.text or "" for t in fld.findall(".//" + qn("w:t")))
        self.lines.append(f"{self._prefix}R{summary} [FIELD {instr!r}] cached={cached!r}")

    def _on_fld_char(self, el: etree._Element, summary: str) -> None:
        fld_type = el.get(qn("w:fldCharType"))
        if fld_type == "begin":
            self.flush()
            self._field = {"instr": [], "cached": [], "stage": "instr", "summary": summary}
        elif fld_type == "separate" and self._field is not None:
            self._field["stage"] = "cached"
        elif fld_type == "end" and self._field is not None:
            instr = "".join(self._field["instr"]).strip()
            cached = "".join(self._field["cached"])
            self.lines.append(
                f"{self._prefix}R{self._field['summary']} [FIELD {instr!r}] cached={cached!r}"
            )
            self._field = None

    def _emit_object(self, el: etree._Element, summary: str, is_drawing: bool) -> None:
        """DRAWING/PICT + рекурсия в textbox (w:txbxContent не виден python-docx API)."""
        self.flush()
        if is_drawing:
            extent = el.find(".//" + qn("wp:extent"))
            dims = f" {extent.get('cx')}x{extent.get('cy')}" if extent is not None else ""
            self.lines.append(f"{self._prefix}R{summary} [DRAWING{dims}]")
        else:
            self.lines.append(f"{self._prefix}R{summary} [PICT]")
        for txbx in el.findall(".//" + qn("w:txbxContent")):
            self.lines.append(f"{self._prefix}{_IND}TXBX")
            self.lines.extend(walk_block(txbx, self.ctx, self.level + 2))
