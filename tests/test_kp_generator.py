"""Тесты kp_generator: контекст, имя файла, генерация DOCX."""
from __future__ import annotations

import re
import zipfile
from datetime import date
from io import BytesIO

from docx import Document

from src.generators.kp_generator import (
    build_filename,
    build_main_scale_label,
    build_template_context,
    generate_kp,
)


def _state(**overrides) -> dict:
    base = {
        "kp_number": "47141",
        "kp_date": date(2026, 4, 22),
        "kp_valid_days": 15,
        "total_term_days": 35,
        "manager_id": "makarov_av",
        "client_name": "АО «Гипсобетон»",
        "model_line": "С",
        "model_max": 80,
        "model_length": 18,
        "model_id": "vesta-с-80-18",
        "model_price": 2_450_000,
        "warranty_months": 36,
        "is_dual_range": False,
        "sensor_id": "zemic_dhm9b_30t",
        "indicator_id": "titan_3cs",
        "construction_beam": "Двутавр 30Б1",
        "construction_beam_count": 8,
        "construction_center_beam": "Швеллер №12",
        "construction_center_beam_count": 2,
        "construction_deck_mm": 10,
        "construction_underlining_mm": 4,
        "options": {},
        "spec_items_overrides": {},
        "payment_preset_id": "v1_prepay_postpay",
        "payment_split_state": {},
        "payment_percents": {},
        "payment_days": 5,
        "payment_custom_text": "",
        "payment_v1_prepay": 50,
        "payment_v2_prepay": 30,
        "payment_v2_preship": 40,
        "payment_v3_days": 15,
        "payment_v3_trigger_id": "after_installation",
    }
    base.update(overrides)
    return base


# --- build_main_scale_label ---


def test_build_main_scale_label_single_range():
    model = {"verification_division_kg": 20}
    assert build_main_scale_label(model, is_dual_range=False) == "20 кг"


def test_build_main_scale_label_dual_range():
    model = {
        "verification_division_kg": 50,
        "dual_range": {
            "w1": {"max_load_t": 60, "e_kg": 20},
            "w2": {"max_load_t": 80, "e_kg": 50},
        },
    }
    label = build_main_scale_label(model, is_dual_range=True)
    assert "20 кг (до 60 т)" in label
    assert "50 кг (от 60 до 80 т)" in label
    assert "\n" in label
    assert " / " not in label


# --- build_template_context ---


def test_build_template_context_keys(prices):
    state = _state()
    ctx = build_template_context(state, prices)
    expected_keys = {
        "client_name", "kp_number", "kp_date", "kp_valid_days",
        "warranty_text", "max_load_t", "platform_size",
        "main_scale_label", "construction_description",
        "model_full_name",
        "sensor_label", "sensor_temp_range",
        "indicator_label", "indicator_temp_range",
        "spec_items", "total_price", "total_term_days", "vat_percent",
        "payment_terms_block",
        "manager_full_name", "manager_phone", "manager_email",
    }
    assert expected_keys <= set(ctx.keys()), (
        f"Missing keys: {expected_keys - set(ctx.keys())}"
    )
    assert "division_info" not in ctx


def test_build_template_context_values(prices):
    state = _state()
    ctx = build_template_context(state, prices)
    assert ctx["client_name"] == "АО «Гипсобетон»"
    assert ctx["kp_number"] == "47141"
    assert ctx["kp_date"] == "22.04.2026"
    assert ctx["kp_valid_days"] == "15 дней"
    assert ctx["warranty_text"] == "36 месяцев"
    assert ctx["vat_percent"] == "22"
    assert ctx["manager_full_name"] == "Макаров Антон"
    assert isinstance(ctx["spec_items"], list)
    assert all("name" in i and "price" in i and "term_days" in i for i in ctx["spec_items"])


def test_pluralize_kp_valid_days_and_warranty(prices):
    state = _state(kp_valid_days=21, warranty_months=24)
    ctx = build_template_context(state, prices)
    assert ctx["kp_valid_days"] == "21 день"
    assert ctx["warranty_text"] == "24 месяца"


# --- equipment_info (datчик/терминал) ---


def test_build_template_context_has_equipment_keys(prices):
    state = _state()
    ctx = build_template_context(state, prices)
    new_keys = {
        "model_full_name",
        "sensor_label", "sensor_temp_range",
        "indicator_label", "indicator_temp_range",
    }
    assert new_keys.issubset(ctx.keys())
    assert "division_info" not in ctx


def test_sensor_indicator_labels_match_state(prices):
    state = _state(sensor_id="zemic_bm14g_30t", indicator_id="titan_12s")
    ctx = build_template_context(state, prices)
    assert ctx["sensor_label"] == "Zemic BM14G-30t"
    assert ctx["sensor_temp_range"] == "-30...+40"


def test_unknown_sensor_id_falls_back_to_default(prices):
    state = _state(sensor_id="несуществующий")
    ctx = build_template_context(state, prices)
    assert ctx["sensor_label"] == "Zemic DHM9B-30t"


# --- build_filename ---


def test_build_filename_cyrillic_preserved():
    """Кириллица в имени клиента сохраняется без транслитерации."""
    state = _state(client_name="ООО Гипсобетон")
    name = build_filename(state)
    # Формат: НомерКП_НазваниеКомпании_МодельВесов_ДатаКП.docx
    assert name.startswith("47141_")
    assert name.endswith("_2026-04-22.docx")
    assert "Гипсобетон" in name
    # Транслитерации быть не должно
    assert "Gipsobeton" not in name


def test_build_filename_empty_client_uses_fallback():
    state = _state(client_name="")
    name = build_filename(state)
    assert "б-н" in name


def test_build_filename_includes_model():
    state = _state()
    name = build_filename(state)
    assert "80-18" in name  # часть model_full_name


def test_build_filename_strips_org_forms_and_quotes():
    """ООО, АО, кавычки убираются — остаётся только название."""
    for client_name in [
        "ООО «Счастливая Корова»",
        "АО Счастливая Корова",
        "ЗАО «Счастливая Корова»",
    ]:
        state = _state(client_name=client_name)
        name = build_filename(state)
        assert "Счастливая Корова" in name, f"Не найдено название в: {name}"
        assert "ООО" not in name
        assert "АО" not in name
        assert "«" not in name
        assert "»" not in name


def test_build_filename_sanitizes_forbidden_chars():
    """Недопустимые символы ФС заменяются на _, кириллица остаётся."""
    state = _state(kp_number="КП:2026/001", client_name="Рога & Копыта")
    name = build_filename(state)
    assert name.startswith("КП_2026_001_")
    assert "Рога" in name
    assert "/" not in name
    assert ":" not in name


# --- generate_kp ---


def test_generate_kp_returns_zip_bytes(prices):
    """Сгенерированные байты — валидный zip (DOCX)."""
    state = _state()
    docx = generate_kp(state, prices)
    assert isinstance(docx, bytes)
    assert len(docx) > 1000
    assert docx.startswith(b"PK")  # zip signature


def test_generate_kp_no_remaining_jinja(prices):
    """В готовом DOCX нет неподставленных {{ или {% (включая колонтитул)."""
    state = _state()
    docx = generate_kp(state, prices)
    doc = Document(BytesIO(docx))

    parts: list[str] = []
    for p in doc.paragraphs:
        parts.append("".join(r.text or "" for r in p.runs))
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append("".join(r.text or "" for r in p.runs))
    for s in doc.sections:
        f = s.footer
        if f is None:
            continue
        for p in f.paragraphs:
            parts.append("".join(r.text or "" for r in p.runs))
        for t in f.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        parts.append("".join(r.text or "" for r in p.runs))

    text = " ".join(parts)
    assert "{{" not in text, "Остались плейсхолдеры {{...}}"
    assert "{%" not in text, "Остались jinja-теги {%...%}"


def test_generated_docx_preserves_tx_labels(prices):
    """Подписи в строках 10/11 таблицы ТХ не затёрты плейсхолдерами."""
    state = _state(model_id="vesta-с-80-18")
    docx = generate_kp(state, prices)
    doc = Document(BytesIO(docx))
    tx = doc.tables[2]
    assert tx.rows[10].cells[0].text.strip() == "Максимальная нагрузка, т"
    assert tx.rows[11].cells[0].text.strip().startswith(
        "Цена поверочного деления"
    )


def test_generated_docx_description_cell_is_empty(prices):
    """В строке «Описание весов» колонка значения пуста."""
    state = _state()
    docx = generate_kp(state, prices)
    doc = Document(BytesIO(docx))
    tx = doc.tables[2]
    assert "Описание весов" in tx.rows[9].cells[0].text
    assert tx.rows[9].cells[2].text.strip() == ""


def test_section_title_includes_model_name(prices):
    """Заголовок 'Характеристики ВЕСТА-С-80-18' использует model_full_name."""
    state = _state(model_id="vesta-с-80-18")
    docx = generate_kp(state, prices)
    doc = Document(BytesIO(docx))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Характеристики ВЕСТА-С-80-18" in full_text


def test_kp_valid_days_paragraph_is_bold(prices):
    """Срок действия КП — жирным шрифтом."""
    state = _state()
    docx = generate_kp(state, prices)
    doc = Document(BytesIO(docx))
    valid_paras = [
        p for p in doc.paragraphs if "Срок действия настоящего" in p.text
    ]
    assert valid_paras, "Параграф 'Срок действия настоящего' не найден"
    assert any(r.bold for r in valid_paras[0].runs)


def test_spec_item_name_is_richtext_with_8pt_subline(prices):
    """Имя модели в spec_items_fmt — RichText, вспомогательные строки 8pt.

    Ограждение НЕ должно попадать в имя, даже если fence включён.
    """
    state = _state(
        model_id="vesta-с-80-18",
        options={
            "fence_norma_18": {
                "enabled": True, "price": 128_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 128_000, "dealer_is_synthetic": False,
                "block": "fences",
            }
        },
    )
    ctx = build_template_context(state, prices)
    name = ctx["spec_items"][0]["name"]
    # docxtpl.RichText сериализуется в XML с size в half-points
    xml = str(name)
    assert "ВЕСТА-С-80-18" in xml
    assert "Датчики:" in xml
    assert "Терминал" in xml
    assert "Ограждение" not in xml, f"Ограждение не должно быть в имени весов: {xml[:300]}"
    # 8pt = half-points 16
    assert 'w:sz w:val="16"' in xml or 'sz w:val="16"' in xml, (
        f"8pt size не найден в RichText XML: {xml[:300]}"
    )


# --- 1.5b-fix3: шрифты, оранжевый «ВЕСТА», тире в payment-блоке ---


def test_generated_docx_specification_has_no_arial_runs(prices):
    """В run-ах body нет явного Arial — спецификация в PT Sans (1.5b-fix3)."""
    state = _state()
    docx = generate_kp(state, prices)
    with zipfile.ZipFile(BytesIO(docx)) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    assert 'w:ascii="Arial"' not in xml, (
        "Найден явный Arial в run-ах — регрессия rPr"
    )


def test_generated_docx_vesta_is_orange(prices):
    """Слово «ВЕСТА» в заголовке имеет color D04514 (1.5b-fix3)."""
    state = _state()
    docx = generate_kp(state, prices)
    with zipfile.ZipFile(BytesIO(docx)) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    pattern = re.compile(
        r'<w:r[^>]*>(?:(?!</w:r>).)*?w:val="D04514"(?:(?!</w:r>).)*?'
        r'<w:t[^>]*>ВЕСТА</w:t>(?:(?!</w:r>).)*?</w:r>',
        re.DOTALL,
    )
    assert pattern.search(xml), "Оранжевая «ВЕСТА» (D04514) не найдена"


def test_generated_docx_payment_lines_have_dashes(prices):
    """Блок «Условия поставки» содержит ≥ 2 строк с «— » (1.5b-fix3)."""
    state = _state(
        payment_preset_id="split_by_items",
        options={
            "foundation_s_f_18": {
                "enabled": True, "price": 1_900_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 1_900_000, "dealer_is_synthetic": False,
                "block": "foundations",
            },
        },
    )
    docx = generate_kp(state, prices)
    doc = Document(BytesIO(docx))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert text.count("— ") >= 2, (
        f"Меньше 2 строк с «— » в payment-блоке: {text!r}"
    )


def test_generate_kp_payment_block_contains_split_lines(prices):
    """split_by_items: убедимся, что строки про «по уведомлению» и
    «по факту готовности фундамента» есть в тексте."""
    state = _state(
        payment_preset_id="split_by_items",
        options={
            "foundation_s_f_18": {
                "enabled": True, "price": 1_900_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 1_900_000, "dealer_is_synthetic": False,
                "block": "foundations",
            },
        },
    )
    docx = generate_kp(state, prices)
    doc = Document(BytesIO(docx))
    text = " ".join(
        "".join(r.text or "" for r in p.runs) for p in doc.paragraphs
    )
    assert "по уведомлению" in text
    assert "фундамент" in text


def test_payment_block_paragraph_has_no_list_numbering(prices):
    """Параграф с «— Предоплата:» в готовом DOCX не содержит numPr —
    иначе Word добавляет list-маркер перед первой строкой, и получается
    двойное тире (1.5b-fix4)."""
    state = _state()
    docx = generate_kp(state, prices)
    with zipfile.ZipFile(BytesIO(docx)) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    # После рендера «— Предоплата:» — первая строка блока условий оплаты.
    marker = "Предоплата:"
    idx = xml.find(marker)
    assert idx > 0, f"'{marker}' не найден в document.xml"
    para_start = xml.rfind("<w:p ", 0, idx)
    para_end = xml.find("</w:p>", idx) + len("</w:p>")
    para_xml = xml[para_start:para_end]
    assert "<w:numPr>" not in para_xml, (
        "Параграф с «Предоплата:» содержит numPr — Word добавит лишний list-маркер"
    )


# --- per-item срок исполнения с vMerge ---


def test_generated_docx_term_days_with_vmerge(prices):
    """Полный набор позиций при total=30 → vMerge для весов+опций, отдельные
    ячейки для фундамента/монтажа/поверки/доставки.

    Состав: scales + foundation + install + delivery + verification + opts.
    T_default = 20+10+3+1+1 = 35. user=30 < T_default → пропорциональный скейл
    вариативных (variable_target=25, variable_default=30):
      scales = round(20*25/30) = 17, foundation = round(10*25/30) = 8.
    Итого total = 17+8+3+1+1 = 30.

    Ожидается:
      - 1-я content-row (модель): cell[2].text="17" + <w:vMerge w:val="restart"/>
      - 2-я / 3-я (frame, fence): cell[2].text="" + <w:vMerge/> (continue)
      - 4-я (foundation): "8", без vMerge
      - 5-я (install): "3", без vMerge
      - 6-я (delivery): "1", без vMerge
      - 7-я (verification): "1", без vMerge
    """
    state = _state(
        total_term_days=30,
        options={
            "frame_18": {
                "enabled": True, "price": 130_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 130_000, "dealer_is_synthetic": False,
                "block": "frames",
            },
            "fence_norma_18": {
                "enabled": True, "price": 128_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 128_000, "dealer_is_synthetic": False,
                "block": "fences",
            },
            "foundation_s_f_18": {
                "enabled": True, "price": 280_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 280_000, "dealer_is_synthetic": False,
                "block": "foundations",
            },
            "install_default": {
                "enabled": True, "price": 200_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 200_000, "dealer_is_synthetic": False,
                "block": "install",
            },
            "verification_default": {
                "enabled": True, "price": 30_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 30_000, "dealer_is_synthetic": False,
                "block": "verification",
            },
            "delivery_default": {
                "enabled": True, "price": 70_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 70_000, "dealer_is_synthetic": False,
                "block": "delivery",
            },
        },
    )
    docx = generate_kp(state, prices)
    doc = Document(BytesIO(docx))

    spec_table = None
    for t in doc.tables:
        if t.rows[0].cells[0].text.strip() == "Наименование":
            spec_table = t
            break
    assert spec_table is not None, "Таблица спецификации не найдена"

    # Читаем напрямую из XML: row.cells «склеивает» vMerge-ячейки и
    # возвращает дубликаты — это сбивает проверку.
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    qn = lambda tag: f"{{{W}}}{tag}"  # noqa: E731

    trs = spec_table._tbl.findall(qn("tr"))
    assert len(trs) >= 9, f"Ожидалось ≥9 строк, получено {len(trs)}"

    def col2_info(tr):
        tcs = tr.findall(qn("tc"))
        tc = tcs[2]
        text = "".join(t.text or "" for t in tc.iter(qn("t")))
        tcPr = tc.find(qn("tcPr"))
        vm = tcPr.find(qn("vMerge")) if tcPr is not None else None
        if vm is None:
            return text, "none"
        val = vm.get(qn("val"))
        return text, "restart" if val == "restart" else "continue"

    # Порядок content-rows совпадает со spec_items:
    # vesta-с-80-18 (модель), frame_18, fence_norma_18, foundation_s_f_18,
    # install_default, delivery_default, verification_default
    # (OPTION_BLOCKS_ORDER: install → delivery → verification).

    # row 1 = модель: vMerge=restart, value="17"
    text, vm = col2_info(trs[1])
    assert vm == "restart"
    assert text == "17"

    # row 2 = frame, row 3 = fence: vMerge=continue, value=""
    text, vm = col2_info(trs[2])
    assert vm == "continue"
    assert text == ""
    text, vm = col2_info(trs[3])
    assert vm == "continue"
    assert text == ""

    # row 4 = foundation: без vMerge, value="8"
    text, vm = col2_info(trs[4])
    assert vm == "none"
    assert text == "8"

    # row 5 = install: без vMerge, value="3"
    text, vm = col2_info(trs[5])
    assert vm == "none"
    assert text == "3"

    # row 6 = delivery: без vMerge, value="1"
    text, vm = col2_info(trs[6])
    assert vm == "none"
    assert text == "1"

    # row 7 = verification: без vMerge, value="1"
    text, vm = col2_info(trs[7])
    assert vm == "none"
    assert text == "1"

    # Маркеров в спек-таблице нет
    full_text = "".join(
        t.text or "" for t in spec_table._tbl.iter(qn("t"))
    )
    assert "⟦MERGE" not in full_text


def test_generated_docx_no_foundation_term_days(prices):
    """Без фундамента, total=30. Состав: scales+install+delivery+verification.

    T_default = 20+3+1+1 = 25. user=30 > T_default → скейл вариативных вверх.
    variable_target = 25, variable_default = 20, scales = round(20*25/20) = 25.
    """
    state = _state(
        total_term_days=30,
        options={
            "frame_18": {
                "enabled": True, "price": 130_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 130_000, "dealer_is_synthetic": False,
                "block": "frames",
            },
            "install_default": {
                "enabled": True, "price": 200_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 200_000, "dealer_is_synthetic": False,
                "block": "install",
            },
            "verification_default": {
                "enabled": True, "price": 30_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 30_000, "dealer_is_synthetic": False,
                "block": "verification",
            },
            "delivery_default": {
                "enabled": True, "price": 70_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 70_000, "dealer_is_synthetic": False,
                "block": "delivery",
            },
        },
    )
    docx = generate_kp(state, prices)
    doc = Document(BytesIO(docx))

    spec_table = None
    for t in doc.tables:
        if t.rows[0].cells[0].text.strip() == "Наименование":
            spec_table = t
            break
    assert spec_table is not None

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    qn = lambda tag: f"{{{W}}}{tag}"  # noqa: E731

    # Модель в первой content-row → "25" (variable_target=25 целиком на scales)
    trs = spec_table._tbl.findall(qn("tr"))
    tcs = trs[1].findall(qn("tc"))
    text = "".join(t.text or "" for t in tcs[2].iter(qn("t")))
    assert text == "25"


def test_generate_kp_raises_term_too_small(prices):
    """total_term_days=3 при полном наборе → TermDaysTooSmallError из generate_kp."""
    import pytest as _pytest

    from src.term_days import TermDaysTooSmallError

    state = _state(
        total_term_days=3,
        options={
            "install_default": {
                "enabled": True, "price": 200_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 200_000, "dealer_is_synthetic": False,
                "block": "install",
            },
            "verification_default": {
                "enabled": True, "price": 30_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 30_000, "dealer_is_synthetic": False,
                "block": "verification",
            },
            "delivery_default": {
                "enabled": True, "price": 70_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 70_000, "dealer_is_synthetic": False,
                "block": "delivery",
            },
        },
    )
    with _pytest.raises(TermDaysTooSmallError) as exc:
        generate_kp(state, prices)
    d = exc.value.details
    assert d["total"] == 3
    # fixed_total = 3+1+1 = 5; active_variable = {scales}; min = 5+1 = 6
    assert d["min"] == 6
    assert d["fixed_total"] == 5
    assert d["variable_count"] == 1
    assert d["install"] == 3


def test_generated_docx_full_set_with_canopy_and_orion(prices):
    """Расширенный состав: scales+frame+orion+canopy+foundation+install+delivery+verification.

    T_default = 20+5+25+10+3+1+1 = 65 (без user override → дефолты по ролям).
    Каждая роль (orion/canopy/foundation/install/delivery/verification) — своя
    строка со своим числом, без vMerge. ОРИОН НЕ сливается с весами.
    """
    state = _state(
        total_term_days=None,
        options={
            "frame_18": {
                "enabled": True, "price": 130_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 130_000, "dealer_is_synthetic": False,
                "block": "frames",
            },
            "orion_standard": {
                "enabled": True, "price": 380_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 380_000, "dealer_is_synthetic": False,
                "block": "pak_orion",
            },
            "canopy_turnkey_18": {
                "enabled": True, "price": 1_500_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 1_500_000, "dealer_is_synthetic": False,
                "block": "canopy",
            },
            "foundation_s_f_18": {
                "enabled": True, "price": 280_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 280_000, "dealer_is_synthetic": False,
                "block": "foundations",
            },
            "install_default": {
                "enabled": True, "price": 200_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 200_000, "dealer_is_synthetic": False,
                "block": "install",
            },
            "delivery_default": {
                "enabled": True, "price": 70_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 70_000, "dealer_is_synthetic": False,
                "block": "delivery",
            },
            "verification_default": {
                "enabled": True, "price": 30_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 30_000, "dealer_is_synthetic": False,
                "block": "verification",
            },
        },
    )
    docx = generate_kp(state, prices)
    doc = Document(BytesIO(docx))

    spec_table = None
    for t in doc.tables:
        if t.rows[0].cells[0].text.strip() == "Наименование":
            spec_table = t
            break
    assert spec_table is not None

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    qn = lambda tag: f"{{{W}}}{tag}"  # noqa: E731

    def col2(tr):
        tc = tr.findall(qn("tc"))[2]
        text = "".join(t.text or "" for t in tc.iter(qn("t")))
        tcPr = tc.find(qn("tcPr"))
        vm = tcPr.find(qn("vMerge")) if tcPr is not None else None
        if vm is None:
            return text, "none"
        return text, "restart" if vm.get(qn("val")) == "restart" else "continue"

    trs = spec_table._tbl.findall(qn("tr"))

    # row 1 = модель: restart "20"
    text, vm = col2(trs[1])
    assert vm == "restart" and text == "20"
    # row 2 = frame (роль None) → continue
    text, vm = col2(trs[2])
    assert vm == "continue" and text == ""
    # row 3 = orion → отдельная строка "5"
    text, vm = col2(trs[3])
    assert vm == "none" and text == "5"
    # row 4 = canopy → "25"
    text, vm = col2(trs[4])
    assert vm == "none" and text == "25"
    # row 5 = foundation → "10"
    text, vm = col2(trs[5])
    assert vm == "none" and text == "10"
    # row 6 = install → "3"
    text, vm = col2(trs[6])
    assert vm == "none" and text == "3"
    # row 7 = delivery → "1"
    text, vm = col2(trs[7])
    assert vm == "none" and text == "1"
    # row 8 = verification → "1"
    text, vm = col2(trs[8])
    assert vm == "none" and text == "1"


def test_generated_docx_foundation_label_uses_model_line(prices):
    """Лейбл фундамента в DOCX подменён на букву линейки выбранной модели."""
    state = _state(
        model_id="vesta-с-80-18",
        model_line="С",
        total_term_days=None,
        options={
            "foundation_s_f_18": {
                "enabled": True, "price": 280_000, "qty": 1,
                "customer_side": False, "is_on_request": False,
                "retail": 280_000, "dealer_is_synthetic": False,
                "block": "foundations",
            },
        },
    )
    docx = generate_kp(state, prices)
    doc = Document(BytesIO(docx))

    spec_table = None
    for t in doc.tables:
        if t.rows[0].cells[0].text.strip() == "Наименование":
            spec_table = t
            break
    assert spec_table is not None

    # Найти строку фундамента: ищем в первой колонке "Фундамент"
    foundation_text = ""
    for row in spec_table.rows[1:]:
        cell_text = row.cells[0].text
        if "Фундамент" in cell_text:
            foundation_text = cell_text
            break
    assert foundation_text, "Строка фундамента не найдена в spec-таблице"
    assert "ВЕСТА-С" in foundation_text
    assert "С/Ф/П" not in foundation_text


def test_validity_paragraph_uses_pt_sans(prices):
    """Параграф «Срок действия КП» использует PT Sans (не Times New Roman) и жирный."""
    state = _state()
    docx = generate_kp(state, prices)
    doc = Document(BytesIO(docx))

    target = None
    for para in doc.paragraphs:
        if "Срок действия настоящего" in para.text:
            target = para
            break

    assert target is not None, "Параграф со сроком действия не найден"

    fonts_found: set[str] = set()
    for run in target.runs:
        rpr = run._r.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr")
        if rpr is not None:
            rfonts = rpr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts")
            if rfonts is not None:
                W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                for attr in ("ascii", "hAnsi", "cs"):
                    val = rfonts.get(f"{{{W}}}{attr}")
                    if val:
                        fonts_found.add(val)

    assert any("PT Sans" in f for f in fonts_found), (
        f"Ожидался PT Sans в rFonts, найдено: {fonts_found}"
    )
    assert any(run.bold for run in target.runs), (
        "Срок действия КП должен быть жирным"
    )


# --- 1.2: PT Sans 11pt в ячейке температурного диапазона ---


def test_temperature_cell_uses_pt_sans_11pt(prices):
    """Tbl 2 (ТХ), row 1 — ячейка с температурными диапазонами должна
    рендериться в PT Sans 11pt. До фикса (1.2) первый run эталона не имел
    явных rFonts/sz → Word наследовал Calibri 11pt по дефолту."""
    state = _state()
    docx = generate_kp(state, prices)
    doc = Document(BytesIO(docx))

    tx_table = doc.tables[2]
    cell_2 = tx_table.rows[1].cells[2]  # значения {{ sensor_temp_range }}\n{{ indicator_temp_range }}
    cell_0 = tx_table.rows[1].cells[0]  # подписи строки

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    qn = lambda tag: f"{{{W}}}{tag}"  # noqa: E731

    for cell, label in [(cell_2, "значения температур"), (cell_0, "подписи температур")]:
        runs = cell._tc.findall(f".//{qn('r')}")
        assert runs, f"{label}: ни одного <w:r> в ячейке"
        for r in runs:
            rpr = r.find(qn("rPr"))
            assert rpr is not None, f"{label}: <w:r> без rPr"
            rfonts = rpr.find(qn("rFonts"))
            assert rfonts is not None, f"{label}: rPr без rFonts"
            assert rfonts.get(qn("ascii")) == "PT Sans", (
                f"{label}: rFonts ascii={rfonts.get(qn('ascii'))!r} ≠ 'PT Sans'"
            )
            sz = rpr.find(qn("sz"))
            assert sz is not None, f"{label}: rPr без sz"
            assert sz.get(qn("val")) == "22", (
                f"{label}: sz val={sz.get(qn('val'))!r} ≠ '22' (11pt)"
            )


# --- 2.7: V1 в DOCX — текст «— Предоплата 30%» / «— Доплата 70%» ---


def test_v1_block_renders_in_docx(prices):
    """V1 (Аванс+Постоплата) с prepay=30 → в DOCX блок условий поставки
    содержит «Предоплата: 30%» и «Доплата: 70%»."""
    state = _state(
        payment_preset_id="v1_prepay_postpay",
        payment_v1_prepay=30,
        payment_days=5,
    )
    docx = generate_kp(state, prices)
    doc = Document(BytesIO(docx))
    text = " ".join(
        "".join(r.text or "" for r in p.runs) for p in doc.paragraphs
    )
    assert "Предоплата: 30%" in text
    assert "Доплата: 70%" in text
    assert "5 банковских дней" in text


def test_v3_block_renders_in_docx(prices):
    """V3 (100% постоплата) — 30 дней после акта приёмки."""
    state = _state(
        payment_preset_id="v3_postpay_only",
        payment_v3_days=30,
        payment_v3_trigger_id="after_act",
    )
    docx = generate_kp(state, prices)
    doc = Document(BytesIO(docx))
    text = " ".join(
        "".join(r.text or "" for r in p.runs) for p in doc.paragraphs
    )
    assert "100% оплаты" in text
    assert "30 банковских дней" in text
    assert "после подписания акта приёмки" in text
