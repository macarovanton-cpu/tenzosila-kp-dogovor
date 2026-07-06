"""Тесты фичи «Количество автовесов > 1» (model_qty).

Строки спецификации остаются per-unit; итоги и оплата — от подытог_за_1 × qty.
Каждый кейс проверяется на qty=2 И на регресс qty=1.
"""
from __future__ import annotations

import os
import tempfile
from datetime import date
from io import BytesIO

from docx import Document

from src.contracts.filler import fill_spec_with_items
from src.contracts.payment_line import build_lines_from_snapshot
from src.contracts.supply_filler import build_supply_context
from src.generators.kp_generator import build_template_context, generate_kp
from src.state import initial_state
from src.storage.kp_restore import reconstruct_kp_state
from src.storage.snapshot_builder import build_kp_snapshot


def _spec_table_rows(doc) -> list[list[str]]:
    """Строки таблицы, содержащей ИТОГО (текст ячеек)."""
    for t in doc.tables:
        joined = " ".join(c.text for r in t.rows for c in r.cells)
        if "ИТОГО" in joined:
            return [[c.text.strip() for c in r.cells] for r in t.rows]
    return []


# --- КП: контекст subtotal_N, строки per-unit ---


def test_kp_context_subtotal_scales_with_qty(prices):
    ctx1 = build_template_context(initial_state() | {"model_qty": 1}, prices)
    ctx2 = build_template_context(initial_state() | {"model_qty": 2}, prices)

    # Цены строк спецификации (per-unit) идентичны — qty их не трогает.
    assert [i["price"] for i in ctx1["spec_items"]] == [
        i["price"] for i in ctx2["spec_items"]
    ]
    assert ctx1["total_price"] == ctx1["subtotal_per_1"]
    assert ctx2["subtotal_per_1"] == ctx1["subtotal_per_1"]

    per1 = int(ctx2["subtotal_per_1"].replace(" ", ""))
    n = int(ctx2["subtotal_n"].replace(" ", ""))
    assert n == per1 * 2
    assert ctx2["qty_scales"] == 2


def test_kp_docx_qty1_single_total_row(prices):
    doc = Document(BytesIO(generate_kp(initial_state() | {"model_qty": 1}, prices)))
    rows = _spec_table_rows(doc)
    labels = [r[0] for r in rows]
    assert "ИТОГО" in labels
    assert not any("за 1 весы" in lbl for lbl in labels)


def test_kp_docx_qty2_two_total_rows(prices):
    doc = Document(BytesIO(generate_kp(initial_state() | {"model_qty": 2}, prices)))
    rows = _spec_table_rows(doc)
    labels = [r[0] for r in rows]
    assert "ИТОГО за 1 весы" in labels
    assert "ИТОГО за 2 весов" in labels
    # Сумма «за N» = «за 1» × 2.
    per1 = next(int(r[1].replace(" ", "")) for r in rows if r[0] == "ИТОГО за 1 весы")
    n = next(int(r[1].replace(" ", "")) for r in rows if r[0] == "ИТОГО за 2 весов")
    assert n == per1 * 2


# --- Оплата договора: график от subtotal_N ---


def _split_payment() -> dict:
    return {
        "preset_id": "split_by_items",
        "days": 5,
        "split_state": {"scales": {"prepay": 40, "postpay": 60}},
    }


def test_payment_lines_scale_with_qty():
    spec = [{"name": "Весы", "total": 3_050_000, "payment_group": "scales", "item_key": "model"}]

    lines1 = build_lines_from_snapshot(_split_payment(), spec, 1)
    lines2 = build_lines_from_snapshot(_split_payment(), spec, 2)

    # Регресс qty=1: 40% от 3 050 000.
    assert lines1[0].amount == 1_220_000
    assert lines1[0].base_amount == 3_050_000
    # qty=2: 40% от subtotal_N = 6 100 000 (эталон пользователя).
    assert lines2[0].amount == 2_440_000
    assert lines2[0].base_amount == 6_100_000


# --- Supply-контекст: цена договора и «в количестве N шт» ---


def _supply_items() -> list[dict]:
    return [
        {"name": "Весы автомобильные ВЕСТА-С-60-18-Ц", "total": 2_000_000,
         "price": 2_000_000, "quantity": 1, "payment_group": "scales", "item_key": "model"},
        {"name": "Доставка", "total": 200_000, "price": 200_000,
         "quantity": 1, "payment_group": "delivery", "item_key": "delivery"},
    ]


def _supply_ctx(qty: int) -> dict:
    return build_supply_context(
        {"ЗАКАЗЧИК_КРАТКОЕ_НАИМЕНОВАНИЕ": "ООО Тест"},
        _supply_items(), {"items": _supply_items()}, [],
        {"contract_number": "1", "object_address": "г. Тест"},
        date(2026, 7, 6), model_qty=qty,
    )


def test_supply_contract_total_scales_with_qty():
    c1, c2 = _supply_ctx(1), _supply_ctx(2)
    assert c1["СУММА_ЦИФРАМИ"] == "2 200 000"
    assert c2["СУММА_ЦИФРАМИ"] == "4 400 000"
    assert "1шт." in c1["ТОВАР_НАИМЕНОВАНИЕ"]
    assert "2шт." in c2["ТОВАР_НАИМЕНОВАНИЕ"]
    # Служебные ключи для второй строки ИТОГО appendix_1.
    assert c2["_qty_scales"] == 2
    assert c2["_itogo_per_1"] == "2 200 000"
    assert c2["_itogo_n"] == "4 400 000"
    assert c1["_qty_scales"] == 1


# --- Спецификация: вторая строка ИТОГО ---


def test_fill_spec_qty1_single_row():
    items = [{"name": "Весы", "total": 2_000_000, "metadata": {}},
             {"name": "Доставка", "total": 200_000, "metadata": {}}]
    out = os.path.join(tempfile.gettempdir(), "_t_spec_q1.docx")
    try:
        fill_spec_with_items("templates/contracts/spec_v2.docx", {}, items, out, model_qty=1)
        rows = _spec_table_rows(Document(out))
        labels = [r[0] for r in rows]
        assert not any("за 1 весы" in lbl for lbl in labels)
        assert any("2 200 000" in r[1] for r in rows)
    finally:
        os.path.exists(out) and os.unlink(out)


def test_fill_spec_qty2_two_rows():
    items = [{"name": "Весы", "total": 2_000_000, "metadata": {}},
             {"name": "Доставка", "total": 200_000, "metadata": {}}]
    out = os.path.join(tempfile.gettempdir(), "_t_spec_q2.docx")
    try:
        fill_spec_with_items("templates/contracts/spec_v2.docx", {}, items, out, model_qty=2)
        rows = _spec_table_rows(Document(out))
        labels = [r[0] for r in rows]
        assert "ИТОГО за 1 весы" in labels
        assert "ИТОГО за 2 весов" in labels
        # Строки позиций не тронуты (per-unit).
        assert any(r[0] == "Весы" and "2 000 000" in r[1] for r in rows)
    finally:
        os.path.exists(out) and os.unlink(out)


# --- Снапшот round-trip ---


def test_snapshot_roundtrip_model_qty(prices):
    state = initial_state() | {"model_qty": 3, "model_price": 2_000_000}
    snap = build_kp_snapshot(state)
    assert snap["model"]["qty"] == 3

    out = reconstruct_kp_state({"data": snap}, prices)
    assert out["model_qty"] == 3


def test_snapshot_legacy_without_qty_defaults_to_1(prices):
    state = initial_state() | {"model_price": 2_000_000}
    snap = build_kp_snapshot(state)
    del snap["model"]["qty"]  # эмуляция старого снапшота

    out = reconstruct_kp_state({"data": snap}, prices)
    assert out["model_qty"] == 1
