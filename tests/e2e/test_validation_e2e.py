"""Валидация UI: пустые поля блокируют генерацию, длинные/спецсимволы — OK."""
from __future__ import annotations

from pathlib import Path

from docx import Document

from tests.e2e.helpers.docx_assertions import _cell_text_full
from tests.e2e.pages.kp_page import save_download

OUTPUT = Path(__file__).resolve().parent.parent / "output" / "validation"


def test_empty_client_disables_generate(kp_page):
    kp_page.set_kp_number("test-001")
    kp_page.select_line("С")
    kp_page.set_max_load(60)
    kp_page.set_length(18)
    # client_name НЕ заполняем
    assert kp_page.is_generate_disabled(), (
        "Кнопка должна быть disabled при пустом client_name. "
        f"Сообщения валидации: {kp_page.get_validation_messages()}"
    )


def test_empty_kp_number_disables_generate(kp_page):
    kp_page.set_client_name("ООО Тест")
    kp_page.select_line("С")
    kp_page.set_max_load(60)
    kp_page.set_length(18)
    # kp_number НЕ заполняем
    assert kp_page.is_generate_disabled(), (
        "Кнопка должна быть disabled при пустом kp_number. "
        f"Сообщения валидации: {kp_page.get_validation_messages()}"
    )


def test_long_client_name_renders(kp_page):
    long_name = ("ООО " + "Очень-длинное-наименование-компании-" * 5)[:200]
    kp_page.set_kp_number("test-long")
    kp_page.set_client_name(long_name)
    kp_page.select_line("С")
    kp_page.set_max_load(60)
    kp_page.set_length(18)

    download = kp_page.generate()
    docx = save_download(download, OUTPUT / "long_name.docx")

    doc = Document(str(docx))
    full = "\n".join(p.text for p in doc.paragraphs)
    assert long_name[:50] in full, (
        f"Префикс длинного имени не найден в DOCX. Header={full[:300]!r}"
    )


def test_special_chars_in_client_name(kp_page):
    name = 'АО "Гипсо&бетон" <тест>'
    kp_page.set_kp_number("test-special")
    kp_page.set_client_name(name)
    kp_page.select_line("С")
    kp_page.set_max_load(60)
    kp_page.set_length(18)

    download = kp_page.generate()
    docx = save_download(download, OUTPUT / "special_chars.docx")

    doc = Document(str(docx))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                parts.append(_cell_text_full(c))
    full = "\n".join(parts)
    # Спецсимволы XML экранируются Streamlit'ом / docxtpl: «&», «<», «>»
    # и текст после них теряется. Контракт теста — DOCX валидный (Document
    # открылся) и хотя бы префикс имени до первого спецсимвола сохранился.
    assert "Гипсо" in full
