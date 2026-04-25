"""Парсинг и ассерты по сгенерированному DOCX.

Источник истины — реальный sample КП-Гипсобетон:
- Шапка (клиент, номер, дата) лежит в `doc.paragraphs[0..1]`, а не в tables[0].
- ТХ — `doc.tables[2]` (17×3); ключевые строки [10] (max_load), [11] (main_scale),
  [12] (platform), [13] (конструкция).
- Спецификация — `doc.tables[3]` (3 колонки: Наименование, Цена, Срок),
  row[0] = заголовок, row[1] = модель (имя через RichText — читать через XML),
  row[N+1] = «ИТОГО».
- Блок оплаты — параграф, начинающийся с «— Предоплата:» (Listing с `\n` → `<w:br/>`).
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import _Cell


_MONEY_RE = re.compile(r"[^\d]")


def _parse_money(s: str) -> int | None:
    digits = _MONEY_RE.sub("", s or "")
    return int(digits) if digits else None


def _cell_text_full(cell: _Cell) -> str:
    """Все текст-узлы ячейки (включая RichText-runs) через XML.

    `cell.text` обрывается на пустых параграфах между runs RichText.
    Идём через `<w:t>` и склеиваем переносами.
    """
    parts = [el.text for el in cell._tc.iter(qn("w:t")) if el.text]
    return "\n".join(parts)


def assert_header(
    docx_path: Path, *, client_name: str, kp_number: str, kp_date: str | None = None
) -> None:
    doc = Document(str(docx_path))
    paras = [p.text for p in doc.paragraphs]
    head1 = paras[0] if paras else ""
    head2 = paras[1] if len(paras) > 1 else ""
    assert client_name in head1, (
        f"В первом параграфе нет клиента: ожидали {client_name!r}, "
        f"в документе {head1!r}"
    )
    assert kp_number in head2, (
        f"Во втором параграфе нет номера КП {kp_number!r}: {head2!r}"
    )
    if kp_date:
        assert kp_date in head2, (
            f"Дата {kp_date!r} не найдена в шапке: {head2!r}"
        )


def assert_tx(
    docx_path: Path,
    *,
    max_load_t: int,
    platform_length_m: int | None = None,
    main_scale_contains: list[str] | None = None,
) -> None:
    doc = Document(str(docx_path))
    tx = doc.tables[2]
    rows = [[_cell_text_full(c) for c in r.cells] for r in tx.rows]

    # row[10] — Максимальная нагрузка
    label_max = rows[10][0].lower()
    assert "максимальная нагрузка" in label_max, (
        f"Ожидали row[10][0]='Максимальная нагрузка...', получили {label_max!r}"
    )
    assert str(max_load_t) in rows[10][2], (
        f"max_load {max_load_t} не найден в row[10][2]={rows[10][2]!r}"
    )

    # row[11] — Цена поверочного деления
    if main_scale_contains:
        scale = rows[11][2]
        for token in main_scale_contains:
            assert token in scale, (
                f"main_scale_label не содержит {token!r}; "
                f"row[11][2]={scale!r}"
            )

    # row[12] — Платформа
    if platform_length_m is not None:
        platform = rows[12][2]
        expected = f"{platform_length_m}×3"
        assert expected in platform, (
            f"platform_size: ждали '{expected}', получили {platform!r}"
        )


def assert_spec(
    docx_path: Path,
    *,
    expected_item_names: list[str],
    expected_total: int | None = None,
    require_model_full_name: str | None = None,
) -> None:
    """Проверить спецификацию.

    expected_item_names — подстроки имён позиций, которые должны быть в
    каждой строке (по порядку). Базовая модель — первая строка.
    """
    doc = Document(str(docx_path))
    spec = doc.tables[3]
    rows = [[_cell_text_full(c) for c in r.cells] for r in spec.rows]
    assert rows[0][0].strip().lower() == "наименование", (
        f"Заголовок спецификации не на месте: {rows[0]!r}"
    )

    body = rows[1:-1]  # без header и без ИТОГО
    assert len(body) == len(expected_item_names), (
        f"Количество позиций не совпало: ожидали {len(expected_item_names)}, "
        f"в DOCX {len(body)}. Имена в DOCX: "
        f"{[r[0][:60] for r in body]}"
    )
    for actual_row, expected_substr in zip(body, expected_item_names):
        actual_name = actual_row[0]
        assert expected_substr in actual_name, (
            f"Не нашли {expected_substr!r} в имени позиции: {actual_name!r}"
        )

    if require_model_full_name is not None:
        assert require_model_full_name in body[0][0], (
            f"В первой строке спеки нет full-name модели "
            f"{require_model_full_name!r}: {body[0][0]!r}"
        )

    total_row = rows[-1]
    # «ИТОГО» в DOCX иногда рендерится с переносом внутри слова — нормализуем.
    total_label = re.sub(r"\s+", "", total_row[0]).lower()
    assert "итого" in total_label, (
        f"Последняя строка спеки не ИТОГО: {total_row!r}"
    )
    if expected_total is not None:
        actual_total = _parse_money(total_row[1])
        assert actual_total == expected_total, (
            f"ИТОГО: ждали {expected_total}, в DOCX {actual_total} "
            f"(текст {total_row[1]!r})"
        )


def find_payment_paragraph(doc) -> str:
    for p in doc.paragraphs:
        text = p.text
        low = text.lower()
        # Корневой match (родительный/именительный): «предоплат», «оплат».
        starts_with_dash = text.lstrip().startswith("—")
        if "предоплат" in low or (starts_with_dash and "оплат" in low):
            return text
    return ""


def assert_payment(
    docx_path: Path,
    *,
    contains: list[str] | None = None,
    min_dash_lines: int | None = None,
    exact_dash_lines: int | None = None,
) -> None:
    doc = Document(str(docx_path))
    block = find_payment_paragraph(doc)
    assert block, "Параграф с условиями оплаты не найден"
    if contains:
        for token in contains:
            assert token in block, (
                f"В блоке оплаты нет {token!r}; полный текст: {block!r}"
            )
    if min_dash_lines is not None or exact_dash_lines is not None:
        # Считаем только строки, начинающиеся с «— » (после переносов).
        lines = [
            ln for ln in block.split("\n")
            if ln.lstrip().startswith("—")
        ]
        if exact_dash_lines is not None:
            assert len(lines) == exact_dash_lines, (
                f"Ожидали ровно {exact_dash_lines} строк с «— », "
                f"получили {len(lines)}: {lines!r}"
            )
        if min_dash_lines is not None:
            assert len(lines) >= min_dash_lines, (
                f"Ожидали минимум {min_dash_lines} строк с «— », "
                f"получили {len(lines)}: {lines!r}"
            )


def assert_full_text_contains(docx_path: Path, *substrings: str) -> None:
    doc = Document(str(docx_path))
    parts: list[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                parts.append(_cell_text_full(c))
    text = "\n".join(parts)
    for s in substrings:
        assert s in text, (
            f"Подстрока {s!r} не найдена в DOCX. Длина текста {len(text)}."
        )


__all__ = [
    "assert_header",
    "assert_tx",
    "assert_spec",
    "assert_payment",
    "assert_full_text_contains",
    "find_payment_paragraph",
]
