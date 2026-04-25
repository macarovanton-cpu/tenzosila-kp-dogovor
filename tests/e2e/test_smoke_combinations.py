"""Smoke по 5 линейкам: для каждой — минимальная конфигурация генерит DOCX."""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from tests.e2e.helpers.docx_assertions import _cell_text_full
from tests.e2e.pages.kp_page import save_download

OUTPUT = Path(__file__).resolve().parent.parent / "output" / "smoke"


@pytest.mark.parametrize(
    "line, max_load, length, model_id",
    [
        ("С", 60, 18, "vesta-с-60-18"),
        ("СЛ", 60, 18, "vesta-сл-60-18"),
        ("Ф", 60, 18, "vesta-ф-60-18"),
        ("ФЛ", 60, 18, "vesta-фл-60-18"),
        pytest.param(
            "П", 80, 18, "vesta-п-80-18",
            marks=pytest.mark.xfail(
                strict=False,
                reason="vesta-п-80-18 помечена data_incomplete в models.json",
            ),
        ),
    ],
)
def test_smoke_each_line_generates_docx(
    kp_page, line, max_load, length, model_id
):
    OUTPUT.mkdir(parents=True, exist_ok=True)
    expected_full = f"ВЕСТА-{line}-{max_load}-{length}"
    client = f"Smoke {line}"
    kp_no = f"smoke-{line}-{max_load}-{length}"

    kp_page.set_kp_number(kp_no)
    kp_page.set_client_name(client)
    kp_page.select_line(line)
    kp_page.set_max_load(max_load)
    kp_page.set_length(length)

    download = kp_page.generate()
    docx = save_download(download, OUTPUT / f"smoke_{model_id}.docx")

    doc = Document(str(docx))
    parts: list[str] = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                parts.append(_cell_text_full(c))
    full = "\n".join(parts)

    assert expected_full in full, (
        f"Не нашли {expected_full!r} в сгенерированном DOCX"
    )
    assert client in full, (
        f"Не нашли клиента {client!r} в DOCX"
    )
