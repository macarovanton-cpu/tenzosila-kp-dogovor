"""State-переходы: смена линейки, повторное переключение опций.

Каскад в state.py:on_cascade_change сбрасывает options + dual_range
при смене model_id. Эти тесты проверяют, что после переходов state
не повреждён и DOCX по-прежнему генерируется корректно.
"""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from tests.e2e.helpers.docx_assertions import _cell_text_full
from tests.e2e.pages.kp_page import save_download

OUTPUT = Path(__file__).resolve().parent.parent / "output" / "state"


def test_options_persist_after_compatible_line_change(kp_page):
    """Сейчас on_cascade_change при смене линейки сбрасывает все options.
    Проверяем: после С → Ф пере-включаем доставку → DOCX генерится с доставкой.
    """
    kp_page.set_kp_number("trans-line")
    kp_page.set_client_name("Тест переключения линейки")
    kp_page.select_line("С")
    kp_page.set_max_load(80)
    kp_page.set_length(18)
    kp_page.toggle_option("Доставка", "Доставка весов до объекта", on=True)

    # Смена линейки. Колбэк on_cascade_change сбрасывает options.
    kp_page.select_line("Ф")
    # Перевключаем доставку под новую линейку — UI rerender'ится с новым model_id.
    kp_page.toggle_option("Доставка", "Доставка весов до объекта", on=True)

    download = kp_page.generate()
    docx = save_download(download, OUTPUT / "transition_line.docx")

    doc = Document(str(docx))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                parts.append(_cell_text_full(c))
    full = "\n".join(parts)
    assert "ВЕСТА-Ф-80-18" in full, (
        "Не подхватилась новая линейка Ф после переключения"
    )
    spec_text = "\n".join(_cell_text_full(c) for r in doc.tables[3].rows for c in r.cells)
    assert "Доставка весов до объекта" in spec_text


def test_dual_range_resets_when_unsupported(kp_page):
    """Tests for models without dual_range support — отложен до расширения scope."""
    pytest.skip(
        "Все scope-модели MVP имеют dual_range. Тест активируется при добавлении М1/М2/К."
    )


def test_repeated_option_toggle_no_state_corruption(kp_page):
    """ON/OFF/ON по «Доставка» → доставка корректно в спеке ровно одна."""
    kp_page.set_kp_number("toggle-001")
    kp_page.set_client_name("Тест-toggle")
    kp_page.select_line("С")
    kp_page.set_max_load(60)
    kp_page.set_length(18)

    kp_page.toggle_option("Доставка", "Доставка весов до объекта", on=True)
    kp_page.toggle_option("Доставка", "Доставка весов до объекта", on=False)
    kp_page.toggle_option("Доставка", "Доставка весов до объекта", on=True)

    download = kp_page.generate()
    docx = save_download(download, OUTPUT / "toggle.docx")

    doc = Document(str(docx))
    spec = doc.tables[3]
    spec_rows = [_cell_text_full(c) for r in spec.rows for c in r.cells]
    text = "\n".join(spec_rows)
    occurrences = text.count("Доставка весов до объекта")
    assert occurrences == 1, (
        f"Доставка должна быть ровно одна в спеке, нашли {occurrences}: {text!r}"
    )
