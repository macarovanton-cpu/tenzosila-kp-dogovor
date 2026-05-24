"""Integration-тесты: build_spec_v2_data → fill_spec_v2 → проверка DOCX."""
import json
import os

import pytest
from docx import Document

from src.contracts.filler import get_unfilled_placeholders
from src.contracts.from_kp import build_spec_v2_data
from src.contracts.spec_v2_filler import fill_spec_v2

SPEC_V2_PATH = os.path.normpath(os.path.join(
    os.path.dirname(__file__), '..', '..', 'templates', 'contracts', 'spec_v2.docx'
))
DATA_DIR = os.path.normpath(os.path.join(
    os.path.dirname(__file__), '..', '..', 'data'
))

_REQUISITES = {
    "ДОГОВОР_НОМЕР": "Т-INT/2026",
    "ДОГОВОР_ДАТА_ПОЛНАЯ": "20.05.2026",
    "ДОГОВОР_ДЕНЬ": "20",
    "ДОГОВОР_МЕСЯЦ": "мая",
    "ДОГОВОР_ГОД": "2026",
    "СПЕЦ_НОМЕР": "1",
    "СПЕЦ_АДРЕС_ОБЪЕКТА": "г. Тест, ул. Интеграции, 1",
    "ЗАКАЗЧИК_ДИРЕКТОР_ДОЛЖНОСТЬ": "Генеральный директор",
    "ЗАКАЗЧИК_ПОЛНОЕ_НАИМЕНОВАНИЕ": "ООО «Интеграция»",
    "ЗАКАЗЧИК_КРАТКОЕ_НАИМЕНОВАНИЕ": "ООО «Интеграция»",
    "ЗАКАЗЧИК_ДИРЕКТОР_ФИО_КРАТКОЕ": "Тестов Т.Т.",
    "ЗАКАЗЧИК_ДИРЕКТОР_ИНИЦИАЛЫ": "Т.Т. Тестов",
    "ЗАКАЗЧИК_ДИРЕКТОР_ДОЛЖНОСТЬ_РП": "генерального директора",
    "ЗАКАЗЧИК_ДИРЕКТОР_ФИО_РП": "Тестова Тимофея Тимофеевича",
    "ЗАКАЗЧИК_ОСНОВАНИЕ": "Устава",
    "ДИРЕКТОР_ПРИЧАСТИЕ": "действующего",
}


@pytest.fixture(scope="module")
def json_data():
    with open(os.path.join(DATA_DIR, "models.json"), encoding="utf-8") as f:
        models = json.load(f)
    with open(os.path.join(DATA_DIR, "prices.json"), encoding="utf-8") as f:
        prices = json.load(f)
    with open(os.path.join(DATA_DIR, "payment_terms.json"), encoding="utf-8") as f:
        payment_terms = json.load(f)
    with open(os.path.join(DATA_DIR, "equipment_specs.json"), encoding="utf-8") as f:
        equipment_specs = json.load(f)
    return models, prices, payment_terms, equipment_specs


def _make_kp(
    model_id: str, line: str, max_t: int, length: int, price: int,
    options: dict | None = None, flags: dict | None = None,
) -> dict:
    return {
        "kp_number": "КП-INT-001",
        "model_id": model_id,
        "data": {
            "model": {"line": line, "max": max_t, "length": length, "price": price},
            "equipment": {"sensor_id": "zemic_dhm9b_30t", "indicator_id": "titan_3cs"},
            "options": options or {},
            "payment": {"preset_id": "split_by_items", "days": 5, "split_state": {}},
            "flags": flags or {},
            "delivery_address": "г. Тест",
        },
    }


class TestExample1_SL40_DeliveryOnly:
    """ВЕСТА-СЛ-40-18 — только поставка.

    Ожидаем: 1 строку сроков, ТТХ СЛ-40, секция 7, нет контр. листа.
    """

    def test_no_unfilled_placeholders(self, tmp_path, json_data):
        models, prices, pt, es = json_data
        kp = _make_kp("vesta-сл-40-18", "СЛ", 40, 18, 1_900_000)
        data, items, deal = build_spec_v2_data(kp, prices, models, pt, es)
        data.update(_REQUISITES)
        output = str(tmp_path / "ex1.docx")
        fill_spec_v2(SPEC_V2_PATH, data, items, deal, output)
        unfilled = get_unfilled_placeholders(output)
        assert unfilled == [], f"Unfilled: {unfilled}"

    def test_one_term_line(self, tmp_path, json_data):
        models, prices, pt, es = json_data
        kp = _make_kp("vesta-сл-40-18", "СЛ", 40, 18, 1_900_000)
        data, items, deal = build_spec_v2_data(kp, prices, models, pt, es)
        assert len(data["_terms_lines"]) == 1

    def test_final_section_only(self, tmp_path, json_data):
        models, prices, pt, es = json_data
        kp = _make_kp("vesta-сл-40-18", "СЛ", 40, 18, 1_900_000)
        data, items, deal = build_spec_v2_data(kp, prices, models, pt, es)
        data.update(_REQUISITES)
        output = str(tmp_path / "ex1b.docx")
        fill_spec_v2(SPEC_V2_PATH, data, items, deal, output)
        doc = Document(output)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "7. Заключительные положения" in all_text
        assert "Контрольный лист" not in all_text


class TestExample2_C80_ContractorFull:
    """ВЕСТА-С-80-18 — contractor_full + монтаж.

    Ожидаем: 3 строки сроков, Прил.№1, секции 4-5-7.
    """

    def test_no_unfilled_placeholders(self, tmp_path, json_data):
        models, prices, pt, es = json_data
        kp = _make_kp("vesta-с-80-18", "С", 80, 18, 3_700_000, options={
            "foundation_s_f_18": {"qty": 1, "price": 350_000, "customer_side": False},
            "install_default": {"qty": 1, "price": 80_000, "customer_side": False},
            "verification_default": {"qty": 1, "price": 30_000, "customer_side": False},
            "delivery_default": {"qty": 1, "price": 50_000, "customer_side": False},
        })
        data, items, deal = build_spec_v2_data(kp, prices, models, pt, es)
        data.update(_REQUISITES)
        output = str(tmp_path / "ex2.docx")
        fill_spec_v2(SPEC_V2_PATH, data, items, deal, output)
        unfilled = get_unfilled_placeholders(output)
        assert unfilled == [], f"Unfilled: {unfilled}"

    def test_three_term_lines(self, tmp_path, json_data):
        models, prices, pt, es = json_data
        kp = _make_kp("vesta-с-80-18", "С", 80, 18, 3_700_000, options={
            "foundation_s_f_18": {"qty": 1, "price": 350_000, "customer_side": False},
            "install_default": {"qty": 1, "price": 80_000, "customer_side": False},
            "verification_default": {"qty": 1, "price": 30_000, "customer_side": False},
        })
        data, items, deal = build_spec_v2_data(kp, prices, models, pt, es)
        assert len(data["_terms_lines"]) == 3

    def test_sections_4_5_7(self, tmp_path, json_data):
        models, prices, pt, es = json_data
        kp = _make_kp("vesta-с-80-18", "С", 80, 18, 3_700_000, options={
            "foundation_s_f_18": {"qty": 1, "price": 350_000, "customer_side": False},
            "install_default": {"qty": 1, "price": 80_000, "customer_side": False},
            "verification_default": {"qty": 1, "price": 30_000, "customer_side": False},
        })
        data, items, deal = build_spec_v2_data(kp, prices, models, pt, es)
        data.update(_REQUISITES)
        output = str(tmp_path / "ex2b.docx")
        fill_spec_v2(SPEC_V2_PATH, data, items, deal, output)
        doc = Document(output)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "4. Обязательства Подрядчика" in all_text
        assert "5. Обязательства Заказчика" in all_text
        assert "7. Заключительные положения" in all_text


class TestExample3_C100_CustomerBuilds:
    """ВЕСТА-С-100-24 — customer_builds + монтаж + ОРИОН.

    Ожидаем: Прил.№1 + контр. лист, все 4 секции.
    """

    def test_no_unfilled_placeholders(self, tmp_path, json_data):
        models, prices, pt, es = json_data
        kp = _make_kp("vesta-с-100-24", "С", 100, 24, 5_500_000, options={
            "install_default": {"qty": 1, "price": 100_000, "customer_side": False},
            "verification_default": {"qty": 1, "price": 40_000, "customer_side": False},
            "orion_turnkey_1": {"qty": 1, "price": 200_000, "customer_side": False},
        }, flags={"foundation_scope_override": "customer_builds"})
        data, items, deal = build_spec_v2_data(kp, prices, models, pt, es)
        deal["scope_overrides"] = {"foundation_scope": "customer_builds"}
        data.update(_REQUISITES)
        output = str(tmp_path / "ex3.docx")
        fill_spec_v2(SPEC_V2_PATH, data, items, deal, output)
        unfilled = get_unfilled_placeholders(output)
        assert unfilled == [], f"Unfilled: {unfilled}"

    def test_foundation_check_present(self, tmp_path, json_data):
        models, prices, pt, es = json_data
        kp = _make_kp("vesta-с-100-24", "С", 100, 24, 5_500_000, options={
            "install_default": {"qty": 1, "price": 100_000, "customer_side": False},
            "verification_default": {"qty": 1, "price": 40_000, "customer_side": False},
        })
        data, items, deal = build_spec_v2_data(kp, prices, models, pt, es)
        deal["scope_overrides"] = {"foundation_scope": "customer_builds"}
        data.update(_REQUISITES)
        output = str(tmp_path / "ex3b.docx")
        fill_spec_v2(SPEC_V2_PATH, data, items, deal, output)
        doc = Document(output)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Контрольный лист" in all_text

    def test_all_four_sections(self, tmp_path, json_data):
        models, prices, pt, es = json_data
        kp = _make_kp("vesta-с-100-24", "С", 100, 24, 5_500_000, options={
            "install_default": {"qty": 1, "price": 100_000, "customer_side": False},
            "verification_default": {"qty": 1, "price": 40_000, "customer_side": False},
            "orion_turnkey_1": {"qty": 1, "price": 200_000, "customer_side": False},
        })
        data, items, deal = build_spec_v2_data(kp, prices, models, pt, es)
        deal["scope_overrides"] = {"foundation_scope": "customer_builds"}
        data.update(_REQUISITES)
        output = str(tmp_path / "ex3c.docx")
        fill_spec_v2(SPEC_V2_PATH, data, items, deal, output)
        doc = Document(output)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "4. Обязательства Подрядчика" in all_text
        assert "5. Обязательства Заказчика" in all_text
        assert "6. Особые условия" in all_text
        assert "7. Заключительные положения" in all_text
