"""Тесты конвертера PDF-чертежей фундаментов в DOCX-приложения."""
from docx import Document
from docx.oxml.ns import qn

from scripts.pdf_to_fundament_docx import (
    FIRST_PAGE_IMAGE_MAX_H_EMU,
    NEXT_PAGE_IMAGE_MAX_H_EMU,
    TEXTBOX_ANCHOR_BASE_ID,
    TEXTBOX_DOC_PR_BASE_ID,
    TEXTBOX_HEIGHT_EMU,
    TEXTBOX_WIDTH_EMU,
    USABLE_W_EMU,
    WPS_TXBX_TAG,
    _clone_textbox,
    _image_emu,
    _load_reference_parts,
    _render_pages,
    convert,
)
from tests.contracts.pdf_to_fundament_helpers import (
    REFERENCE,
    WP14_ANCHOR_ID_ATTR,
    XML_NS,
    docx_root,
    docx_xml,
    make_a3_landscape_pdf,
    png_size,
)


def _xml_text(element) -> str:
    """Собрать текст из XML-элемента DOCX."""
    return "".join(node.text or "" for node in element.iter(qn("w:t")))


def test_load_reference_parts_copies_appendix_block_without_check_marker():
    doc = Document(REFERENCE)

    header_block, img_ppr, textbox_drawing = _load_reference_parts(doc)

    assert [_xml_text(p).lstrip("\t") for p in header_block] == [
        "Приложение №{{ПРИЛОЖЕНИЕ_НОМЕР}} к Спецификации №{{СПЕЦ_НОМЕР}} "
        "от {{ДОГОВОР_ДАТА_ПОЛНАЯ}} г.",
        "",
        "",
        "Строительное задание на фундамент Весов",
        "",
    ]
    assert "APPENDIX_FOUNDATION_CHECK" not in "\n".join(_xml_text(p) for p in header_block)

    assert img_ppr.tag == qn("w:pPr")
    assert textbox_drawing.tag == qn("w:drawing")
    assert textbox_drawing.find(".//" + WPS_TXBX_TAG) is not None
    assert "{{ЗАКАЗЧИК_ДИРЕКТОР_ИНИЦИАЛЫ}}" in _xml_text(textbox_drawing)


def test_render_pages_returns_png_bytes_for_each_pdf_page(tmp_path):
    pdf_path = tmp_path / "a3_landscape.pdf"
    make_a3_landscape_pdf(pdf_path, pages=2)

    pages = _render_pages(pdf_path)

    assert len(pages) == 2
    for png in pages:
        width, height = png_size(png)
        assert width > 3000
        assert height > 2000
        assert 1.40 < width / height < 1.43


def test_image_emu_uses_full_width_for_a3_landscape():
    width, height = _image_emu(4200, 2970, page_idx=0)

    assert width == USABLE_W_EMU
    assert height == round(USABLE_W_EMU * 2970 / 4200)
    assert height < FIRST_PAGE_IMAGE_MAX_H_EMU


def test_image_emu_caps_height_for_tall_pages():
    first_width, first_height = _image_emu(1000, 3000, page_idx=0)
    next_width, next_height = _image_emu(1000, 3000, page_idx=1)

    assert first_height == FIRST_PAGE_IMAGE_MAX_H_EMU
    assert first_width == round(FIRST_PAGE_IMAGE_MAX_H_EMU * 1000 / 3000)
    assert next_height == NEXT_PAGE_IMAGE_MAX_H_EMU
    assert next_width == round(NEXT_PAGE_IMAGE_MAX_H_EMU * 1000 / 3000)
    assert first_width <= USABLE_W_EMU
    assert next_width <= USABLE_W_EMU


def test_clone_textbox_preserves_stamp_and_sets_unique_ids():
    doc = Document(REFERENCE)
    _, _, textbox_drawing = _load_reference_parts(doc)

    first = _clone_textbox(textbox_drawing, page_idx=0)
    second = _clone_textbox(textbox_drawing, page_idx=1)

    first_doc_pr = first.find(".//" + qn("wp:docPr"))
    second_doc_pr = second.find(".//" + qn("wp:docPr"))
    assert first_doc_pr is not None
    assert second_doc_pr is not None
    assert first_doc_pr.get("id") == str(TEXTBOX_DOC_PR_BASE_ID)
    assert second_doc_pr.get("id") == str(TEXTBOX_DOC_PR_BASE_ID + 1)

    first_anchor = first.find(".//" + qn("wp:anchor"))
    second_anchor = second.find(".//" + qn("wp:anchor"))
    assert first_anchor is not None
    assert second_anchor is not None
    assert first_anchor.get(WP14_ANCHOR_ID_ATTR) == f"{TEXTBOX_ANCHOR_BASE_ID:08X}"
    assert second_anchor.get(WP14_ANCHOR_ID_ATTR) == f"{TEXTBOX_ANCHOR_BASE_ID + 1:08X}"

    assert "{{ЗАКАЗЧИК_ДИРЕКТОР_ИНИЦИАЛЫ}}" in _xml_text(first)
    assert "{{ЗАКАЗЧИК_ДИРЕКТОР_ИНИЦИАЛЫ}}" in _xml_text(second)
    assert first.find(".//" + WPS_TXBX_TAG) is not None
    assert first.find(".//" + qn("wp:positionH") + "/" + qn("wp:posOffset")).text == "642620"
    assert first.find(".//" + qn("wp:positionV") + "/" + qn("wp:posOffset")).text == "4172585"


def test_convert_builds_docx_from_reference_with_two_pdf_pages(tmp_path):
    pdf_path = tmp_path / "foundation.pdf"
    out_path = tmp_path / "foundation.docx"
    make_a3_landscape_pdf(pdf_path, pages=2)

    convert(pdf_path, out_path)

    doc = Document(out_path)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Приложение №{{ПРИЛОЖЕНИЕ_НОМЕР}}" in all_text
    assert "Строительное задание на фундамент Весов" in all_text
    assert "APPENDIX_FOUNDATION_CHECK" not in all_text

    import zipfile

    with zipfile.ZipFile(out_path) as docx:
        names = set(docx.namelist())
    assert "word/header1.xml" in names

    document_xml = docx_xml(out_path, "word/document.xml")
    header_text = "\n".join(p.text for p in doc.sections[0].header.paragraphs)
    assert "{{ДОГОВОР_НОМЕР}}" in header_text
    assert "{{ДОГОВОР_ДАТА_ПОЛНАЯ}}" in header_text
    assert "w:headerReference" in document_xml
    assert document_xml.count("<wp:inline") == 2
    assert document_xml.count("<wps:txbx") == 2
    assert document_xml.count('w:type="page"') == 1
    root = docx_root(out_path, "word/document.xml")
    paragraphs = root.xpath("//w:body/w:p", namespaces=XML_NS)
    image_idxs = [i for i, p in enumerate(paragraphs) if p.xpath(".//wp:inline", namespaces=XML_NS)]
    break_idxs = [i for i, p in enumerate(paragraphs) if p.xpath(".//w:br[@w:type='page']", namespaces=XML_NS)]
    assert break_idxs == image_idxs[1:]
    assert not root.xpath("//w:pPr/w:pageBreakBefore", namespaces=XML_NS)
    assert document_xml.count("APPENDIX_FOUNDATION_CHECK") == 0


def test_convert_preserves_page_geometry_and_image_limits(tmp_path):
    pdf_path = tmp_path / "foundation.pdf"
    out_path = tmp_path / "foundation.docx"
    make_a3_landscape_pdf(pdf_path, pages=2)

    convert(pdf_path, out_path)

    root = docx_root(out_path, "word/document.xml")
    pg_sz = root.find(".//w:sectPr/w:pgSz", namespaces=XML_NS)
    pg_mar = root.find(".//w:sectPr/w:pgMar", namespaces=XML_NS)
    assert pg_sz is not None
    assert pg_mar is not None
    assert pg_sz.get(qn("w:w")) == "11906"
    assert pg_sz.get(qn("w:h")) == "16838"
    assert pg_mar.get(qn("w:left")) == "1418"
    assert pg_mar.get(qn("w:right")) == "424"
    assert pg_mar.get(qn("w:top")) == "567"
    assert pg_mar.get(qn("w:bottom")) == "0"

    image_extents = root.xpath(".//wp:inline/wp:extent", namespaces=XML_NS)
    assert len(image_extents) == 2
    for page_idx, extent in enumerate(image_extents):
        max_h = FIRST_PAGE_IMAGE_MAX_H_EMU if page_idx == 0 else NEXT_PAGE_IMAGE_MAX_H_EMU
        assert int(extent.get("cx")) <= USABLE_W_EMU
        assert int(extent.get("cy")) <= max_h

    textbox_anchors = root.xpath(".//wp:anchor[.//wps:txbx]", namespaces=XML_NS)
    assert len(textbox_anchors) == 2
    doc_pr_ids = [int(a.find("wp:docPr", namespaces=XML_NS).get("id")) for a in textbox_anchors]
    anchor_ids = [anchor.get(WP14_ANCHOR_ID_ATTR) for anchor in textbox_anchors]
    assert doc_pr_ids == [TEXTBOX_DOC_PR_BASE_ID, TEXTBOX_DOC_PR_BASE_ID + 1]
    assert anchor_ids == [
        f"{TEXTBOX_ANCHOR_BASE_ID:08X}",
        f"{TEXTBOX_ANCHOR_BASE_ID + 1:08X}",
    ]
    for anchor in textbox_anchors:
        assert anchor.find("wp:positionH/wp:posOffset", namespaces=XML_NS).text == "642620"
        assert anchor.find("wp:positionV/wp:posOffset", namespaces=XML_NS).text == "4172585"
        wp_extent = anchor.find("wp:extent", namespaces=XML_NS)
        shape_extent = anchor.find(".//a:xfrm/a:ext", namespaces=XML_NS)
        wp_size = (int(wp_extent.get("cx")), int(wp_extent.get("cy")))
        shape_size = (int(shape_extent.get("cx")), int(shape_extent.get("cy")))
        assert wp_size == shape_size == (TEXTBOX_WIDTH_EMU, TEXTBOX_HEIGHT_EMU)
        assert 1_933_575 < wp_size[0] and 90 * 36_000 <= wp_size[0] <= 110 * 36_000
        placeholder = "{{\u0417\u0410\u041a\u0410\u0417\u0427\u0418\u041a_\u0414\u0418\u0420\u0415\u041a\u0422\u041e\u0420_\u0418\u041d\u0418\u0426\u0418\u0410\u041b\u042b}}"
        assert placeholder in _xml_text(anchor)
