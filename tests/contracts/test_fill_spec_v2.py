"""Тесты fill_spec_v2 — спецификация с динамическими clauses."""
import os
import re

import pytest
from docx import Document

from src.contracts.spec_v2_filler import fill_spec_v2

SPEC_V2_PATH = os.path.normpath(os.path.join(
    os.path.dirname(__file__), '..', '..', 'templates', 'contracts', 'spec_v2.docx'
))

MOCK_DATA = {
    "ДОГОВОР_НОМЕР": "Т-001/2026",
    "ДОГОВОР_ДАТА_ПОЛНАЯ": "15.03.2026",
    "ДОГОВОР_ДЕНЬ": "15",
    "ДОГОВОР_МЕСЯЦ": "марта",
    "ДОГОВОР_ГОД": "2026",
    "СПЕЦ_НОМЕР": "1",
    "СПЕЦ_НДС": "22",
    "СПЕЦ_ОПЛАТА_П1": "Предоплата 30%",
    "СПЕЦ_ОПЛАТА_П2": "Оплата по отгрузке 70%",
    "СПЕЦ_ОПЛАТА_П3": "",
    "СПЕЦ_ОПЛАТА_П4": "",
    "СПЕЦ_ОПЛАТА_П5": "",
    "СПЕЦ_ОПЛАТА_П6": "",
    "СПЕЦ_СРОК_ПОСТАВКИ": "30",
    "СПЕЦ_СРОК_ФУНДАМЕНТ": "20",
    "СПЕЦ_СРОК_МОНТАЖ": "10",
    "СПЕЦ_АДРЕС_ОБЪЕКТА": "г. Тест",
    "СПЕЦ_МОДЕЛЬ_КРАТКОЕ": "ВЕСТА-С-60-18",
    "СПЕЦ_МАКС_НАГРУЗКА": "60",
    "ЗАКАЗЧИК_ДИРЕКТОР_ДОЛЖНОСТЬ": "Директор",
    "ЗАКАЗЧИК_ПОЛНОЕ_НАИМЕНОВАНИЕ": "ООО «Тест»",
    "ЗАКАЗЧИК_КРАТКОЕ_НАИМЕНОВАНИЕ": "ООО «Тест»",
    "ЗАКАЗЧИК_ДИРЕКТОР_ФИО_КРАТКОЕ": "Тестов Т.Т.",
    "ЗАКАЗЧИК_ДИРЕКТОР_ИНИЦИАЛЫ": "Т.Т. Тестов",
    "ЗАКАЗЧИК_ДИРЕКТОР_ДОЛЖНОСТЬ_РП": "директора",
    "ЗАКАЗЧИК_ДИРЕКТОР_ФИО_РП": "Тестова Тимофея Тимофеевича",
    "ЗАКАЗЧИК_ОСНОВАНИЕ": "Устава",
    "ДИРЕКТОР_ПРИЧАСТИЕ": "действующего",
}


def _item(id: str, metadata: dict | None = None) -> dict:
    return {
        "id": id, "name": f"Позиция {id}", "unit": "компл",
        "quantity": 1.0, "price_per_unit": 100_000.0, "total": 100_000.0,
        "payment_group": None, "is_custom": False, "source": "preset",
        "metadata": metadata or {},
    }


_CLAUSE_HEADER_TITLES = (
    "Обязательства Подрядчика",
    "Обязательства Заказчика",
    "Особые условия",
    "Заключительные положения",
)


def _all_text(doc: Document) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


def _flat_clause_numbers(doc: Document) -> list[int]:
    numbers: list[int] = []
    in_clauses_block = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if "Выезд строительной" in text:
            in_clauses_block = True
            continue
        if "Технические характеристики" in text:
            break
        if not in_clauses_block:
            continue
        if "Технические характеристики" in text or "Комплект поставки" in text:
            continue
        match = re.match(r"^(\d+)\.\s", text)
        if match:
            numbers.append(int(match.group(1)))
    return numbers


def _heading_text(doc: Document, needle: str) -> str:
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            return paragraph.text.strip()
    raise AssertionError(f"Параграф не найден: {needle}")


def _assert_no_clause_headers(all_text: str) -> None:
    for title in _CLAUSE_HEADER_TITLES:
        assert title not in all_text


class TestFillSpecV2Minimal:
    """Минимальный кейс: поставка без монтажа → только секция 7 (final)."""

    def test_only_final_section(self, tmp_path):
        items = [_item("weights"), _item("delivery")]
        deal = {"items": items, "delivery_address": "г. Тест"}
        output = str(tmp_path / "spec_v2_min.docx")

        fill_spec_v2(SPEC_V2_PATH, MOCK_DATA, items, deal, output)

        doc = Document(output)
        all_text = _all_text(doc)

        assert _flat_clause_numbers(doc) == [4, 5]
        assert _heading_text(doc, "Технические характеристики").startswith("6. ")
        assert _heading_text(doc, "Комплект поставки").startswith("7. ")
        _assert_no_clause_headers(all_text)

    def test_markers_removed(self, tmp_path):
        items = [_item("weights")]
        deal = {"items": items, "delivery_address": "г. Тест"}
        output = str(tmp_path / "spec_v2_markers.docx")

        fill_spec_v2(SPEC_V2_PATH, MOCK_DATA, items, deal, output)

        doc = Document(output)
        all_text = _all_text(doc)
        assert "CLAUSE_SECTION" not in all_text

    def test_items_table_has_correct_rows(self, tmp_path):
        items = [_item("weights"), _item("delivery")]
        deal = {"items": items, "delivery_address": "г. Тест"}
        output = str(tmp_path / "spec_v2_table.docx")

        fill_spec_v2(SPEC_V2_PATH, MOCK_DATA, items, deal, output)

        doc = Document(output)
        table = doc.tables[0]
        # header + 2 items + total = 4
        assert len(table.rows) == 4


class TestFillSpecV2CellContent:
    """Проверка содержимого ячеек таблицы позиций (Item 2 fix)."""

    def test_items_cell_names(self, tmp_path):
        items = [_item("weights"), _item("delivery")]
        deal = {"items": items, "delivery_address": "г. Тест"}
        output = str(tmp_path / "spec_v2_names.docx")

        fill_spec_v2(SPEC_V2_PATH, MOCK_DATA, items, deal, output)

        doc = Document(output)
        table = doc.tables[0]
        assert table.rows[1].cells[0].text == "Позиция weights"
        assert table.rows[2].cells[0].text == "Позиция delivery"

    def test_items_cell_amounts(self, tmp_path):
        items = [_item("weights"), _item("delivery")]
        deal = {"items": items, "delivery_address": "г. Тест"}
        output = str(tmp_path / "spec_v2_amounts.docx")

        fill_spec_v2(SPEC_V2_PATH, MOCK_DATA, items, deal, output)

        doc = Document(output)
        table = doc.tables[0]
        assert "100 000" in table.rows[1].cells[1].text
        assert "100 000" in table.rows[2].cells[1].text

    def test_customer_side_shows_zakazchik(self, tmp_path):
        items = [
            _item("weights"),
            _item("delivery", {"customer_side": True}),
        ]
        deal = {"items": items, "delivery_address": "г. Тест"}
        output = str(tmp_path / "spec_v2_cust.docx")

        fill_spec_v2(SPEC_V2_PATH, MOCK_DATA, items, deal, output)

        doc = Document(output)
        table = doc.tables[0]
        assert "ЗАКАЗЧИК" in table.rows[2].cells[1].text


class TestFillSpecV2Medium:
    """Средний кейс: монтаж → секции 4, 5, 7 (7 пунктов)."""

    def _make_deal(self):
        items = [
            _item("weights"),
            _item("installation", {"scope": "full"}),
            _item("verification"),
        ]
        return {
            "items": items,
            "scope_overrides": {},
            "flags": {},
            "delivery_address": "г. Кемерово, пр-т Кузнецкий, 15",
        }, items

    def test_clause_section_headers_absent(self, tmp_path):
        deal, items = self._make_deal()
        output = str(tmp_path / "spec_v2_med.docx")

        fill_spec_v2(SPEC_V2_PATH, MOCK_DATA, items, deal, output)

        doc = Document(output)
        all_text = _all_text(doc)

        _assert_no_clause_headers(all_text)
        assert _flat_clause_numbers(doc) == list(range(4, 11))

    def test_seven_clause_numbers(self, tmp_path):
        deal, items = self._make_deal()
        output = str(tmp_path / "spec_v2_med_cnt.docx")

        fill_spec_v2(SPEC_V2_PATH, MOCK_DATA, items, deal, output)

        doc = Document(output)
        all_text = _all_text(doc)

        assert _flat_clause_numbers(doc) == list(range(4, 11))
        assert "4.1." not in all_text

    def test_delivery_address_substituted(self, tmp_path):
        deal, items = self._make_deal()
        output = str(tmp_path / "spec_v2_med_addr.docx")

        fill_spec_v2(SPEC_V2_PATH, MOCK_DATA, items, deal, output)

        doc = Document(output)
        all_text = _all_text(doc)
        assert "г. Кемерово" in all_text


class TestFillSpecV2FoundationInstall:
    """Контроль: фундамент + монтаж + поверка → 4-13, ТТХ 14, комплект 15."""

    def _make_deal(self):
        items = [
            _item("weights"),
            _item("foundation", {"scope": "fundament_jb"}),
            _item("installation", {"scope": "full"}),
            _item("verification"),
        ]
        return {
            "items": items,
            "scope_overrides": {},
            "flags": {},
            "delivery_address": "г. Тест",
        }, items

    def test_flat_numbering_cross_reference_tth_and_kit(self, tmp_path):
        deal, items = self._make_deal()
        output = str(tmp_path / "spec_v2_foundation_install.docx")

        fill_spec_v2(SPEC_V2_PATH, MOCK_DATA, items, deal, output)

        doc = Document(output)
        all_text = _all_text(doc)

        _assert_no_clause_headers(all_text)
        assert _flat_clause_numbers(doc) == list(range(4, 14))
        assert "п.п. 7-11" in all_text
        assert _heading_text(doc, "Технические характеристики").startswith("14. ")
        assert _heading_text(doc, "Комплект поставки").startswith("15. ")


class TestPaymentSection:
    """Маркер {{PAYMENT_SECTION}} → строки оплаты."""

    def test_payment_lines_in_output(self, tmp_path):
        data = dict(MOCK_DATA, _payment_lines=[
            "2.1 Предоплата 30% в течение 5 дней",
            "2.2 Оплата по отгрузке 70%",
        ])
        items = [_item("weights")]
        deal = {"items": items, "delivery_address": "г. Тест"}
        output = str(tmp_path / "spec_v2_pay.docx")
        fill_spec_v2(SPEC_V2_PATH, data, items, deal, output)
        doc = Document(output)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Предоплата 30%" in all_text
        assert "Оплата по отгрузке" in all_text

    def test_payment_marker_removed(self, tmp_path):
        data = dict(MOCK_DATA, _payment_lines=["Оплата 100%"])
        items = [_item("weights")]
        deal = {"items": items, "delivery_address": "г. Тест"}
        output = str(tmp_path / "spec_v2_pay2.docx")
        fill_spec_v2(SPEC_V2_PATH, data, items, deal, output)
        doc = Document(output)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "PAYMENT_SECTION" not in all_text


class TestTermsSection:
    """Маркер {{TERMS_SECTION}} → строки сроков."""

    def test_terms_lines_in_output(self, tmp_path):
        data = dict(MOCK_DATA, _terms_lines=[
            "- поставка Весов: в течение 20 рабочих дней",
            "- монтаж: в течение 3 рабочих дней",
        ])
        items = [_item("weights")]
        deal = {"items": items, "delivery_address": "г. Тест"}
        output = str(tmp_path / "spec_v2_terms.docx")
        fill_spec_v2(SPEC_V2_PATH, data, items, deal, output)
        doc = Document(output)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "20 рабочих дней" in all_text
        assert "TERMS_SECTION" not in all_text


class TestKitSection:
    """Таблица комплекта поставки — клонирование строк."""

    def test_kit_rows(self, tmp_path):
        kit = [
            {"name": "Платформа", "qty": "1"},
            {"name": "Датчик", "qty": "8"},
            {"name": "Терминал", "qty": "1"},
        ]
        data = dict(MOCK_DATA, _kit_items=kit)
        items = [_item("weights")]
        deal = {"items": items, "delivery_address": "г. Тест"}
        output = str(tmp_path / "spec_v2_kit.docx")
        fill_spec_v2(SPEC_V2_PATH, data, items, deal, output)
        doc = Document(output)
        kit_table = doc.tables[2]
        # header + 3 kit items = 4
        assert len(kit_table.rows) == 4

    def test_kit_cell_text(self, tmp_path):
        kit = [{"name": "Платформа сплошного типа", "qty": "1"}]
        data = dict(MOCK_DATA, _kit_items=kit)
        items = [_item("weights")]
        deal = {"items": items, "delivery_address": "г. Тест"}
        output = str(tmp_path / "spec_v2_kit2.docx")
        fill_spec_v2(SPEC_V2_PATH, data, items, deal, output)
        doc = Document(output)
        kit_table = doc.tables[2]
        assert "Платформа" in kit_table.rows[1].cells[0].text


class TestFoundationCheck:
    """Контрольный лист — при customer_builds."""

    def test_check_present_when_customer_builds(self, tmp_path):
        data = dict(MOCK_DATA)
        items = [_item("weights")]
        deal = {
            "items": items,
            "scope_overrides": {"foundation_scope": "customer_builds"},
            "flags": {},
            "delivery_address": "г. Тест",
        }
        output = str(tmp_path / "spec_v2_check.docx")
        fill_spec_v2(SPEC_V2_PATH, data, items, deal, output)
        doc = Document(output)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Контрольный лист" in all_text
        assert "APPENDIX_FOUNDATION_CHECK" not in all_text

    def test_check_absent_when_contractor(self, tmp_path):
        data = dict(MOCK_DATA)
        items = [
            _item("weights"),
            _item("foundation", {"scope": "fundament_jb"}),
        ]
        deal = {
            "items": items,
            "scope_overrides": {},
            "flags": {},
            "delivery_address": "г. Тест",
        }
        output = str(tmp_path / "spec_v2_nocheck.docx")
        fill_spec_v2(SPEC_V2_PATH, data, items, deal, output)
        doc = Document(output)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Контрольный лист" not in all_text
        assert "APPENDIX_FOUNDATION_CHECK" not in all_text


class TestTTXSection:
    """ТТХ плейсхолдеры подставлены в таблицу."""

    def test_tth_values_in_table(self, tmp_path):
        data = dict(MOCK_DATA, **{
            "ТТХ_НАГРУЗКА_НА_ОСЬ": "14",
            "ТТХ_РАССТОЯНИЕ_ДО_ТЕРМИНАЛА": "не более 50 м",
            "ТТХ_ДИСКРЕТНОСТЬ_БЛОК": "10\n20",
            "ТТХ_ГАБАРИТЫ": "18×3",
            "ТТХ_ТЕМПЕРАТУРА": "От -30 до +40",
        })
        items = [_item("weights")]
        deal = {"items": items, "delivery_address": "г. Тест"}
        output = str(tmp_path / "spec_v2_tth.docx")
        fill_spec_v2(SPEC_V2_PATH, data, items, deal, output)
        doc = Document(output)
        tth_table = doc.tables[1]
        assert tth_table.rows[2].cells[2].text == "14"
        assert "50 м" in tth_table.rows[3].cells[2].text
        assert "18×3" in tth_table.rows[5].cells[2].text
        assert "-30" in tth_table.rows[6].cells[2].text

    def test_no_hardcoded_11(self, tmp_path):
        data = dict(MOCK_DATA, **{
            "ТТХ_НАГРУЗКА_НА_ОСЬ": "14",
            "ТТХ_РАССТОЯНИЕ_ДО_ТЕРМИНАЛА": "не более 50 м",
            "ТТХ_ДИСКРЕТНОСТЬ_БЛОК": "10\n20",
            "ТТХ_ГАБАРИТЫ": "18×3",
            "ТТХ_ТЕМПЕРАТУРА": "От -30 до +40",
        })
        items = [_item("weights")]
        deal = {"items": items, "delivery_address": "г. Тест"}
        output = str(tmp_path / "spec_v2_tth2.docx")
        fill_spec_v2(SPEC_V2_PATH, data, items, deal, output)
        doc = Document(output)
        # Старый хардкод "11" для нагрузки на ось не должен остаться
        assert doc.tables[1].rows[2].cells[2].text != "11"

    def test_partial_tth_overrides_fill_missing_boundaries(self, tmp_path):
        data = dict(MOCK_DATA, **{
            "ТТХ_НАГРУЗКА_НА_ОСЬ": "18",
            "ТТХ_ДИСКРЕТНОСТЬ_1": "20",
            "ТТХ_ДИСКРЕТНОСТЬ_2": "50",
        })
        items = [_item("weights", {"line": "С", "max": 80, "length": 18})]
        deal = {"items": items, "delivery_address": "г. Тест"}
        output = str(tmp_path / "spec_v2_tth_partial.docx")

        fill_spec_v2(SPEC_V2_PATH, data, items, deal, output)

        doc = Document(output)
        discreteness_label = doc.tables[1].rows[4].cells[1].text
        assert "ТТХ_ГРАНИЦА" not in discreteness_label
        assert "от 0 до 60т" in discreteness_label
        assert "от 60 до 80т" in discreteness_label


class TestYearPlaceholder:
    """Год {{ТЕКУЩИЙ_ГОД}} заменяется."""

    def test_no_hardcoded_2026(self, tmp_path):
        data = dict(MOCK_DATA, **{"ТЕКУЩИЙ_ГОД": "2025"})
        items = [_item("weights")]
        deal = {"items": items, "delivery_address": "г. Тест"}
        output = str(tmp_path / "spec_v2_year.docx")
        fill_spec_v2(SPEC_V2_PATH, data, items, deal, output)
        doc = Document(output)
        # Table 3 (signatures) should not have hardcoded 2026
        sig_text = doc.tables[3].rows[0].cells[0].text
        assert "2026" not in sig_text
        assert "2025" in sig_text


class TestAppendixNumber:
    """Приложение №{{ПРИЛОЖЕНИЕ_НОМЕР}} заполнен."""

    def test_appendix_filled(self, tmp_path):
        data = dict(MOCK_DATA, ПРИЛОЖЕНИЕ_НОМЕР="1")
        items = [_item("weights")]
        deal = {"items": items, "delivery_address": "г. Тест"}
        output = str(tmp_path / "spec_v2_app.docx")
        fill_spec_v2(SPEC_V2_PATH, data, items, deal, output)
        doc = Document(output)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Приложение №1" in all_text
        assert "ПРИЛОЖЕНИЕ_НОМЕР" not in all_text


class TestFillSpecV2Max:
    """Максимальный кейс: фундамент+монтаж+ОРИОН → 14 плоских пунктов."""

    def _make_deal(self):
        items = [
            _item("weights"),
            _item("foundation", {"scope": "fundament_jb"}),
            _item("installation", {"scope": "full"}),
            _item("verification"),
            _item("orion"),
            _item("orion_install"),
        ]
        return {
            "items": items,
            "scope_overrides": {},
            "flags": {},
            "delivery_address": "г. Тест",
        }, items

    def test_clause_section_headers_absent(self, tmp_path):
        deal, items = self._make_deal()
        output = str(tmp_path / "spec_v2_max.docx")

        fill_spec_v2(SPEC_V2_PATH, MOCK_DATA, items, deal, output)

        doc = Document(output)
        all_text = _all_text(doc)

        _assert_no_clause_headers(all_text)

    def test_fourteen_clauses(self, tmp_path):
        deal, items = self._make_deal()
        output = str(tmp_path / "spec_v2_max_cnt.docx")

        fill_spec_v2(SPEC_V2_PATH, MOCK_DATA, items, deal, output)

        doc = Document(output)
        assert _flat_clause_numbers(doc) == list(range(4, 18))

    def test_key_clause_texts(self, tmp_path):
        deal, items = self._make_deal()
        output = str(tmp_path / "spec_v2_max_txt.docx")

        fill_spec_v2(SPEC_V2_PATH, MOCK_DATA, items, deal, output)

        doc = Document(output)
        all_text = _all_text(doc)

        assert "Подрядчик обеспечивает подготовку" in all_text
        assert "автокран" in all_text
        assert "ОРИОН" in all_text
        assert "7-15" in all_text

    def test_items_table_intact(self, tmp_path):
        deal, items = self._make_deal()
        output = str(tmp_path / "spec_v2_max_tbl.docx")

        fill_spec_v2(SPEC_V2_PATH, MOCK_DATA, items, deal, output)

        doc = Document(output)
        table = doc.tables[0]
        # header + 6 items + total = 8
        assert len(table.rows) == 8
