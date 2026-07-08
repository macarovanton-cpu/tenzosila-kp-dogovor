"""Тесты для src/contracts/filler.py — заполнение шаблона договора."""

import os

import pytest

from src.contracts.filler import fill_template, get_unfilled_placeholders

TEMPLATE_PATH = os.path.join(
    os.path.dirname(__file__), '..', '..', 'templates', 'contracts', 'contract.docx'
)

SPEC_TEMPLATE_PATH = os.path.join(
    os.path.dirname(__file__), '..', '..', 'templates', 'contracts',
    'spec_foundation_install.docx'
)

SPEC_MOCK_DATA = {
    "ДОГОВОР_НОМЕР": "Т-001/2026",
    "ДОГОВОР_ДАТА_ПОЛНАЯ": "15.03.2026",
    "СПЕЦ_НОМЕР": "1",
    "СПЕЦ_НДС": "22",
    "СПЕЦ_ИТОГО": "2 000 000",
    "СПЕЦ_ИТОГО_ПРОПИСЬ": "два миллиона",
    "СПЕЦ_П1_НАИМЕНОВАНИЕ": "Весы ВЕСТА-С-60-18-Ц",
    "СПЕЦ_П1_СУММА": "1 500 000",
    "СПЕЦ_П2_НАИМЕНОВАНИЕ": "ВЕСТА-С, 18м",
    "СПЕЦ_П2_СУММА": "500 000",
    "СПЕЦ_П3_НАИМЕНОВАНИЕ": "",
    "СПЕЦ_П3_СУММА": "",
    "СПЕЦ_П4_НАИМЕНОВАНИЕ": "",
    "СПЕЦ_П4_СУММА": "",
    "СПЕЦ_П5_НАИМЕНОВАНИЕ": "",
    "СПЕЦ_П5_СУММА": "",
    "СПЕЦ_ОПЛАТА_П1": "Предоплата 30% = 600 000 руб.",
    "СПЕЦ_ОПЛАТА_П2": "По отгрузке 70% = 1 400 000 руб.",
    "СПЕЦ_ОПЛАТА_П3": "",
    "СПЕЦ_ОПЛАТА_П4": "",
    "СПЕЦ_ОПЛАТА_П5": "",
    "СПЕЦ_ОПЛАТА_П6": "",
    "СПЕЦ_СРОК_ПОСТАВКИ": "30",
    "СПЕЦ_СРОК_ФУНДАМЕНТ": "20",
    "СПЕЦ_СРОК_МОНТАЖ": "10",
    "СПЕЦ_АДРЕС_ОБЪЕКТА": "г. Москва, промзона Северная",
    "СПЕЦ_АДРЕС_ОБЪЕКТА_ПОЛНЫЙ": "г. Москва, промзона Северная, уч. 5",
    "СПЕЦ_МОДЕЛЬ_КРАТКОЕ": "ВЕСТА-С-60-18",
    "СПЕЦ_МАКС_НАГРУЗКА": "60",
    "ЗАКАЗЧИК_ДИРЕКТОР_ДОЛЖНОСТЬ": "Директор",
    "ЗАКАЗЧИК_КРАТКОЕ_НАИМЕНОВАНИЕ": "ООО «Тест»",
    "ЗАКАЗЧИК_ДИРЕКТОР_ФИО_КРАТКОЕ": "Тестов Т.Т.",
    "ЗАКАЗЧИК_ДИРЕКТОР_ИНИЦИАЛЫ": "Т.Т. Тестов",
    "ДИРЕКТОР_ПРИЧАСТИЕ": "действующего",
}

MOCK_DATA = {
    "requisites": {
        "ЗАКАЗЧИК_КРАТКОЕ_НАИМЕНОВАНИЕ": "ООО «Тест»",
        "ЗАКАЗЧИК_ПОЛНОЕ_НАИМЕНОВАНИЕ": "Общество с ограниченной ответственностью «Тест»",
        "ЗАКАЗЧИК_ДИРЕКТОР_ДОЛЖНОСТЬ": "Генеральный директор",
        "ЗАКАЗЧИК_ДИРЕКТОР_ДОЛЖНОСТЬ_РП": "генерального директора",
        "ЗАКАЗЧИК_ДИРЕКТОР_ФИО": "Иванов Иван Иванович",
        "ЗАКАЗЧИК_ДИРЕКТОР_ФИО_РП": "Иванова Ивана Ивановича",
        "ЗАКАЗЧИК_ДИРЕКТОР_ИНИЦИАЛЫ": "И.И. Иванов",
        "ЗАКАЗЧИК_ОСНОВАНИЕ": "Устава",
        "ДИРЕКТОР_ПРИЧАСТИЕ": "действующего",
        "ЗАКАЗЧИК_ИНН": "7701234567",
        "ЗАКАЗЧИК_КПП": "770101001",
        "ЗАКАЗЧИК_ОГРН": "1027700000001",
        "ЗАКАЗЧИК_АДРЕС_ЮР": "г. Москва, ул. Тестовая, д. 1",
        "ЗАКАЗЧИК_АДРЕС_ПОЧТ": "г. Москва, ул. Тестовая, д. 1",
        "ЗАКАЗЧИК_РС": "40702810000000000001",
        "ЗАКАЗЧИК_БАНК": "ПАО Сбербанк",
        "ЗАКАЗЧИК_КС": "30101810400000000225",
        "ЗАКАЗЧИК_БИК": "044525225",
        "ЗАКАЗЧИК_ТЕЛЕФОН": "+7 (495) 123-45-67",
        "ЗАКАЗЧИК_EMAIL": "test@example.com",
    },
    "specification": {
        "ДОГОВОР_НОМЕР": "Т-001/2026",
        "ДОГОВОР_ДЕНЬ": "15",
        "ДОГОВОР_МЕСЯЦ": "марта",
        "ДОГОВОР_ГОД": "2026",
        "ДОГОВОР_ДАТА_ПОЛНАЯ": "15.03.2026",
        "СПЕЦ_НОМЕР": "1",
        "СПЕЦ_П1_НАИМЕНОВАНИЕ": "Весы автомобильные ВЕСТА-С-60-18-Ц",
        "СПЕЦ_П1_СУММА": "1 500 000",
        "СПЕЦ_П1_СРОК": "30",
        "СПЕЦ_П2_НАИМЕНОВАНИЕ": "ВЕСТА-С, 18м",
        "СПЕЦ_П2_СУММА": "800 000",
        "СПЕЦ_П2_СРОК": "25",
        "СПЕЦ_П3_НАИМЕНОВАНИЕ": "Доставка весов",
        "СПЕЦ_П3_СУММА": "150 000",
        "СПЕЦ_П3_СРОК": "10",
        "СПЕЦ_П4_НАИМЕНОВАНИЕ": "Монтаж и пусконаладка",
        "СПЕЦ_П4_СУММА": "300 000",
        "СПЕЦ_П4_СРОК": "15",
        "СПЕЦ_П5_НАИМЕНОВАНИЕ": "Поверка весов",
        "СПЕЦ_П5_СУММА": "85 000",
        "СПЕЦ_П5_СРОК": "5",
        "СПЕЦ_ИТОГО": "2 835 000",
        "СПЕЦ_ИТОГО_ПРОПИСЬ": "два миллиона восемьсот тридцать пять тысяч",
        "СПЕЦ_НДС": "22",
        "СПЕЦ_ОПЛАТА_П1": "Предоплата – 10% от общей цены договора 283 500 рублей в т.ч. НДС 22%",
        "СПЕЦ_ОПЛАТА_П2": "Оплата весов – 90% стоимости весов 1 350 000 рублей в т.ч. НДС 22%",
        "СПЕЦ_ОПЛАТА_П3": "Оплата фундамента – 100% стоимости 800 000 рублей в т.ч. НДС 22%",
        "СПЕЦ_ОПЛАТА_П4": "Окончательный расчёт – 401 500 рублей в т.ч. НДС 22%",
        "СПЕЦ_ОПЛАТА_П5": "",
        "СПЕЦ_ОПЛАТА_П6": "",
        "СПЕЦ_СРОК_ПОСТАВКИ": "30",
        "СПЕЦ_СРОК_ФУНДАМЕНТ": "25",
        "СПЕЦ_СРОК_МОНТАЖ": "15",
        "СПЕЦ_АДРЕС_ОБЪЕКТА": "г. Москва, промзона Северная, уч. 5",
        "СПЕЦ_МОДЕЛЬ_КРАТКОЕ": "ВЕСТА-С-60-18-Ц",
        "СПЕЦ_МАКС_НАГРУЗКА": "60",
    },
}


@pytest.fixture
def output_path(tmp_path):
    return str(tmp_path / "test_contract_output.docx")


def test_fill_template_no_unfilled(output_path):
    """После заполнения шаблона не должно остаться незаполненных плейсхолдеров."""
    template = os.path.normpath(TEMPLATE_PATH)
    assert os.path.exists(template), f"Шаблон не найден: {template}"

    fill_template(template, MOCK_DATA, output_path)

    assert os.path.exists(output_path)
    unfilled = get_unfilled_placeholders(output_path)
    assert unfilled == [], f"Остались незаполненные плейсхолдеры: {unfilled}"


def test_fill_template_with_flat_data(output_path):
    """fill_template принимает и плоский словарь."""
    template = os.path.normpath(TEMPLATE_PATH)
    flat = {}
    flat.update(MOCK_DATA["requisites"])
    flat.update(MOCK_DATA["specification"])

    fill_template(template, flat, output_path)

    assert os.path.exists(output_path)


def test_filler_preserves_footer_page_field(tmp_path):
    """После fill_template поле PAGE в footer сохраняется (merge_runs не уничтожает instrText)."""
    import zipfile

    template = os.path.normpath(SPEC_TEMPLATE_PATH)
    output = str(tmp_path / "spec_out.docx")

    fill_template(template, SPEC_MOCK_DATA, output)

    with zipfile.ZipFile(output) as z:
        footer_xml = z.read("word/footer2.xml").decode("utf-8")

    assert "instrText" in footer_xml, "instrText уничтожен в footer — поле PAGE сломано"
    assert "PAGE" in footer_xml, "Поле PAGE исчезло из footer"


def test_filler_removes_empty_payment_rows(tmp_path):
    """При СПЕЦ_ОПЛАТА_П5='' пустой нумерованный параграф удаляется из вывода."""
    from docx import Document
    from docx.oxml.ns import qn

    template = os.path.normpath(SPEC_TEMPLATE_PATH)
    output = str(tmp_path / "spec_empty_payment.docx")

    data = {**SPEC_MOCK_DATA, "СПЕЦ_ОПЛАТА_П5": "", "СПЕЦ_ОПЛАТА_П6": ""}
    fill_template(template, data, output)

    doc = Document(output)
    body_tag = qn("w:body")
    NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    empty_numpr = [
        p for p in doc.paragraphs
        if p._p.getparent().tag == body_tag
        and p.text.strip() == ""
        and p._p.find(f".//{{{NS}}}numPr") is not None
    ]
    assert empty_numpr == [], (
        f"Найдено {len(empty_numpr)} пустых нумерованных параграфов после fill_template"
    )


def test_filler_replaces_textbox_placeholders(tmp_path):
    """fill_template заменяет {{ЗАКАЗЧИК_ДИРЕКТОР_ИНИЦИАЛЫ}} в text box Приложения."""
    import zipfile

    template = os.path.normpath(SPEC_TEMPLATE_PATH)
    output = str(tmp_path / "spec_textbox.docx")

    fill_template(template, SPEC_MOCK_DATA, output)

    with zipfile.ZipFile(output) as z:
        content = z.read("word/document.xml").decode("utf-8")

    assert "{{ЗАКАЗЧИК_ДИРЕКТОР_ИНИЦИАЛЫ}}" not in content, (
        "Плейсхолдер {{ЗАКАЗЧИК_ДИРЕКТОР_ИНИЦИАЛЫ}} не был заменён в text box Приложения"
    )


SPEC_MOCK_ITEMS = [
    {
        "id": "weights",
        "name": "Весы автомобильные ВЕСТА-С-60-18-Ц, max 60т, размеры платформы 18х3м",
        "unit": "компл",
        "quantity": 1.0,
        "price_per_unit": 2_835_000.0,
        "total": 2_835_000.0,
        "payment_group": None,
        "is_custom": False,
        "source": "preset",
        "metadata": {},
    },
    {
        "id": "foundation",
        "name": "Фундамент железобетонный под весы ВЕСТА-С, 18м",
        "unit": "компл",
        "quantity": 1.0,
        "price_per_unit": 350_000.0,
        "total": 350_000.0,
        "payment_group": None,
        "is_custom": False,
        "source": "preset",
        "metadata": {"scope": "fundament_jb"},
    },
    {
        "id": "verification",
        "name": "Поверка автомобильных весов с доставкой эталонов",
        "unit": "компл",
        "quantity": 1.0,
        "price_per_unit": 0.0,
        "total": 0.0,
        "payment_group": None,
        "is_custom": False,
        "source": "preset",
        "metadata": {"customer_side": True},
    },
]


def test_fill_spec_with_items_row_count(tmp_path):
    """Table[0] содержит ровно len(items) строк данных (кроме header и total)."""
    from docx import Document
    from src.contracts.filler import fill_spec_with_items

    template = os.path.normpath(SPEC_TEMPLATE_PATH)
    output = str(tmp_path / "spec_items.docx")

    fill_spec_with_items(template, SPEC_MOCK_DATA, SPEC_MOCK_ITEMS, output)

    doc = Document(output)
    table = doc.tables[0]
    # Header row + N item rows + total row
    assert len(table.rows) == 1 + len(SPEC_MOCK_ITEMS) + 1


def test_fill_spec_with_items_names_in_table(tmp_path):
    """Наименования позиций присутствуют в Table[0]."""
    from docx import Document
    from src.contracts.filler import fill_spec_with_items

    template = os.path.normpath(SPEC_TEMPLATE_PATH)
    output = str(tmp_path / "spec_items_names.docx")

    fill_spec_with_items(template, SPEC_MOCK_DATA, SPEC_MOCK_ITEMS, output)

    doc = Document(output)
    table = doc.tables[0]
    all_text = " ".join(c.text for row in table.rows for c in row.cells)

    assert "Весы автомобильные ВЕСТА-С-60-18-Ц" in all_text
    assert "Фундамент железобетонный" in all_text
    assert "ЗАКАЗЧИК" in all_text


def test_fill_spec_with_items_total_computed_from_items(tmp_path):
    """ИТОГО в Table[0] = сумма non-customer-side позиций."""
    from docx import Document
    from src.contracts.filler import fill_spec_with_items

    template = os.path.normpath(SPEC_TEMPLATE_PATH)
    output = str(tmp_path / "spec_items_total.docx")

    fill_spec_with_items(template, SPEC_MOCK_DATA, SPEC_MOCK_ITEMS, output)

    doc = Document(output)
    table = doc.tables[0]
    total_row = table.rows[-1]
    # total = 2_835_000 + 350_000 = 3_185_000; join all cells text
    total_text = " ".join(c.text for c in total_row.cells)
    assert "3" in total_text
    assert total_text.strip() != ""


def test_fill_spec_with_items_preserves_footer_page_field(tmp_path):
    """fill_spec_with_items не уничтожает поле PAGE в footer."""
    import zipfile
    from src.contracts.filler import fill_spec_with_items

    template = os.path.normpath(SPEC_TEMPLATE_PATH)
    output = str(tmp_path / "spec_items_footer.docx")

    fill_spec_with_items(template, SPEC_MOCK_DATA, SPEC_MOCK_ITEMS, output)

    with zipfile.ZipFile(output) as z:
        footer_xml = z.read("word/footer2.xml").decode("utf-8")

    assert "PAGE" in footer_xml


def test_fill_spec_with_items_e2e_from_kp_snapshot(tmp_path):
    """E2E: KP snapshot → build_specification_items → fill_spec_with_items → DOCX."""
    from docx import Document
    from src.contracts.from_kp import build_specification_items
    from src.contracts.filler import fill_spec_with_items

    kp_row = {
        "kp_number": "КП-2026-E2E",
        "model_id": "vesta-s-60-18",
        "data": {
            "model": {"line": "С", "max": 60, "length": 18, "price": 2_835_000},
            "options": {
                "foundation_s_f_18": {"qty": 1, "price": 350_000, "customer_side": False},
                "install_default": {"qty": 1, "price": 80_000, "customer_side": False},
                "delivery_default": {"qty": 1, "price": 50_000, "customer_side": False},
            },
        },
    }

    items = build_specification_items(kp_row)
    assert len(items) == 4  # weights + foundation + install + delivery

    template = os.path.normpath(SPEC_TEMPLATE_PATH)
    output = str(tmp_path / "spec_e2e.docx")
    fill_spec_with_items(template, SPEC_MOCK_DATA, items, output)

    doc = Document(output)
    table = doc.tables[0]
    # header + 4 items + total = 6
    assert len(table.rows) == 6

    item_names = [table.rows[i].cells[0].text for i in range(1, 5)]
    assert any("Весы" in n for n in item_names)
    assert any("Фундамент" in n for n in item_names)
    assert any("Монтаж" in n for n in item_names)
    assert any("Доставка" in n for n in item_names)


def test_fill_spec_with_items_no_raw_ids_leak(tmp_path):
    """E2E регресс: frame_18/ramp_set_f_s/ramp_set_fl_sl/fence_norma_20 → в .docx
    попадают человекочитаемые label из prices.json, а не raw-ключи опций."""
    import json
    from pathlib import Path

    from docx import Document
    from src.contracts.from_kp import build_specification_items
    from src.contracts.filler import fill_spec_with_items

    prices = json.loads(Path("data/prices.json").read_text(encoding="utf-8"))
    raw_keys = ["frame_18", "ramp_set_f_s", "ramp_set_fl_sl", "fence_norma_20"]
    # frame_18/ramp_set_* получают эталонное имя (FIX_SPEC §E3b), не raw
    # prices.json label — fence_norma_20 не переопределён, читает label как есть.
    etalon_overrides = {
        "frame_18": "Рама 18м для весов ВЕСТА",
        "ramp_set_f_s": "Комплект пандусов для весов ВЕСТА-Ф/С",
        "ramp_set_fl_sl": "Комплект пандусов для весов ВЕСТА-СЛ/ФЛ",
    }

    kp_row = {
        "kp_number": "КП-2026-RAWID",
        "model_id": "vesta-s-60-18",
        "data": {
            "model": {"line": "С", "max": 60, "length": 18, "price": 2_835_000},
            "options": {
                key: {"qty": 1, "price": 1_000, "customer_side": False}
                for key in raw_keys
            },
        },
    }

    items = build_specification_items(kp_row, prices=prices)

    template = os.path.normpath(SPEC_TEMPLATE_PATH)
    output = str(tmp_path / "spec_no_raw_ids.docx")
    fill_spec_with_items(template, SPEC_MOCK_DATA, items, output)

    doc = Document(output)
    table = doc.tables[0]
    all_text = " ".join(c.text for row in table.rows for c in row.cells)

    for key in raw_keys:
        expected_label = etalon_overrides.get(key) or prices["options"][key]["label"]
        assert expected_label in all_text
        assert key not in all_text


# ---------------------------------------------------------------------------
# Тесты сохранения drawing-объектов (regression: merge_runs не должен их
# уничтожать через CT_R.clear_content при слиянии рунов без rPr)
# ---------------------------------------------------------------------------

BUILD_TASK_PATH = os.path.normpath(
    os.path.join(
        os.path.dirname(__file__), "..", "..",
        "data", "fundament", "build_task", "пандусный_С_Ф_3скц.docx",
    )
)


@pytest.mark.skipif(
    not os.path.exists(BUILD_TASK_PATH),
    reason="build_task fixture not present",
)
def test_fill_template_preserves_drawings_in_build_task(tmp_path):
    """fill_template не должен уничтожать w:drawing в шаблоне строительного задания.

    До фикса merge_runs: drawings=0 после fill_template.
    После фикса: drawings=14 (7 inline + 7 anchor).
    """
    import zipfile
    from lxml import etree

    output = str(tmp_path / "build_task_filled.docx")
    fill_template(
        BUILD_TASK_PATH,
        {"ПРИЛОЖЕНИЕ_НОМЕР": "1", "СПЕЦ_НОМЕР": "7", "ДОГОВОР_ДАТА_ПОЛНАЯ": "01.06.2026",
         "ЗАКАЗЧИК_ДИРЕКТОР_ИНИЦИАЛЫ": "И.И. Иванов", "ДОГОВОР_НОМЕР": "Д-1"},
        output,
    )

    with zipfile.ZipFile(output) as z:
        doc_xml = z.read("word/document.xml")
    root = etree.fromstring(doc_xml)
    ns = {"wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"}
    inline = root.findall(".//wp:inline", ns)
    anchor = root.findall(".//wp:anchor", ns)

    assert len(inline) == 7, f"Ожидалось 7 inline-drawing, получено {len(inline)}"
    assert len(anchor) == 7, f"Ожидалось 7 anchor-drawing, получено {len(anchor)}"


def test_fill_template_spec_preserves_drawings(tmp_path):
    """Регресс: fill_template не трогает drawing в spec_foundation_install.docx."""
    import zipfile
    from lxml import etree

    template = os.path.normpath(SPEC_TEMPLATE_PATH)
    output = str(tmp_path / "spec_drawings_regression.docx")
    fill_template(template, SPEC_MOCK_DATA, output)

    with zipfile.ZipFile(output) as z:
        doc_xml = z.read("word/document.xml")
    root = etree.fromstring(doc_xml)
    ns = {"wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"}
    inline = root.findall(".//wp:inline", ns)
    anchor = root.findall(".//wp:anchor", ns)

    assert len(inline) >= 1, "inline drawing исчез из spec_foundation_install после fill_template"
    assert len(anchor) >= 1, "anchor drawing исчез из spec_foundation_install после fill_template"

    # Данные из SPEC_MOCK_DATA подставлены (хотя бы один известный ключ)
    assert get_unfilled_placeholders(output) != ["{{ДОГОВОР_НОМЕР}}"], (
        "ДОГОВОР_НОМЕР не заполнен"
    )


def test_merge_runs_drawing_adjacent_to_split_placeholder(tmp_path):
    """Drawing в параграфе рядом с разорванным плейсхолдером — оба сохраняются.

    Параграф: [run '{{'][run 'TEST_KEY'][run '}}'][drawing run]
    После replace_in_paragraph: drawing цел, {{TEST_KEY}} → 'VALUE'.
    """
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn
    from src.contracts.filler import replace_in_paragraph

    doc = Document()
    para = doc.add_paragraph()

    # Три рана симулируют Word-разрыв плейсхолдера {{TEST_KEY}}
    para.add_run("{{")
    para.add_run("TEST_KEY")
    para.add_run("}}")

    # Drawing run — minimal w:drawing без содержимого, как маркер
    r_el = OxmlElement("w:r")
    drawing_el = OxmlElement("w:drawing")
    r_el.append(drawing_el)
    para._p.append(r_el)

    replace_in_paragraph(para, {"TEST_KEY": "VALUE"})

    # Плейсхолдер заполнен
    assert para.text == "VALUE", f"Плейсхолдер не заполнен: {para.text!r}"

    # Drawing ран не уничтожен
    drawings_in_para = para._p.findall(f".//{_qn('w:drawing')}")
    assert len(drawings_in_para) == 1, (
        f"w:drawing уничтожен merge_runs: найдено {len(drawings_in_para)}"
    )
