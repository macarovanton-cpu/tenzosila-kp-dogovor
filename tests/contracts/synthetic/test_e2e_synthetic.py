"""E2E-тесты на 5 синтетических кейсах.

Каждый тест:
1. Прогоняет extract_from_files (реальный AI-вызов через OpenRouter)
2. Заполняет шаблоны договора и спецификации
3. Проверяет чек-лист A1–D3

Тесты намеренно ловят баги — FAIL ожидаем.
"""

import json
import re
from pathlib import Path

import pytest
from docx import Document

from src.contracts.extractor import extract_from_files
from src.contracts.filler import fill_template, get_unfilled_placeholders
from src.contracts.utils import format_date_parts, infer_director_gender, number_to_words

from tests.contracts.synthetic.generate_card import CARDS, generate_all as gen_cards
from tests.contracts.synthetic.generate_kp import CASES, generate_all as gen_kp

FIXTURES_DIR = Path(__file__).parent / "fixtures"
OUTPUT_DIR = Path(__file__).parent / "output"
ROOT = Path(__file__).resolve().parents[3]
CONTRACT_TPL = ROOT / "templates" / "contracts" / "contract.docx"
SPEC_TPL = ROOT / "templates" / "contracts" / "spec_foundation_install.docx"
RESULTS_PATH = OUTPUT_DIR / "test_results.json"


# ---------------------------------------------------------------------------
# Утилиты
# ---------------------------------------------------------------------------


class SoftAssert:
    """Собирает все ошибки, не прерывая тест на первом assert."""

    def __init__(self, case_id: int):
        self.case_id = case_id
        self.results: dict[str, dict] = {}

    def check(self, check_id: str, ok: bool, message: str) -> None:
        self.results[check_id] = {
            "ok": ok,
            "message": "" if ok else f"[Case {self.case_id}] {check_id}: {message}",
        }

    def assert_all(self) -> None:
        fails = [r["message"] for r in self.results.values() if not r["ok"]]
        if fails:
            pytest.fail(f"{len(fails)} проверок не пройдено:\n" + "\n".join(fails))


def _parse_money(s: str) -> int | None:
    """'2 835 000' → 2835000"""
    if not s:
        return None
    digits = re.sub(r"[^\d]", "", s)
    return int(digits) if digits else None


def _extract_docx_text(path: str) -> str:
    """Полный текст DOCX: параграфы + ячейки таблиц."""
    doc = Document(path)
    parts: list[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    for section in doc.sections:
        for p in section.header.paragraphs:
            parts.append(p.text)
        for p in section.footer.paragraphs:
            parts.append(p.text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Фикстуры
# ---------------------------------------------------------------------------


@pytest.fixture(scope="session")
def generated_fixtures():
    """Генерирует синтетические PDF и DOCX один раз."""
    gen_kp()
    gen_cards()


# ---------------------------------------------------------------------------
# Чек-лист A–D
# ---------------------------------------------------------------------------


def _check_A1(sa: SoftAssert, contract_path: str, spec_path: str) -> None:
    for label, path in [("contract", contract_path), ("spec", spec_path)]:
        unfilled = get_unfilled_placeholders(path)
        sa.check("A1", len(unfilled) == 0,
                 f"Незаполненные плейсхолдеры в {label}: {unfilled}")


def _check_A2(sa: SoftAssert, text: str) -> None:
    forbidden = ["110/2026", "ЛСЗ", "Лебедянский", "29.05.2026"]
    for token in forbidden:
        sa.check("A2", token not in text,
                 f"Токен из примера '{token}' найден в выходном документе")


def _check_A3(sa: SoftAssert, text: str) -> None:
    sa.check("A3", "ТПК «Тензосила»" in text,
             "'ТПК «Тензосила»' не найдено")
    sa.check("A3", "Компания Тензосила" not in text,
             "'Компания Тензосила' найдено (артефакт шаблона)")


def _check_B1(sa: SoftAssert, flat: dict) -> None:
    keys = [f"СПЕЦ_П{i}_СУММА" for i in range(1, 6)]
    row_sum = 0
    for k in keys:
        v = _parse_money(flat.get(k, ""))
        if v is not None:
            row_sum += v
    total = _parse_money(flat.get("СПЕЦ_ИТОГО", ""))
    sa.check("B1", total is not None and row_sum == total,
             f"Сумма строк ({row_sum}) != ИТОГО ({total})")


def _check_B2(sa: SoftAssert, flat: dict) -> None:
    total = _parse_money(flat.get("СПЕЦ_ИТОГО", ""))
    words = flat.get("СПЕЦ_ИТОГО_ПРОПИСЬ", "")
    if total is None:
        sa.check("B2", False, "ИТОГО пуст — невозможно проверить пропись")
        return
    expected = number_to_words(total)
    sa.check("B2", expected == words,
             f"number_to_words({total}) = '{expected}', "
             f"а ИТОГО_ПРОПИСЬ = '{words}'")


def _check_B3(sa: SoftAssert, flat: dict) -> None:
    total = _parse_money(flat.get("СПЕЦ_ИТОГО", ""))
    if total is None:
        sa.check("B3", False, "ИТОГО пуст")
        return
    payment_sum = 0
    found_any = False
    for i in range(1, 7):
        text = flat.get(f"СПЕЦ_ОПЛАТА_П{i}", "")
        if not text:
            continue
        amounts = re.findall(r"(\d[\d\s]+\d)\s*(?:\(|рубл|руб)", text)
        for amt_str in amounts:
            v = _parse_money(amt_str)
            if v and v < total:
                payment_sum += v
                found_any = True
    if found_any:
        sa.check("B3", payment_sum == total,
                 f"Сумма этапов оплаты ({payment_sum}) != ИТОГО ({total})")
    else:
        sa.check("B3", False,
                 "Не удалось извлечь суммы из условий оплаты")


def _check_C1(sa: SoftAssert, flat: dict, kp_case) -> None:
    model_short = flat.get("СПЕЦ_МОДЕЛЬ_КРАТКОЕ", "")
    sa.check("C1",
             kp_case.model in model_short or model_short in kp_case.model,
             f"Модель: ожидали '{kp_case.model}', получили '{model_short}'")


def _check_C2(sa: SoftAssert, flat: dict, kp_case) -> None:
    p1 = flat.get("СПЕЦ_П1_НАИМЕНОВАНИЕ", "")
    load = flat.get("СПЕЦ_МАКС_НАГРУЗКА", "")
    combined = f"{p1} {load}"
    ok = (kp_case.sensor_model in combined
          or str(kp_case.sensor_count) in combined)
    sa.check("C2", ok,
             f"Датчики {kp_case.sensor_model} {kp_case.sensor_count} шт. "
             f"не найдены в П1_НАИМЕНОВАНИЕ='{p1}' / "
             f"МАКС_НАГРУЗКА='{load}'")


def _check_C3(sa: SoftAssert, flat: dict, kp_case) -> None:
    extracted = _parse_money(flat.get("СПЕЦ_ИТОГО", ""))
    sa.check("C3", extracted == kp_case.total,
             f"ИТОГО: ожидали {kp_case.total}, получили {extracted}")


def _check_C4(sa: SoftAssert, flat: dict, kp_case) -> None:
    payments = " ".join(
        flat.get(f"СПЕЦ_ОПЛАТА_П{i}", "") for i in range(1, 7)
    )
    for cond in kp_case.payment_conditions:
        if "банковских" in cond:
            sa.check("C4", "банковских" in payments,
                     "'банковских дней' утеряно в условиях оплаты")
            return
        if "рабочих" in cond:
            sa.check("C4", "рабочих" in payments,
                     "'рабочих дней' утеряно в условиях оплаты")
            return
    sa.check("C4", True, "")


def _check_C5(sa: SoftAssert, flat: dict, kp_case) -> None:
    extracted_sums = []
    for i in range(1, 6):
        v = _parse_money(flat.get(f"СПЕЦ_П{i}_СУММА", ""))
        if v:
            extracted_sums.append(v)
    for row in kp_case.spec_rows:
        sa.check("C5", row.price in extracted_sums,
                 f"Строка '{row.name}' ({row.price}) "
                 f"не найдена в суммах спеки {extracted_sums}")


def _check_D1(sa: SoftAssert, flat: dict, card) -> None:
    extracted = flat.get("ЗАКАЗЧИК_КРАТКОЕ_НАИМЕНОВАНИЕ", "")
    sa.check("D1", card.short_name == extracted,
             f"Краткое: ожидали '{card.short_name}', получили '{extracted}'")


def _check_D2(sa: SoftAssert, flat: dict, card) -> None:
    fio = flat.get("ЗАКАЗЧИК_ДИРЕКТОР_ФИО", "")
    title = flat.get("ЗАКАЗЧИК_ДИРЕКТОР_ДОЛЖНОСТЬ", "")
    sa.check("D2", card.director_fio == fio,
             f"ФИО: ожидали '{card.director_fio}', получили '{fio}'")
    sa.check("D2", card.director_title == title,
             f"Должность: ожидали '{card.director_title}', получили '{title}'")


def _check_D3(sa: SoftAssert, flat: dict, card) -> None:
    checks = [
        ("ЗАКАЗЧИК_ИНН", card.inn, "ИНН"),
        ("ЗАКАЗЧИК_КПП", card.kpp, "КПП"),
        ("ЗАКАЗЧИК_ОГРН", card.ogrn, "ОГРН"),
        ("ЗАКАЗЧИК_РС", card.bank_account, "Р/с"),
        ("ЗАКАЗЧИК_БИК", card.bik, "БИК"),
    ]
    for key, expected, label in checks:
        got = flat.get(key, "")
        sa.check("D3", expected == got,
                 f"{label}: ожидали '{expected}', получили '{got}'")


# ---------------------------------------------------------------------------
# Тесты
# ---------------------------------------------------------------------------


CASE_PARAMS = [
    pytest.param(kp, card, id=f"case_{kp.case_id}")
    for kp, card in zip(CASES, CARDS)
]


@pytest.mark.usefixtures("generated_fixtures")
@pytest.mark.parametrize("kp_case,card_data", CASE_PARAMS)
def test_contract_e2e(kp_case, card_data):
    """Полный E2E-тест для одного синтетического кейса."""
    sa = SoftAssert(kp_case.case_id)

    kp_path = str(FIXTURES_DIR / f"case_{kp_case.case_id}_kp.pdf")
    card_path = str(FIXTURES_DIR / f"case_{kp_case.case_id}_card.docx")

    # 1. AI-извлечение
    data = extract_from_files(kp_path, card_path)

    # Сохраняем сырые данные для отладки
    raw_path = OUTPUT_DIR / f"case_{kp_case.case_id}_raw.json"
    raw_path.parent.mkdir(parents=True, exist_ok=True)
    with open(raw_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    # 2. Плоский словарь + ручные поля
    flat: dict[str, str] = {}
    flat.update(data.get("requisites", {}))
    flat.update(data.get("specification", {}))
    date_parts = format_date_parts("15.05.2026")
    flat.update(date_parts)
    flat["ДОГОВОР_НОМЕР"] = f"Т-{kp_case.case_id:03d}/2026"
    flat["СПЕЦ_АДРЕС_ОБЪЕКТА"] = f"г. Тестовый, промзона №{kp_case.case_id}"
    flat["СПЕЦ_НОМЕР"] = "1"
    _pol = flat.get("ЗАКАЗЧИК_ДИРЕКТОР_ПОЛ", "")
    if not _pol:
        _fio = flat.get("ЗАКАЗЧИК_ДИРЕКТОР_ФИО", "")
        _pol = infer_director_gender(_fio) if _fio else "male"
    flat["ДИРЕКТОР_ПРИЧАСТИЕ"] = "действующей" if _pol == "female" else "действующего"

    # 3. Заполнение шаблонов
    contract_out = str(OUTPUT_DIR / f"case_{kp_case.case_id}_contract.docx")
    spec_out = str(OUTPUT_DIR / f"case_{kp_case.case_id}_spec.docx")
    fill_template(str(CONTRACT_TPL), flat, contract_out)
    fill_template(str(SPEC_TPL), flat, spec_out)

    # 4. Извлечение текста для проверок
    contract_text = _extract_docx_text(contract_out)
    spec_text = _extract_docx_text(spec_out)
    full_text = contract_text + "\n" + spec_text

    # 5. Чек-лист
    _check_A1(sa, contract_out, spec_out)
    _check_A2(sa, full_text)
    _check_A3(sa, full_text)

    _check_B1(sa, flat)
    _check_B2(sa, flat)
    _check_B3(sa, flat)

    _check_C1(sa, flat, kp_case)
    _check_C2(sa, flat, kp_case)
    _check_C3(sa, flat, kp_case)
    _check_C4(sa, flat, kp_case)
    _check_C5(sa, flat, kp_case)

    _check_D1(sa, flat, card_data)
    _check_D2(sa, flat, card_data)
    _check_D3(sa, flat, card_data)

    # 6. Сохраняем результаты чек-листа
    results_all = {}
    if RESULTS_PATH.exists():
        with open(RESULTS_PATH, "r", encoding="utf-8") as f:
            results_all = json.load(f)
    results_all[str(kp_case.case_id)] = {
        k: v["ok"] for k, v in sa.results.items()
    }
    with open(RESULTS_PATH, "w", encoding="utf-8") as f:
        json.dump(results_all, f, ensure_ascii=False, indent=2)

    sa.assert_all()
