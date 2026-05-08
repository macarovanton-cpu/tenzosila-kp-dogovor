"""Генератор синтетических карточек контрагентов (DOCX)."""

from dataclasses import dataclass
from pathlib import Path

from docx import Document

FIXTURES_DIR = Path(__file__).parent / "fixtures"


@dataclass
class CardData:
    """Полные данные карточки контрагента."""

    case_id: int
    short_name: str
    full_name: str
    inn: str
    kpp: str
    ogrn: str
    okpo: str
    legal_address: str
    postal_address: str
    bank_account: str
    bank_name: str
    corr_account: str
    bik: str
    director_title: str
    director_fio: str
    director_basis: str
    phone: str
    email: str


CARDS: list[CardData] = [
    CardData(
        case_id=1,
        short_name='ООО «АгроТрейд»',
        full_name='Общество с ограниченной ответственностью «АгроТрейд»',
        inn='2312345678', kpp='231201001', ogrn='1152312000123', okpo='12345678',
        legal_address='350000, г. Краснодар, ул. Красная, д. 45, офис 201',
        postal_address='350000, г. Краснодар, ул. Красная, д. 45, офис 201',
        bank_account='40702810400260012345',
        bank_name='ПАО Сбербанк г. Краснодар',
        corr_account='30101810600000000602',
        bik='040349602',
        director_title='Генеральный директор',
        director_fio='Петров Андрей Викторович',
        director_basis='Устава',
        phone='+7 (861) 234-56-78',
        email='info@agrotrade-krd.ru',
    ),
    CardData(
        case_id=2,
        short_name='ООО «МеталлИнвест»',
        full_name='Общество с ограниченной ответственностью «МеталлИнвест»',
        inn='7451234567', kpp='745101001', ogrn='1167451000234', okpo='23456789',
        legal_address='454000, г. Челябинск, пр. Ленина, д. 87, офис 412',
        postal_address='454000, г. Челябинск, пр. Ленина, д. 87, офис 412',
        bank_account='40702810200000054321',
        bank_name='Филиал «Уральский» ПАО «Промсвязьбанк»',
        corr_account='30101810700000000877',
        bik='047501877',
        director_title='Генеральный директор',
        director_fio='Сидоров Константин Михайлович',
        director_basis='Устава',
        phone='+7 (351) 345-67-89',
        email='office@metallinvest74.ru',
    ),
    CardData(
        case_id=3,
        short_name='АО «УралГорнодобыча»',
        full_name='Акционерное общество «УралГорнодобыча»',
        inn='6678901234', kpp='667801001', ogrn='1206600012345', okpo='34567890',
        legal_address='620000, г. Екатеринбург, ул. Малышева, д. 101, стр. 3',
        postal_address='620000, г. Екатеринбург, а/я 567',
        bank_account='40702810900000067890',
        bank_name='Уральский филиал АО «Альфа-Банк»',
        corr_account='30101810100000000964',
        bik='046577964',
        director_title='Генеральный директор',
        director_fio='Кузнецов Дмитрий Андреевич',
        director_basis='Устава',
        phone='+7 (343) 456-78-90',
        email='info@uralgornodobycha.ru',
    ),
    CardData(
        case_id=4,
        short_name='ЗАО «СтройКомплекс»',
        full_name='Закрытое акционерное общество «СтройКомплекс»',
        inn='6950012345', kpp='695001001', ogrn='1036900012345', okpo='45678901',
        legal_address='170000, г. Тверь, ул. Советская, д. 23',
        postal_address='170000, г. Тверь, ул. Советская, д. 23',
        bank_account='40702810100000098765',
        bank_name='ПАО «ВТБ» филиал в г. Тверь',
        corr_account='30101810145250000411',
        bik='042809411',
        director_title='Генеральный директор',
        director_fio='Морозов Виктор Сергеевич',
        director_basis='Устава',
        phone='+7 (4822) 56-78-90',
        email='info@stroykompleks-tver.ru',
    ),
    CardData(
        case_id=5,
        short_name='ООО «ЗерноЛогистика»',
        full_name='Общество с ограниченной ответственностью «ЗерноЛогистика»',
        inn='6312345678', kpp='631201001', ogrn='1186312000567', okpo='56789012',
        legal_address='443000, г. Самара, ул. Молодогвардейская, д. 33, офис 15',
        postal_address='443000, г. Самара, ул. Молодогвардейская, д. 33, офис 15',
        bank_account='40702810300000011223',
        bank_name='ПАО Сбербанк Самарское отделение',
        corr_account='30101810200000000607',
        bik='043601607',
        director_title='Генеральный директор',
        director_fio='Васильев Игорь Николаевич',
        director_basis='Устава',
        phone='+7 (846) 678-90-12',
        email='logistics@zernolog.ru',
    ),
]


def generate_card_docx(card: CardData, output_path: Path) -> None:
    """Генерирует DOCX-карточку контрагента с таблицей реквизитов."""
    doc = Document()
    doc.add_heading('Карточка контрагента', level=1)

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    fields = [
        ('Краткое наименование', card.short_name),
        ('Полное наименование', card.full_name),
        ('ИНН', card.inn),
        ('КПП', card.kpp),
        ('ОГРН', card.ogrn),
        ('ОКПО', card.okpo),
        ('Юридический адрес', card.legal_address),
        ('Почтовый адрес', card.postal_address),
        ('Расчетный счет', card.bank_account),
        ('Банк', card.bank_name),
        ('Корреспондентский счет', card.corr_account),
        ('БИК', card.bik),
        ('Должность руководителя', card.director_title),
        ('ФИО руководителя', card.director_fio),
        ('Основание полномочий', card.director_basis),
        ('Телефон', card.phone),
        ('Email', card.email),
    ]

    for label, value in fields:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value

    doc.save(str(output_path))


def generate_all() -> None:
    """Генерирует все 5 карточек в fixtures/."""
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)
    for card in CARDS:
        path = FIXTURES_DIR / f"case_{card.case_id}_card.docx"
        generate_card_docx(card, path)


if __name__ == "__main__":
    generate_all()
    print(f"Сгенерировано {len(CARDS)} карточек в {FIXTURES_DIR}")
