"""Тесты склейки спецификации с внешними приложениями."""
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from src.contracts.compose import compose_spec_with_attachments, compose_supply
from src.contracts.filler import get_unfilled_placeholders
from src.contracts.requisites_parser import parse_requisites
from src.contracts.spec_v2_filler import _make_clause_para
from src.contracts.supply_filler import _buyer_context


def _make_spec(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Начало спецификации")
    doc.add_paragraph("Конец спецификации")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Поставщик"
    table.rows[0].cells[1].text = "Заказчик"
    doc.save(path)


def _make_appendix(path: Path, title: str, page_break_before: bool = False) -> None:
    doc = Document()
    p = doc.add_paragraph(f"Приложение №{{{{ПРИЛОЖЕНИЕ_НОМЕР}}}}. {title}")
    p.paragraph_format.page_break_before = page_break_before
    doc.add_paragraph("К Спецификации №{{СПЕЦ_НОМЕР}}")
    doc.add_paragraph("К Договору №{{ДОГОВОР_НОМЕР}} от {{ДОГОВОР_ДАТА_ПОЛНАЯ}}")
    doc.add_paragraph("Утверждаю {{ЗАКАЗЧИК_ДИРЕКТОР_ИНИЦИАЛЫ}}")
    doc.save(path)


def _all_text(doc: Document) -> str:
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return "\n".join(texts)


def _paragraph_text(p_el) -> str:
    return "".join(t.text or "" for t in p_el.findall(".//" + qn("w:t")))


def _has_page_break(p_el) -> bool:
    if p_el.find(".//" + qn("w:pageBreakBefore")) is not None:
        return True
    for br in p_el.findall(".//" + qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def _body_paragraphs(doc: Document) -> list:
    return [el for el in doc.element.body.iterchildren() if el.tag == qn("w:p")]


def _data() -> dict:
    return {
        "СПЕЦ_НОМЕР": "7",
        "ДОГОВОР_НОМЕР": "Д-42",
        "ДОГОВОР_ДАТА_ПОЛНАЯ": "10 июня 2026 г.",
        "ЗАКАЗЧИК_ДИРЕКТОР_ИНИЦИАЛЫ": "И.И. Иванов",
    }


def test_compose_both(tmp_path: Path) -> None:
    spec = tmp_path / "spec.docx"
    build_task = tmp_path / "build_task.docx"
    control_sheet = tmp_path / "control_sheet.docx"
    _make_spec(spec)
    _make_appendix(build_task, "Строительное задание")
    _make_appendix(control_sheet, "Контрольный лист")

    before_count = len(Document(spec).paragraphs)
    compose_spec_with_attachments(
        spec,
        {
            "build_task_path": str(build_task),
            "build_task_source": "manual",
            "include_control_sheet": True,
            "control_sheet_path": str(control_sheet),
        },
        _data(),
    )

    doc = Document(spec)
    text = _all_text(doc)
    assert len(doc.paragraphs) > before_count
    assert "Строительное задание" in text
    assert "Контрольный лист" in text
    assert "Поставщик" in text
    assert "Заказчик" in text


def test_compose_build_task_only(tmp_path: Path) -> None:
    spec = tmp_path / "spec.docx"
    build_task = tmp_path / "build_task.docx"
    control_sheet = tmp_path / "control_sheet.docx"
    _make_spec(spec)
    _make_appendix(build_task, "Строительное задание")
    _make_appendix(control_sheet, "Контрольный лист")

    compose_spec_with_attachments(
        spec,
        {
            "build_task_path": str(build_task),
            "build_task_source": "manual",
            "include_control_sheet": False,
            "control_sheet_path": str(control_sheet),
        },
        _data(),
    )

    text = _all_text(Document(spec))
    assert "Строительное задание" in text
    assert "Контрольный лист" not in text


def test_compose_none_leaves_file_unchanged(tmp_path: Path) -> None:
    spec = tmp_path / "spec.docx"
    _make_spec(spec)
    before = spec.read_bytes()

    compose_spec_with_attachments(
        spec,
        {
            "build_task_path": "",
            "build_task_source": "none",
            "include_control_sheet": False,
            "control_sheet_path": "",
        },
        _data(),
    )

    assert spec.read_bytes() == before


def test_placeholders_filled(tmp_path: Path) -> None:
    spec = tmp_path / "spec.docx"
    build_task = tmp_path / "build_task.docx"
    _make_spec(spec)
    _make_appendix(build_task, "Строительное задание")

    compose_spec_with_attachments(
        spec,
        {
            "build_task_path": str(build_task),
            "build_task_source": "manual",
            "include_control_sheet": False,
            "control_sheet_path": "",
        },
        _data(),
    )

    text = _all_text(Document(spec))
    assert "{{" not in text
    assert get_unfilled_placeholders(str(spec)) == []
    assert "К Спецификации №7" in text
    assert "Д-42" in text


def test_numbering_is_fixed_by_attachment_type(tmp_path: Path) -> None:
    spec = tmp_path / "spec.docx"
    build_task = tmp_path / "build_task.docx"
    control_sheet = tmp_path / "control_sheet.docx"
    _make_spec(spec)
    _make_appendix(build_task, "Строительное задание")
    _make_appendix(control_sheet, "Контрольный лист")

    compose_spec_with_attachments(
        spec,
        {
            "build_task_path": str(build_task),
            "build_task_source": "manual",
            "include_control_sheet": True,
            "control_sheet_path": str(control_sheet),
        },
        _data(),
    )

    text = _all_text(Document(spec))
    assert "Приложение №1. Строительное задание" in text
    assert "Приложение №2. Контрольный лист" in text


# ---------------------------------------------------------------------------
# compose_supply — guard пустого блока оплаты (task #1)
# ---------------------------------------------------------------------------

def _supply_ctx(payment_lines: list[str]) -> dict:
    """Минимальный контекст для compose_supply (рендерит реальные шаблоны)."""
    ctx: dict = {
        "ДОГОВОР_НОМЕР": "Д-1",
        "ДОГОВОР_ДАТА": "1 июня 2026 года",
        "ДОГОВОР_ГОРОД": "г. Воронеж",
        "ТОВАР_НАИМЕНОВАНИЕ": "ВЕСТА-С-60, в количестве 1шт.",
        "СУММА_ЦИФРАМИ": "1 000 000",
        "СУММА_ПРОПИСЬЮ": "Один миллион",
        "СРОК_ПРОИЗВОДСТВА_ДН": "20 (двадцати)",
        "СРОК_ДОСТАВКИ_ДН": "4 (четырёх)",
        "АДРЕС_ПОСТАВКИ": "г. Воронеж",
        "СРОК_ДЕЙСТВИЯ_ДО": "31.12.2026 г.",
        "ТЕКУЩИЙ_ГОД": "2026",
        "spec_rows": [{"name": "ВЕСТА-С-60", "sum": "1 000 000"}],
        "kit_rows": [],
        "PAYMENT_SECTION": "{{PAYMENT_SECTION}}",
        "_payment_lines": payment_lines,
    }
    for key in (
        "ТТХ_MAX", "ТТХ_ОСЬ", "ТТХ_РАССТОЯНИЕ_ТЕРМИНАЛ", "ТТХ_ДИСКРЕТНОСТЬ",
        "ТТХ_ГАБАРИТЫ", "ТТХ_ТЕМПЕРАТУРА", "ТТХ_СВЯЗЬ", "ТТХ_ПИТАНИЕ",
        "ТТХ_МОЩНОСТЬ", "ТТХ_ГОСТ_СТРОКА",
    ):
        ctx[key] = ""
    return ctx


def test_compose_supply_empty_payment_removes_marker(tmp_path: Path) -> None:
    """Пустой _payment_lines → маркер {{PAYMENT_SECTION}} удалён, не остаётся в документе."""
    out = tmp_path / "supply.docx"
    compose_supply(_supply_ctx([]), out)

    text = _all_text(Document(out))
    assert "PAYMENT_SECTION" not in text


def test_compose_supply_with_payment_inserts_lines(tmp_path: Path) -> None:
    """Непустой _payment_lines → строки оплаты вставлены, маркера нет."""
    out = tmp_path / "supply.docx"
    compose_supply(_supply_ctx(["4.2.1. Тестовая строка оплаты."]), out)

    text = _all_text(Document(out))
    assert "Тестовая строка оплаты" in text
    assert "PAYMENT_SECTION" not in text


def test_compose_supply_payment_lines_inherit_marker_format(tmp_path: Path) -> None:
    """Строки оплаты наследуют формат маркера: шрифт (Arial 11) и pPr (интервал 1.5, по ширине)."""
    out = tmp_path / "supply.docx"
    compose_supply(_supply_ctx(["4.2.1. Тестовая строка оплаты."]), out)

    doc = Document(out)
    target = next(p for p in doc.paragraphs if "Тестовая строка оплаты" in p.text)
    assert target.runs, "у параграфа оплаты должен быть хотя бы один ран"
    rpr = target.runs[0]._element.find(qn("w:rPr"))
    assert rpr is not None
    r_fonts = rpr.find(qn("w:rFonts"))
    assert r_fonts is not None
    assert r_fonts.get(qn("w:ascii")) == "Arial"
    sz = rpr.find(qn("w:sz"))
    assert sz is not None
    assert sz.get(qn("w:val")) == "22"

    p_el = target._element
    p_pr = p_el.find(qn("w:pPr"))
    assert p_pr is not None
    spacing = p_pr.find(qn("w:spacing"))
    assert spacing is not None
    assert spacing.get(qn("w:line")) == "360"
    assert spacing.get(qn("w:lineRule")) == "auto"
    jc = p_pr.find(qn("w:jc"))
    assert jc is not None
    assert jc.get(qn("w:val")) == "both"


def test_compose_supply_buyer_phone_from_pasted_requisites_shown(tmp_path: Path) -> None:
    """Реквизиты с «Телефон: +7 (NNN) ...» → строка Тел./факс покупателя видна в договоре.

    Полный путь бага: paste → parse_requisites → _buyer_context → рендер.
    """
    requisites = (
        'ООО "Завод деталей"\n'
        "ИНН 3665123456 КПП 366501001\n"
        "ОГРН 1163668123456\n"
        "Юр.адрес: 394000, г. Воронеж, ул. Ленина, 1\n"
        "Телефон: +7 (473) 214-58-62\n"
    )
    ctx = _supply_ctx([])
    ctx.update(_buyer_context(parse_requisites(requisites)))

    out = tmp_path / "supply.docx"
    compose_supply(ctx, out)

    text = _all_text(Document(out))
    assert "Тел./факс: +7 (473) 214-58-62" in text
    assert "{{" not in text
    assert "{%" not in text


def test_compose_supply_buyer_phone_absent_line_hidden(tmp_path: Path) -> None:
    """Реквизиты без телефона → строка Тел./факс покупателя скрыта (только строка поставщика)."""
    requisites = (
        'ООО "Завод деталей"\n'
        "ИНН 3665123456 КПП 366501001\n"
        "ОГРН 1163668123456\n"
        "Юр.адрес: 394000, г. Воронеж, ул. Ленина, 1\n"
    )
    ctx = _supply_ctx([])
    ctx.update(_buyer_context(parse_requisites(requisites)))

    out = tmp_path / "supply.docx"
    compose_supply(ctx, out)

    text = _all_text(Document(out))
    # Только захардкоженный телефон поставщика (Тензосилы), строки покупателя нет.
    assert text.count("Тел./факс:") == 1


def test_make_clause_para_rpr_and_bold_schema_order() -> None:
    """rPr копируется на каждый ран (включая br), <w:b/> вставляется по схеме OOXML."""
    from docx.oxml import OxmlElement

    marker_rpr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Arial")
    marker_rpr.append(r_fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")
    marker_rpr.append(sz)

    p_el = _make_clause_para("Строка 1\nСтрока 2", bold=True, rpr=marker_rpr)
    runs = p_el.findall(qn("w:r"))
    assert len(runs) == 3  # line1, br, line2

    for run in runs:
        run_rpr = run.find(qn("w:rPr"))
        assert run_rpr is not None
        tags = [child.tag.split("}")[-1] for child in run_rpr]
        assert tags == ["rFonts", "b", "sz"], tags
        assert run_rpr.findall(qn("w:b")) and len(run_rpr.findall(qn("w:b"))) == 1
        assert run_rpr.find(qn("w:rFonts")).get(qn("w:ascii")) == "Arial"

    br_run = runs[1]
    assert br_run.find(qn("w:br")) is not None


def test_make_clause_para_ppr_is_source_of_truth_over_justify() -> None:
    """ppr копируется целиком на параграф; justify игнорируется, если ppr передан."""
    from docx.oxml import OxmlElement

    marker_ppr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "360")
    spacing.set(qn("w:lineRule"), "auto")
    marker_ppr.append(spacing)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "both")
    marker_ppr.append(jc)

    p_el = _make_clause_para("Строка оплаты", justify=False, ppr=marker_ppr)
    p_pr = p_el.find(qn("w:pPr"))
    assert p_pr is not None
    assert p_pr is not marker_ppr  # копия, не та же ссылка
    assert p_pr.find(qn("w:spacing")).get(qn("w:line")) == "360"
    assert p_pr.find(qn("w:jc")).get(qn("w:val")) == "both"

    # justify=True не должен задваивать/переопределять, ppr остаётся источником истины
    p_el2 = _make_clause_para("Строка оплаты", justify=True, ppr=marker_ppr)
    p_pr2 = p_el2.find(qn("w:pPr"))
    assert len(p_pr2.findall(qn("w:jc"))) == 1
    assert p_pr2.find(qn("w:jc")).get(qn("w:val")) == "both"

    # без ppr старая логика justify не меняется
    p_el3 = _make_clause_para("Строка оплаты", justify=True)
    p_pr3 = p_el3.find(qn("w:pPr"))
    assert p_pr3 is not None
    assert p_pr3.find(qn("w:jc")).get(qn("w:val")) == "both"
    assert p_pr3.find(qn("w:spacing")) is None


def test_each_appendix_starts_with_single_page_break(tmp_path: Path) -> None:
    spec = tmp_path / "spec.docx"
    build_task = tmp_path / "build_task.docx"
    control_sheet = tmp_path / "control_sheet.docx"
    _make_spec(spec)
    _make_appendix(build_task, "Строительное задание", page_break_before=True)
    _make_appendix(control_sheet, "Контрольный лист", page_break_before=True)

    compose_spec_with_attachments(
        spec,
        {
            "build_task_path": str(build_task),
            "build_task_source": "manual",
            "include_control_sheet": True,
            "control_sheet_path": str(control_sheet),
        },
        _data(),
    )

    doc = Document(spec)
    paragraphs = _body_paragraphs(doc)
    spec_end_idx = next(
        i for i, p_el in enumerate(paragraphs)
        if _paragraph_text(p_el) == "Конец спецификации"
    )
    appendix_idx = next(
        i for i, p_el in enumerate(paragraphs)
        if "Приложение №1. Строительное задание" in _paragraph_text(p_el)
    )
    control_idx = next(
        i for i, p_el in enumerate(paragraphs)
        if "Приложение №2. Контрольный лист" in _paragraph_text(p_el)
    )
    first_appendix_end_idx = next(
        i for i, p_el in enumerate(paragraphs[appendix_idx + 1:], appendix_idx + 1)
        if "Утверждаю" in _paragraph_text(p_el)
    )
    before_build_task = paragraphs[spec_end_idx + 1:appendix_idx]
    before_control_sheet = paragraphs[first_appendix_end_idx + 1:control_idx]

    assert all(_paragraph_text(p_el).strip() for p_el in before_build_task)
    assert not any(_has_page_break(p_el) for p_el in before_build_task)
    assert all(_paragraph_text(p_el).strip() for p_el in before_control_sheet)
    assert not any(_has_page_break(p_el) for p_el in before_control_sheet)
    assert _has_page_break(paragraphs[appendix_idx])
    assert _has_page_break(paragraphs[control_idx])
