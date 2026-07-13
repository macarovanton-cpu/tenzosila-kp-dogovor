"""Тесты на корректность плейсхолдеров и вёрстки DOCX-шаблонов договора."""
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

CONTRACTS = Path("templates/contracts")


def _header_text(docx_path: Path) -> str:
    doc = Document(docx_path)
    parts = []
    for section in doc.sections:
        for p in section.header.paragraphs:
            parts.append(p.text)
    return "\n".join(parts)


def _all_text(docx_path: Path) -> str:
    doc = Document(docx_path)
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    return "\n".join(parts)


def test_contract_header_has_placeholder():
    text = _header_text(CONTRACTS / "contract.docx")
    assert "{{ДОГОВОР_НОМЕР}}" in text, f"Hardcoded header: {text!r}"
    assert "{{ДОГОВОР_ДАТА_ПОЛНАЯ}}" in text


def test_spec_header_has_placeholder():
    text = _header_text(CONTRACTS / "spec_foundation_install.docx")
    assert "{{ДОГОВОР_НОМЕР}}" in text, f"Hardcoded header: {text!r}"


def test_spec_no_kompaniya_tenzosila():
    text = _all_text(CONTRACTS / "spec_foundation_install.docx")
    assert "Компания Тензосила" not in text, "Нашли захардкоженное название"


def test_spec_has_tpk_tenzosila():
    text = _all_text(CONTRACTS / "spec_foundation_install.docx")
    assert "ТПК" in text, "ООО «ТПК «Тензосила»» должно быть в тексте"


# --- вёрстка ---

def _body_paras(doc):
    body_tag = qn("w:body")
    return [p for p in doc.paragraphs if p._p.getparent().tag == body_tag]


def _find_body_para(doc, contains: str):
    for p in _body_paras(doc):
        if contains in p.text:
            return p
    return None


def _has_para_prop(para, prop_name: str) -> bool:
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        return False
    el = pPr.find(qn(prop_name))
    if el is None:
        return False
    return el.get(qn("w:val"), "1") not in ("0", "false", "off")


def _row_has_cant_split(row) -> bool:
    trPr = row._tr.find(qn("w:trPr"))
    if trPr is None:
        return False
    cs = trPr.find(qn("w:cantSplit"))
    if cs is None:
        return False
    return cs.get(qn("w:val"), "1") not in ("0", "false", "off")


def test_spec_tth_heading_has_page_break_before():
    doc = Document(CONTRACTS / "spec_foundation_install.docx")
    p = _find_body_para(doc, "Технические характеристики")
    assert p is not None, "Параграф 'Технические характеристики' не найден"
    assert _has_para_prop(p, "w:pageBreakBefore"), (
        "tth_heading должен иметь pageBreakBefore"
    )


def test_spec_приложение_has_page_break_before():
    doc = Document(CONTRACTS / "spec_foundation_install.docx")
    p = _find_body_para(doc, "Приложение №{{СПЕЦ_НОМЕР}}")
    assert p is not None, "Параграф 'Приложение №' не найден"
    assert _has_para_prop(p, "w:pageBreakBefore"), "Приложение должно иметь pageBreakBefore"


def test_spec_footer_has_page_field():
    """Шаблон spec_foundation_install.docx содержит поле PAGE в footer2."""
    import zipfile
    with zipfile.ZipFile(CONTRACTS / "spec_foundation_install.docx") as z:
        footer_xml = z.read("word/footer2.xml").decode("utf-8")
    assert "instrText" in footer_xml, "instrText отсутствует в footer2"
    assert "PAGE" in footer_xml, "Поле PAGE отсутствует в footer2"


def test_spec_kit_heading_has_keep_with_next():
    """Параграф kit_heading (Комплект поставки) имеет keepWithNext."""
    doc = Document(CONTRACTS / "spec_foundation_install.docx")
    p = _find_body_para(doc, "Комплект поставки")
    assert p is not None, "Параграф 'Комплект поставки' не найден"
    assert _has_para_prop(p, "w:keepNext"), "kit_heading должен иметь keepNext"


def test_spec_v2_tth_and_kit_headings_are_placeholders_without_numbering():
    """spec_v2.docx хранит номера ТТХ/комплекта через плейсхолдеры, без numPr."""
    doc = Document(CONTRACTS / "spec_v2.docx")
    tth = _find_body_para(doc, "{{ТТХ_НОМЕР}}. Технические характеристики")
    kit = _find_body_para(doc, "{{КОМПЛЕКТ_НОМЕР}}. Комплект поставки")

    assert tth is not None, "tth_heading с {{ТТХ_НОМЕР}} не найден"
    assert kit is not None, "kit_heading с {{КОМПЛЕКТ_НОМЕР}} не найден"
    assert not _has_para_prop(tth, "w:numPr"), "tth_heading не должен иметь numPr"
    assert not _has_para_prop(kit, "w:numPr"), "kit_heading не должен иметь numPr"


def test_spec_v2_payment_and_terms_headings_have_no_numpr():
    """A8/B9: заголовки «2. Порядок оплаты» и «3. Срок поставки» несут номер
    только литеральным текстом, без Word-numPr.

    Двойная нумерация «2. 2.» рисуется Word при рендере — ниже уровня golden-
    дампа. Проверка структурная (на шаблоне), не на дампе (см. B9).
    """
    doc = Document(CONTRACTS / "spec_v2.docx")
    payment = _find_body_para(doc, "Порядок оплаты")
    terms = _find_body_para(doc, "Срок поставки")

    assert payment is not None, "заголовок 'Порядок оплаты' не найден"
    assert terms is not None, "заголовок 'Срок поставки' не найден"
    assert not _has_para_prop(payment, "w:numPr"), (
        "заголовок оплаты не должен иметь numPr (иначе Word печатает «2. 2.»)"
    )
    assert not _has_para_prop(terms, "w:numPr"), (
        "заголовок срока не должен иметь numPr (иначе Word печатает «3. 3.»)"
    )


def test_spec_table2_rows_have_cant_split():
    """Все строки TABLE[2] (Комплект поставки) имеют cantSplit."""
    doc = Document(CONTRACTS / "spec_foundation_install.docx")
    assert len(doc.tables) > 2, "TABLE[2] не найдена в шаблоне"
    for i, row in enumerate(doc.tables[2].rows):
        assert _row_has_cant_split(row), (
            f"TABLE[2] row[{i}] должна иметь cantSplit"
        )


def test_spec_no_kratkoe_placeholder():
    """Плейсхолдер ЗАКАЗЧИК_ДИРЕКТОР_ФИО_КРАТКОЕ отсутствует в шаблоне."""
    import zipfile
    with zipfile.ZipFile(CONTRACTS / "spec_foundation_install.docx") as z:
        for name in z.namelist():
            if name.endswith(".xml"):
                text = z.read(name).decode("utf-8")
                assert "ЗАКАЗЧИК_ДИРЕКТОР_ФИО_КРАТКОЕ" not in text, (
                    f"Устаревший плейсхолдер найден в {name}"
                )


def test_spec_appendix_has_initsialy():
    """Приложение №1 содержит {{ЗАКАЗЧИК_ДИРЕКТОР_ИНИЦИАЛЫ}} (минимум 2 вхождения в XML)."""
    import zipfile
    with zipfile.ZipFile(CONTRACTS / "spec_foundation_install.docx") as z:
        content = z.read("word/document.xml").decode("utf-8")
    assert content.count("ЗАКАЗЧИК_ДИРЕКТОР_ИНИЦИАЛЫ") >= 2, (
        "ЗАКАЗЧИК_ДИРЕКТОР_ИНИЦИАЛЫ должен встречаться минимум 2 раза: "
        "таблица подписей + Приложение №1"
    )


def test_spec_kwn_chain_until_signatures():
    """Цепочка keepNext непрерывна: tth_heading → пустые → kit_heading → подписи."""
    doc = Document(CONTRACTS / "spec_foundation_install.docx")
    body_paras = _body_paras(doc)

    tth_idx = next(
        i for i, p in enumerate(body_paras) if "Технические характеристики" in p.text
    )
    kit_idx = next(
        i for i, p in enumerate(body_paras) if "Комплект поставки" in p.text
    )

    # body_paras от tth_heading до пустого после kit_heading включительно.
    for i in range(tth_idx, kit_idx + 2):
        p = body_paras[i]
        assert _has_para_prop(p, "w:keepNext"), (
            f"body_paras[{i}] {p.text[:40]!r} должен иметь keepNext"
        )

    # body_paras[kit_idx+2] — пустой параграф перед таблицей подписей.
    p61 = body_paras[kit_idx + 2]
    assert _has_para_prop(p61, "w:keepNext"), (
        "Параграф перед таблицей подписей (body_paras[kit_idx+2]) должен иметь keepNext"
    )

    # Таблица подписей (doc.tables[3]) — cantSplit
    assert len(doc.tables) > 3, "doc.tables[3] (подписи) не найдена"
    for i, row in enumerate(doc.tables[3].rows):
        assert _row_has_cant_split(row), f"doc.tables[3] row[{i}] должна иметь cantSplit"


def test_spec_max_empty_paras_between_th_and_kompl():
    """Между таблицей ТХ и kit_heading не более 1 пустого параграфа."""
    doc = Document(CONTRACTS / "spec_foundation_install.docx")
    body_paras = _body_paras(doc)
    tth_idx = next(
        i for i, p in enumerate(body_paras) if "Технические характеристики" in p.text
    )
    kit_idx = next(
        i for i, p in enumerate(body_paras) if "Комплект поставки" in p.text
    )
    # tth_idx+2 .. kit_idx — параграфы после tth_heading до kit_heading.
    between = body_paras[tth_idx + 2 : kit_idx]
    assert len(between) <= 1, (
        f"Между таблицей ТХ и заголовком п.15 должно быть не более 1 пустого параграфа, "
        f"найдено {len(between)}"
    )


# --- плейсхолдер КРАТКОЕ_НАИМЕНОВАНИЕ ---

def _filled_text(template_name: str) -> str:
    """Генерирует документ с тестовой карточкой, возвращает текст всех параграфов и таблиц."""
    import tempfile
    import os
    from src.contracts.filler import fill_template

    data = {
        "ЗАКАЗЧИК_КРАТКОЕ_НАИМЕНОВАНИЕ": "МЕРИДИАН",
        "ЗАКАЗЧИК_ПОЛНОЕ_НАИМЕНОВАНИЕ": 'ООО «МЕРИДИАН»',
        "ЗАКАЗЧИК_ДИРЕКТОР_ДОЛЖНОСТЬ": "Генеральный директор",
        "ЗАКАЗЧИК_ДИРЕКТОР_ДОЛЖНОСТЬ_РП": "генерального директора",
        "ЗАКАЗЧИК_ДИРЕКТОР_ФИО": "Иванов Иван Иванович",
        "ЗАКАЗЧИК_ДИРЕКТОР_ФИО_РП": "Иванова Ивана Ивановича",
        "ЗАКАЗЧИК_ДИРЕКТОР_ИНИЦИАЛЫ": "И.И. Иванов",
        "ЗАКАЗЧИК_ДИРЕКТОР_ПОЛ": "male",
        "ЗАКАЗЧИК_ОСНОВАНИЕ": "Устава",
        "ЗАКАЗЧИК_ИНН": "1234567890",
        "ЗАКАЗЧИК_КПП": "123456789",
        "ЗАКАЗЧИК_ОГРН": "1234567890123",
        "ЗАКАЗЧИК_АДРЕС_ЮР": "г. Москва, ул. Тестовая, 1",
        "ЗАКАЗЧИК_АДРЕС_ПОЧТ": "г. Москва, ул. Тестовая, 1",
        "ЗАКАЗЧИК_РС": "40702810000000000001",
        "ЗАКАЗЧИК_БАНК": "Тестбанк",
        "ЗАКАЗЧИК_КС": "30101810000000000001",
        "ЗАКАЗЧИК_БИК": "044525001",
        "ЗАКАЗЧИК_ТЕЛЕФОН": "+7 (999) 000-00-00",
        "ЗАКАЗЧИК_EMAIL": "test@example.com",
        "ДОГОВОР_НОМЕР": "1",
        "ДОГОВОР_ДАТА_ПОЛНАЯ": "01 января 2026 г.",
        "ДИРЕКТОР_ПРИЧАСТИЕ": "действующего",
        "ИСПОЛНИТЕЛЬ_ДИРЕКТОР_ФИО_РП": "Сенаторова Олега Александровича",
        "СПЕЦ_НОМЕР": "1",
        "СПЕЦ_НАИМЕНОВАНИЕ": "Поставка весового оборудования",
        "СПЕЦ_ОБЩАЯ_СУММА": "1 000 000",
        "СПЕЦ_ОБЩАЯ_СУММА_ПРОПИСЬЮ": "Один миллион рублей 00 копеек",
        "СПЕЦ_ОПЛАТА_П1": "50% — 500 000 руб.",
        "СПЕЦ_ОПЛАТА_П2": "50% — 500 000 руб.",
        "СПЕЦ_ОПЛАТА_П3": "",
        "СПЕЦ_ОПЛАТА_П4": "",
        "СПЕЦ_ОПЛАТА_П5": "",
        "СПЕЦ_ОПЛАТА_П6": "",
    }

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        out = tmp.name

    try:
        fill_template(str(CONTRACTS / template_name), data, out)
        doc = Document(out)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        parts.append(p.text)
        return "\n".join(parts)
    finally:
        if os.path.exists(out):
            os.unlink(out)


def test_contract_preamble_and_signatures_use_short_name():
    """В преамбуле, реквизитах и подписях договора должно быть краткое наименование."""
    text = _filled_text("contract.docx")
    assert "МЕРИДИАН" in text, "Краткое наименование отсутствует в договоре"
    assert 'ООО «МЕРИДИАН»' not in text, "Полное наименование не должно использоваться в договоре"


def test_spec_v2_preamble_and_signatures_use_short_name():
    """В преамбуле и подписях spec_v2 должно быть краткое наименование."""
    text = _filled_text("spec_v2.docx")
    assert "МЕРИДИАН" in text, "Краткое наименование отсутствует в спецификации"
    assert 'ООО «МЕРИДИАН»' not in text, "Полное наименование не должно использоваться в спецификации"


# --- рендер-смоук договора поставки (docxtpl: преамбула в род. падеже) ---

def _supply_context():
    """Собрать полный docxtpl-контекст supply_contract.docx через реальный билдер."""
    from datetime import date

    from src.contracts.supply_filler import build_supply_context

    ctx = {
        "ЗАКАЗЧИК_КРАТКОЕ_НАИМЕНОВАНИЕ": "ООО «Меридиан»",
        "ЗАКАЗЧИК_ИНН": "7701234567",
        "ЗАКАЗЧИК_КПП": "770101001",
        "ЗАКАЗЧИК_ОГРН": "1027700000000",
        "ЗАКАЗЧИК_АДРЕС_ЮР": "г. Москва, ул. Тестовая, 1",
        "ЗАКАЗЧИК_РС": "40702810000000000001",
        "ЗАКАЗЧИК_КС": "30101810400000000225",
        "ЗАКАЗЧИК_БИК": "044525225",
        "ЗАКАЗЧИК_БАНК": "ПАО «Сбербанк России»",
        "ЗАКАЗЧИК_EMAIL": "test@test.ru",
        "ЗАКАЗЧИК_ТЕЛЕФОН": "+7 (473) 000-00-00",
        "ЗАКАЗЧИК_ДИРЕКТОР_ФИО": "Иванов Иван Иванович",
        "ЗАКАЗЧИК_ДИРЕКТОР_ДОЛЖНОСТЬ": "Генеральный директор",
        "ЗАКАЗЧИК_ОСНОВАНИЕ": "Устава",
    }
    items = [{
        "name": "Весы автомобильные ВЕСТА-С-60-20",
        "quantity": 1, "total": 2_000_000, "payment_group": "scales",
    }]
    return build_supply_context(
        ctx, items, deal={}, payment_rows=[],
        manual={"contract_number": "1", "object_address": "г. Воронеж"},
        contract_date=date(2026, 1, 1),
    )


def test_supply_contract_context_covers_all_placeholders():
    """Контекст build_supply_context покрывает ВСЕ плейсхолдеры шаблона.

    Прямая защита от UndefinedError: каждая переменная, на которую ссылается
    supply_contract.docx (вкл. новые ключи преамбулы в {% if %}), есть в контексте.
    """
    from docxtpl import DocxTemplate

    context = _supply_context()
    tpl = DocxTemplate(str(CONTRACTS / "supply_contract.docx"))
    declared = tpl.get_undeclared_template_variables()
    missing = declared - set(context)
    assert not missing, f"Шаблону не хватает ключей контекста: {sorted(missing)}"


def test_supply_contract_renders_preamble_no_raw_tags():
    """Полный compose_supply рендерит без сырых {{ }} / {% %} и со значениями преамбулы."""
    import os
    import tempfile
    from pathlib import Path

    from src.contracts.compose import compose_supply

    context = _supply_context()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        out = tmp.name
    try:
        compose_supply(context, Path(out))
        text = _all_text(Path(out))
    finally:
        if os.path.exists(out):
            os.unlink(out)

    assert "{%" not in text, "В выводе остались сырые jinja-теги {% %}"
    assert "{{" not in text, "В выводе остались сырые плейсхолдеры {{ }}"
    # Преамбула в родительном падеже + орг-форма собрались.
    assert "Иванова Ивана Ивановича" in text, "ФИО_РП не просклонилось"
    assert "генерального директора" in text, "ДОЛЖНОСТЬ_РП не просклонилась"
    assert "действующего" in text
    assert "именуемое" in text
    assert "Устава" in text
