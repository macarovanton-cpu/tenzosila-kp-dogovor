"""Тесты рефакторинга extractor.py — extract_card_data и legacy alias."""
from __future__ import annotations

import json
from unittest.mock import MagicMock, patch

import pytest


@pytest.fixture()
def mock_ai_client():
    with patch("src.contracts.extractor.OpenAI") as mock_cls:
        mock_client = MagicMock()
        mock_cls.return_value = mock_client
        yield mock_client


@pytest.fixture()
def mock_secret_env(monkeypatch):
    monkeypatch.setenv("OPENROUTER_API_KEY", "test-key")


@pytest.fixture()
def minimal_card_docx(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_paragraph("ООО Ромашка")
    doc.add_paragraph("ИНН 7701234567")
    path = tmp_path / "card.docx"
    doc.save(str(path))
    return str(path)


CARD_AI_RESPONSE = {
    "requisites": {
        "ЗАКАЗЧИК_КРАТКОЕ_НАИМЕНОВАНИЕ": "ООО Ромашка",
        "ЗАКАЗЧИК_ПОЛНОЕ_НАИМЕНОВАНИЕ": "Общество с ограниченной ответственностью «Ромашка»",
        "ЗАКАЗЧИК_ИНН": "7701234567",
        "ЗАКАЗЧИК_КПП": "770101001",
        "ЗАКАЗЧИК_ОГРН": "1234567890123",
        "ЗАКАЗЧИК_АДРЕС_ЮР": "г. Москва, ул. Ленина, 1",
        "ЗАКАЗЧИК_АДРЕС_ПОЧТ": "г. Москва, ул. Ленина, 1",
        "ЗАКАЗЧИК_РС": "40702810000000000001",
        "ЗАКАЗЧИК_БАНК": "Сбербанк",
        "ЗАКАЗЧИК_КС": "30101810400000000225",
        "ЗАКАЗЧИК_БИК": "044525225",
        "ЗАКАЗЧИК_ТЕЛЕФОН": "+7 495 123 45 67",
        "ЗАКАЗЧИК_EMAIL": "info@romashka.ru",
        "ЗАКАЗЧИК_ДИРЕКТОР_ДОЛЖНОСТЬ": "Генеральный директор",
        "ЗАКАЗЧИК_ДИРЕКТОР_ФИО": "Иванов Иван Иванович",
        "ЗАКАЗЧИК_ДИРЕКТОР_ДОЛЖНОСТЬ_РП": "генерального директора",
        "ЗАКАЗЧИК_ДИРЕКТОР_ФИО_РП": "Иванова Ивана Ивановича",
        "ЗАКАЗЧИК_ДИРЕКТОР_ИНИЦИАЛЫ": "И.И. Иванов",
        "ЗАКАЗЧИК_ОСНОВАНИЕ": "Устава",
    }
}


class TestExtractCardData:
    def test_returns_only_requisites_key(self, mock_ai_client, mock_secret_env, minimal_card_docx):
        mock_ai_client.chat.completions.create.return_value.choices = [
            MagicMock(message=MagicMock(content=json.dumps(CARD_AI_RESPONSE)))
        ]
        from src.contracts.extractor import extract_card_data
        result = extract_card_data(minimal_card_docx)
        assert "requisites" in result
        assert "specification" not in result

    def test_returns_19_requisite_fields(self, mock_ai_client, mock_secret_env, minimal_card_docx):
        mock_ai_client.chat.completions.create.return_value.choices = [
            MagicMock(message=MagicMock(content=json.dumps(CARD_AI_RESPONSE)))
        ]
        from src.contracts.extractor import extract_card_data
        result = extract_card_data(minimal_card_docx)
        assert len(result["requisites"]) == 19
        assert result["requisites"]["ЗАКАЗЧИК_ИНН"] == "7701234567"

    def test_uses_card_only_prompt(self, mock_ai_client, mock_secret_env, minimal_card_docx):
        """AI вызывается с промтом из extract_card_data.txt, не extract_contract_data.txt."""
        mock_ai_client.chat.completions.create.return_value.choices = [
            MagicMock(message=MagicMock(content=json.dumps(CARD_AI_RESPONSE)))
        ]
        from src.contracts.extractor import extract_card_data
        extract_card_data(minimal_card_docx)
        call_args = mock_ai_client.chat.completions.create.call_args
        system_prompt = call_args[1]["messages"][0]["content"]
        assert "СПЕЦ_П1" not in system_prompt
        assert "ЗАКАЗЧИК_ИНН" in system_prompt

    def test_card_data_accepts_pdf(self, mock_ai_client, mock_secret_env, tmp_path):
        """extract_card_data работает с PDF карточкой (через extract_pdf_text)."""
        mock_ai_client.chat.completions.create.return_value.choices = [
            MagicMock(message=MagicMock(content=json.dumps(CARD_AI_RESPONSE)))
        ]
        pdf_path = tmp_path / "card.pdf"
        pdf_path.write_bytes(b"dummy")
        card_text = "ООО Ромашка ИНН 7701234567 КПП 770101001 ОГРН 1234567890123 р/с 40702810000000000001"
        with patch("src.contracts.extractor.extract_pdf_text", return_value=card_text):
            from src.contracts.extractor import extract_card_data
            result = extract_card_data(str(pdf_path))
        assert "requisites" in result


class TestLegacyAlias:
    def test_extract_from_files_is_alias_for_legacy(self):
        """extract_from_files указывает на ту же функцию что extract_kp_data_legacy."""
        from src.contracts.extractor import extract_from_files, extract_kp_data_legacy
        assert extract_from_files is extract_kp_data_legacy


def _make_pdf(path, pages: list[str]) -> str:
    """PDF с текстовым слоем: одна строка на страницу (pymupdf)."""
    import fitz
    doc = fitz.open()
    for text in pages:
        page = doc.new_page()
        page.insert_text((72, 72), text)
    doc.save(str(path))
    doc.close()
    return str(path)


class TestEmptyInput:
    def test_scan_kp_raises_before_ai(self, mock_ai_client, mock_secret_env, minimal_card_docx, tmp_path):
        """Скан КП (пустой текстовый слой) → NoTextLayerError, LLM не вызывается."""
        from src.contracts.extractor import extract_kp_data_legacy
        from src.contracts.requisites_extract import NoTextLayerError
        kp_path = _make_pdf(tmp_path / "kp.pdf", [""])
        with pytest.raises(NoTextLayerError, match="КП"):
            extract_kp_data_legacy(kp_path, minimal_card_docx)
        mock_ai_client.chat.completions.create.assert_not_called()

    def test_scan_card_pdf_raises_before_ai(self, mock_ai_client, mock_secret_env, tmp_path):
        """Скан карточки (PDF без текста) → NoTextLayerError, LLM не вызывается."""
        from src.contracts.extractor import extract_card_data
        from src.contracts.requisites_extract import NoTextLayerError
        card_path = _make_pdf(tmp_path / "card.pdf", [""])
        with pytest.raises(NoTextLayerError, match="Карточка"):
            extract_card_data(card_path)
        mock_ai_client.chat.completions.create.assert_not_called()


class TestPromptIsCopyOnly:
    def test_prompt_forbids_normalization(self):
        """Промт режима B: LLM только копирует — без пересчёта и достройки (A4)."""
        from src.contracts.extractor import PROMPT_PATH
        prompt = PROMPT_PATH.read_text(encoding="utf-8")
        assert "ДОСЛОВНО" in prompt
        for forbidden in ("пересчитай", "восстанови по контексту", "прописью — формируй"):
            assert forbidden not in prompt, f"промт снова разрешает: {forbidden!r}"


class TestAllPagesRead:
    def test_legacy_reads_pages_beyond_3_5(self, mock_ai_client, mock_secret_env, minimal_card_docx, tmp_path):
        """Хардкод страниц 3–5 снят: текст 1-й и 6-й страниц КП доходит до LLM."""
        mock_ai_client.chat.completions.create.return_value.choices = [
            MagicMock(message=MagicMock(content=json.dumps(CARD_AI_RESPONSE)))
        ]
        filler = "Автомобильные весы ВЕСТА, коммерческое предложение, страница-наполнитель."
        kp_path = _make_pdf(
            tmp_path / "kp.pdf",
            ["MARKER_PAGE_1 " + filler, filler, filler, filler, filler, "MARKER_PAGE_6 " + filler],
        )
        from src.contracts.extractor import extract_kp_data_legacy
        extract_kp_data_legacy(kp_path, minimal_card_docx)
        user_message = mock_ai_client.chat.completions.create.call_args[1]["messages"][1]["content"]
        assert "MARKER_PAGE_1" in user_message
        assert "MARKER_PAGE_6" in user_message
